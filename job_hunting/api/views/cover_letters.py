import logging
import threading

from django.http import HttpResponse
from rest_framework import status
from rest_framework.decorators import action
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from drf_spectacular.utils import (
    extend_schema,
    extend_schema_view,
    OpenApiResponse,
)

from .base import BaseViewSet
from ._schema import _JSONAPI_ITEM, _JSONAPI_WRITE
from ..serializers import CoverLetterSerializer
from job_hunting.api.permissions import IsGuestReadOnly
from job_hunting.lib.ai_client import get_client
from job_hunting.lib.services.cover_letter_service import CoverLetterService
from job_hunting.lib.services.application_prompt_builder import ApplicationPromptBuilder
from job_hunting.lib.models import CareerData
from job_hunting.models import (
    Company,
    JobPost,
    CoverLetter,
    Resume,
)

logger = logging.getLogger(__name__)


@extend_schema_view(
    list=extend_schema(
        tags=["Cover Letters"], summary="List cover letters (authenticated user's only)"
    ),
    retrieve=extend_schema(
        tags=["Cover Letters"],
        summary="Retrieve a cover letter (owner only)",
        responses={
            200: _JSONAPI_ITEM,
            403: OpenApiResponse(description="Forbidden"),
            404: OpenApiResponse(description="Not found"),
        },
    ),
    update=extend_schema(
        tags=["Cover Letters"], summary="Update a cover letter (owner only)"
    ),
    partial_update=extend_schema(
        tags=["Cover Letters"], summary="Partially update a cover letter (owner only)"
    ),
    destroy=extend_schema(
        tags=["Cover Letters"], summary="Delete a cover letter (owner only)"
    ),
)
class CoverLetterViewSet(BaseViewSet):
    model = CoverLetter
    serializer_class = CoverLetterSerializer
    permission_classes = [IsAuthenticated, IsGuestReadOnly]

    def list(self, request):
        items = list(CoverLetter.objects.filter(user_id=request.user.id))
        items = self.paginate(items)
        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = CoverLetter.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Verify ownership
        if obj.user_id != request.user.id:
            return Response(
                {"errors": [{"detail": "Forbidden"}]},
                status=status.HTTP_403_FORBIDDEN,
            )

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def _upsert(self, request, pk, partial=False):
        obj = CoverLetter.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Verify ownership
        if obj.user_id != request.user.id:
            return Response(
                {"errors": [{"detail": "Forbidden"}]},
                status=status.HTTP_403_FORBIDDEN,
            )

        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)

        # Remove any user_id from attrs to prevent ownership changes
        attrs.pop("user_id", None)

        # If resume_id is being changed, validate new resume ownership
        if attrs.get("resume_id", None) is not None:
            new_resume = Resume.objects.filter(pk=attrs["resume_id"]).first()
            if not new_resume or new_resume.user_id != request.user.id:
                return Response(
                    {"errors": [{"detail": "Forbidden"}]},
                    status=status.HTTP_403_FORBIDDEN,
                )

        for k, v in attrs.items():
            setattr(obj, k, v)
        obj.save()
        return Response({"data": ser.to_resource(obj)})

    def destroy(self, request, pk=None):
        obj = CoverLetter.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response(status=204)

        # Verify ownership
        if obj.user_id != request.user.id:
            return Response(
                {"errors": [{"detail": "Forbidden"}]},
                status=status.HTTP_403_FORBIDDEN,
            )

        obj.delete()
        return Response(status=204)

    @extend_schema(
        tags=["Cover Letters"],
        summary="Create a cover letter (or AI-generate if content omitted — requires resume + job-post relationships)",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            202: OpenApiResponse(description="AI generation started — poll the returned resource for state changes"),
            400: OpenApiResponse(description="Missing/invalid resume or job-post"),
            403: OpenApiResponse(description="Resume not owned by user"),
            503: OpenApiResponse(description="AI client not configured"),
        },
    )
    def create(self, request):
        data = request.data if isinstance(request.data, dict) else {}
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(data)
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        node = data.get("data") or {}
        relationships = node.get("relationships") or {}
        attrs_node = node.get("attributes") or {}

        def _first_id(n):
            if isinstance(n, dict):
                d = n.get("data", n)
                if isinstance(d, dict) and "id" in d:
                    return d.get("id")
            return None

        def _rel_id(*keys):
            for k in keys:
                rel = relationships.get(k)
                if rel is not None:
                    rid = _first_id(rel)
                    if rid is not None:
                        return rid
            return None

        # Resolve IDs from attrs (serializer) or relationships fallbacks
        resume_id = attrs.get("resume_id") or _rel_id("resume", "resumes")
        job_post_id = attrs.get("job_post_id") or _rel_id(
            "job-post", "job_post", "jobPost", "job-posts", "jobPosts"
        )
        company_id = attrs.get("company_id") or _rel_id("company", "companies")

        # Additional fallbacks: accept convenience keys at attributes/top-level
        if job_post_id is None:
            job_post_id = (
                attrs_node.get("job_post_id")
                or attrs_node.get("jobPostId")
                or attrs_node.get("job-post-id")
                or node.get("job_post_id")
                or node.get("jobPostId")
                or node.get("job-post-id")
                or data.get("job_post_id")
                or data.get("jobPostId")
                or data.get("job-post-id")
            )
        if resume_id is None:
            resume_id = (
                attrs_node.get("resume_id")
                or attrs_node.get("resumeId")
                or attrs_node.get("resume-id")
                or node.get("resume_id")
                or node.get("resumeId")
                or node.get("resume-id")
                or data.get("resume_id")
                or data.get("resumeId")
                or data.get("resume-id")
            )
        if company_id is None:
            company_id = (
                attrs_node.get("company_id")
                or attrs_node.get("companyId")
                or attrs_node.get("company-id")
                or node.get("company_id")
                or node.get("companyId")
                or node.get("company-id")
                or data.get("company_id")
                or data.get("companyId")
                or data.get("company-id")
            )

        try:
            resume_id = int(resume_id) if resume_id is not None else None
            job_post_id = int(job_post_id) if job_post_id is not None else None
            company_id = int(company_id) if company_id is not None else None
        except (TypeError, ValueError):
            return Response(
                {"errors": [{"detail": "Invalid resume, job-post, or company ID"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        # id=0 means "no resume" → career-data fallback
        if resume_id == 0:
            resume_id = None

        resume = Resume.get(resume_id) if resume_id is not None else None
        job_post = (
            JobPost.objects.filter(pk=job_post_id).first()
            if job_post_id is not None
            else None
        )
        company = Company.get(company_id) if company_id is not None else None

        if job_post_id is not None and not job_post:
            return Response(
                {"errors": [{"detail": "Invalid job-post ID"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        if resume_id is not None and not resume:
            return Response(
                {"errors": [{"detail": "Invalid resume ID"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        if company_id is not None and not company:
            return Response(
                {"errors": [{"detail": "Invalid company ID"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Validate that the referenced resume belongs to the authenticated user
        if resume and resume.user_id != request.user.id:
            return Response(
                {
                    "errors": [
                        {
                            "detail": "Forbidden: resume does not belong to the authenticated user"
                        }
                    ]
                },
                status=status.HTTP_403_FORBIDDEN,
            )

        # Force ownership to authenticated user
        user_id = request.user.id

        # Extract injected prompt for AI generation
        injected_prompt = (
            attrs_node.get("instructions") or attrs_node.get("injected_prompt")
            or node.get("instructions") or data.get("instructions")
            or attrs.get("instructions") or attrs.get("injected_prompt")
        )
        injected_prompt = injected_prompt.strip() if isinstance(injected_prompt, str) else None
        injected_prompt = injected_prompt or None

        # Accept provided content; otherwise, generate via AI
        content = attrs_node.get("content")
        if content:
            cover_letter = CoverLetter(
                content=content,
                user_id=user_id,
                resume_id=(resume.id if resume else None),
                job_post_id=(job_post.id if job_post else None),
                company_id=(company.id if company else None),
            )
            cover_letter.save()
            payload = {"data": ser.to_resource(cover_letter)}
            include_rels = self._parse_include(request)
            if include_rels:
                payload["included"] = self._build_included(
                    [cover_letter], include_rels, request
                )
            return Response(payload, status=status.HTTP_201_CREATED)

        # AI generation path — create a pending record and dispatch async
        if job_post is None:
            return Response(
                {
                    "errors": [
                        {
                            "detail": "Provide 'relationships.job-post' when generating content without providing attributes.content"
                        }
                    ]
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Build career markdown when no resume provided (mirrors score career-data path)
        career_markdown = None
        if resume is None:
            career_data = CareerData.for_user(user_id)
            prompt_builder = ApplicationPromptBuilder(max_section_chars=60000)
            career_markdown = prompt_builder.build_from_career_data(career_data)
            if not career_markdown.strip():
                return Response(
                    {"errors": [{"detail": "No career data found for this user. Add favorite resumes or provide a resume relationship."}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

        client = get_client(required=False)
        if client is None:
            return Response(
                {"errors": [{"detail": "AI client not configured. Set OPENAI_API_KEY."}]},
                status=503,
            )

        cover_letter = CoverLetter.objects.create(
            user_id=user_id,
            resume_id=(resume.id if resume else None),
            job_post_id=job_post.id,
            company_id=(company.id if company else None),
            status="pending",
        )
        cl_id = cover_letter.id
        captured_injected_prompt = injected_prompt

        def _generate():
            import django
            django.db.close_old_connections()
            try:
                cl_service = CoverLetterService(
                    client, job_post, resume=resume, resume_markdown=career_markdown, user_id=user_id
                )
                result = cl_service.generate_cover_letter(injected_prompt=captured_injected_prompt)
                CoverLetter.objects.filter(pk=cl_id).update(
                    content=result.content, status="completed"
                )
            except Exception:
                CoverLetter.objects.filter(pk=cl_id).update(status="failed")

        threading.Thread(target=_generate, daemon=True).start()

        return Response({"data": ser.to_resource(cover_letter)}, status=status.HTTP_202_ACCEPTED)

    @extend_schema(
        tags=["Cover Letters"],
        summary="Export a cover letter as DOCX",
        responses={
            200: OpenApiResponse(description="DOCX file download"),
            403: OpenApiResponse(description="Forbidden"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    @action(
        detail=True,
        methods=["get"],
        url_path="export",
        permission_classes=[IsAuthenticated],
    )
    def export_docx(self, request, pk=None):
        cl = self.model.get(int(pk))
        if not cl:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Verify ownership
        if cl.user_id != request.user.id:
            return Response(
                {"errors": [{"detail": "Forbidden"}]},
                status=status.HTTP_403_FORBIDDEN,
            )

        try:
            from docx import Document  # python-docx
        except Exception:
            return Response(
                {
                    "errors": [
                        {"detail": "DOCX export requires 'python-docx' to be installed"}
                    ]
                },
                status=status.HTTP_501_NOT_IMPLEMENTED,
            )

        from io import BytesIO

        doc = Document()

        # Title
        title_parts = ["Cover Letter"]
        try:
            if getattr(cl, "job_post", None) and getattr(cl.job_post, "title", None):
                title_parts.append(str(cl.job_post.title))
            if getattr(cl, "job_post", None) and getattr(cl.job_post, "company", None):
                if getattr(cl.job_post.company, "name", None):
                    title_parts.append(str(cl.job_post.company.name))
        except Exception:
            pass
        if title_parts:
            doc.add_heading(" - ".join(title_parts), level=1)

        # Body
        content = (cl.content or "").strip()
        if content:
            import re as _re

            for para in [p for p in _re.split(r"\n\s*\n", content) if p.strip()]:
                p = doc.add_paragraph()
                for line in para.splitlines():
                    if p.text:
                        p.add_run("\n")
                    p.add_run(line)

        # Footer/meta
        try:
            meta_bits = []
            if getattr(cl, "created_at", None):
                meta_bits.append(f"Created: {cl.created_at}")
            if getattr(cl, "user", None) and getattr(cl.user, "name", None):
                meta_bits.append(f"Author: {cl.user.name}")
            if meta_bits:
                doc.add_paragraph("\n".join(meta_bits))
        except Exception:
            pass

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        import re as _re2

        filename_parts = ["cover-letter", str(cl.id)]
        try:
            if getattr(cl, "job_post", None) and getattr(cl.job_post, "company", None):
                nm = getattr(cl.job_post.company, "name", "") or ""
                if nm:
                    filename_parts.append(_re2.sub(r"[^A-Za-z0-9_-]+", "-", nm))
        except Exception:
            pass
        filename = "-".join([p for p in filename_parts if p]) + ".docx"

        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp

