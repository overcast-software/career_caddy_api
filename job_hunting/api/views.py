import dateparser
import logging
import math
import os
from django.db.models import Max, Q
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth import get_user_model
from django.conf import settings
from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.permissions import AllowAny, IsAuthenticated, IsAdminUser
from job_hunting.api.permissions import IsGuestReadOnly
from rest_framework.response import Response
from rest_framework.parsers import JSONParser, MultiPartParser
from drf_spectacular.utils import (
    extend_schema,
    extend_schema_view,
    OpenApiResponse,
    inline_serializer,
    OpenApiParameter,
)
from drf_spectacular.types import OpenApiTypes
from rest_framework import serializers as drf_serializers
from .parsers import VndApiJSONParser
from job_hunting.lib.scoring.job_scorer import JobScorer
from job_hunting.lib.ai_client import get_client, set_api_key
from job_hunting.lib.scraper import Scraper
from job_hunting.lib.services.summary_service import SummaryService
from job_hunting.lib.services.cover_letter_service import CoverLetterService
from job_hunting.lib.services.db_export_service import DbExportService
from job_hunting.lib.services.resume_export_service import ResumeExportService
from job_hunting.lib.services.ingest_resume import IngestResume
from job_hunting.lib.services.answer_service import AnswerService
from job_hunting.lib.services.application_prompt_builder import ApplicationPromptBuilder
from job_hunting.lib.models import (
    CareerData,
)
from job_hunting.models import (
    Status,
    Skill,
    Description,
    Certification,
    Education,
    Summary,
    Company,
    ApiKey,
    Question,
    JobPost,
    Answer,
    JobApplication,
    CoverLetter,
    Experience,
    Resume,
    Score,
    Scrape,
    ExperienceDescription,
    ResumeEducation,
    ResumeCertification,
    ResumeSummary,
    ResumeExperience,
    ResumeProject,
    ResumeSkill,
    JobApplicationStatus,
    Project,
    ProjectDescription,
    AiUsage,
    Waitlist,
    Invitation,
)
from .serializers import (
    ApiKeySerializer,
    DjangoUserSerializer,
    ResumeSerializer,
    ScoreSerializer,
    JobPostSerializer,
    ScrapeSerializer,
    CompanySerializer,
    CoverLetterSerializer,
    JobApplicationSerializer,
    SummarySerializer,
    ExperienceSerializer,
    EducationSerializer,
    CertificationSerializer,
    DescriptionSerializer,
    SkillSerializer,
    StatusSerializer,
    JobApplicationStatusSerializer,
    QuestionSerializer,
    AnswerSerializer,
    ProjectSerializer,
    AiUsageSerializer,
    WaitlistSerializer,
    InvitationSerializer,
    TYPE_TO_SERIALIZER,
    _parse_date,
    _resource_base_path,
)

logger = logging.getLogger(__name__)


@extend_schema(
    tags=["Auth"],
    summary="Get current user profile",
    responses={
        200: OpenApiResponse(
            description="JSON:API resource object for the authenticated user",
            response=inline_serializer(
                name="ProfileResponse",
                fields={"data": drf_serializers.DictField()},
            ),
        )
    },
)
@api_view(["GET"])
@permission_classes([IsAuthenticated])
def profile(request):
    """Get the current user's profile information."""
    ser = DjangoUserSerializer()
    resource = ser.to_resource(request.user)
    return Response({"data": resource})


@extend_schema(
    tags=["System"],
    summary="Health check",
    auth=[],
    responses={
        200: OpenApiResponse(
            description="Service is healthy",
            response=inline_serializer(
                name="HealthcheckResponse",
                fields={"healthy": drf_serializers.BooleanField()},
            ),
        )
    },
)
@csrf_exempt
def healthcheck(request):
    """Simple health check endpoint that only reports system health."""
    if request.method == "GET":
        User = get_user_model()
        user_count = User.objects.count()
        return JsonResponse({
            "healthy": True,
            "bootstrap_open": user_count == 0,
            "registration_open": settings.REGISTRATION_OPEN,
        })

    return JsonResponse({"error": "method not allowed"}, status=405)


@csrf_exempt
def guest_session(request):
    """Return a JWT for the guest user — no credentials required."""
    if request.method not in ('GET', 'POST'):
        return JsonResponse({'error': 'method not allowed'}, status=405)

    User = get_user_model()
    try:
        guest = User.objects.select_related('profile_obj').get(username='guest')
        if not guest.profile_obj.is_guest:
            return JsonResponse({'error': 'Guest account not configured.'}, status=400)
    except User.DoesNotExist:
        return JsonResponse({'error': 'Demo mode is not enabled on this server.'}, status=404)

    from rest_framework_simplejwt.tokens import RefreshToken
    refresh = RefreshToken.for_user(guest)
    return JsonResponse({
        'access': str(refresh.access_token),
        'refresh': str(refresh),
    })


@extend_schema(
    tags=["System"],
    summary="Check or perform first-time initialization",
    auth=[],
    methods=["GET"],
    responses={
        200: OpenApiResponse(
            description="Initialization status",
            response=inline_serializer(
                name="InitializeStatusResponse",
                fields={
                    "initialization_needed": drf_serializers.BooleanField(),
                    "status": drf_serializers.CharField(),
                },
            ),
        )
    },
)
@extend_schema(
    tags=["System"],
    summary="Create first superuser and configure the application",
    auth=[],
    methods=["POST"],
    request=inline_serializer(
        name="InitializeRequest",
        fields={
            "username": drf_serializers.CharField(required=False),
            "email": drf_serializers.EmailField(required=False),
            "password": drf_serializers.CharField(required=False),
            "first_name": drf_serializers.CharField(required=False),
            "last_name": drf_serializers.CharField(required=False),
            "openai_api_key": drf_serializers.CharField(required=False),
        },
    ),
    responses={
        201: OpenApiResponse(
            description="Initialized successfully",
            response=inline_serializer(
                name="InitializeResponse",
                fields={
                    "initialized": drf_serializers.BooleanField(),
                    "user": drf_serializers.DictField(),
                },
            ),
        ),
        400: OpenApiResponse(description="Bad request / user creation failed"),
        409: OpenApiResponse(description="Already initialized — superuser already exists"),
    },
)
@csrf_exempt
def initialize(request):
    """Initialize the application with first-time setup."""
    if request.method == "GET":
        # Check if initialization is needed
        try:
            User = get_user_model()
            user_count = User.objects.count()
        except Exception:
            # If Django ORM isn't initialized yet, initialization is needed
            user_count = None

        if user_count is None:
            status_str = "unknown"
            initialization_needed = True
        else:
            initialization_needed = user_count == 0
            status_str = (
                "initialized" if not initialization_needed else "needs_initialization"
            )

        return JsonResponse(
            {
                "initialization_needed": initialization_needed,
                "status": status_str,
            }
        )

    if request.method == "POST":
        # Handle initial setup
        try:
            User = get_user_model()
            user_count = User.objects.count()
        except Exception:
            user_count = None

        # Only allow initialization when no users exist
        if user_count is not None and user_count > 0:
            return JsonResponse(
                {"errors": [{"detail": "Application already initialized"}]}, status=409
            )

        # Parse request data
        try:
            import json

            data = json.loads(request.body.decode("utf-8") or "{}")
        except Exception:
            data = {}

        # Handle both JSON:API and plain JSON formats
        attrs = {}
        if isinstance(data.get("data"), dict):
            attrs = data["data"].get("attributes") or {}
        else:
            attrs = data or {}

        # Extract user creation parameters
        username = attrs.get("username") or attrs.get("name") or "admin"
        email = (attrs.get("email") or None) or None
        password = attrs.get("password") or "admin"
        first_name = attrs.get("first_name") or attrs.get("name") or ""
        last_name = attrs.get("last_name") or ""

        # Extract OpenAI API key
        api_key = (
            attrs.get("openai_api_key")
            or attrs.get("OPENAI_API_KEY")
            or attrs.get("openaiApiKey")
        )

        # Create superuser
        try:
            user = User.objects.create_superuser(
                username=username,
                email=email,
                password=password,
            )
            if first_name:
                user.first_name = first_name
            if last_name:
                user.last_name = last_name
            user.save()
        except Exception as e:
            return JsonResponse(
                {"errors": [{"detail": f"Failed to create user: {str(e)}"}]}, status=400
            )

        # Optionally set OpenAI API key
        meta = {}
        if api_key:
            try:
                set_api_key(str(api_key).strip())
                meta["openai_api_key_saved"] = True
            except Exception as e:
                meta["openai_api_key_saved"] = False
                meta["openai_api_key_error"] = str(e)

        # Return success response
        response_data = {
            "initialized": True,
            "user": {
                "id": user.id,
                "username": user.username,
                "email": user.email,
                "first_name": user.first_name,
                "last_name": user.last_name,
            },
        }
        if meta:
            response_data["meta"] = meta

        return JsonResponse(response_data, status=201)

    return JsonResponse({"error": "method not allowed"}, status=405)


@extend_schema(
    methods=["POST"],
    request=inline_serializer(
        name="WaitlistSignup",
        fields={
            "email": drf_serializers.EmailField(),
        },
    ),
    responses={
        201: OpenApiResponse(description="Added to waiting list"),
        400: OpenApiResponse(description="Missing or invalid email"),
        409: OpenApiResponse(description="Email already on the waiting list"),
    },
    auth=[],
)
@csrf_exempt
def waitlist_signup(request):
    """Public endpoint for joining the waiting list."""
    if request.method == "POST":
        import json
        from job_hunting.models import Waitlist

        try:
            data = json.loads(request.body.decode("utf-8") or "{}")
        except Exception:
            data = {}

        email = (data.get("email") or "").strip().lower()
        if not email or "@" not in email:
            return JsonResponse(
                {"errors": [{"detail": "A valid email address is required."}]},
                status=400,
            )

        if Waitlist.objects.filter(email=email).exists():
            return JsonResponse(
                {"errors": [{"detail": "This email is already on the waiting list."}]},
                status=409,
            )

        Waitlist.objects.create(email=email)

        try:
            from django.core.mail import send_mail
            from django.template.loader import render_to_string

            body = render_to_string(
                "waitlist_email.txt",
                {"frontend_url": settings.FRONTEND_URL},
            )
            send_mail(
                subject="You're on the Career Caddy waiting list",
                message=body,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[email],
            )
        except Exception:
            logger.warning("Failed to send waitlist confirmation to %s", email)

        _notify_admins_new_signup(email, email, method="waitlist")

        return JsonResponse({"success": True, "email": email}, status=201)

    return JsonResponse({"error": "method not allowed"}, status=405)


@csrf_exempt
def password_reset_request(request):
    """Request a password-reset email. Always returns 200 to prevent email enumeration."""
    if request.method != "POST":
        return JsonResponse({"error": "method not allowed"}, status=405)

    import json
    from django.contrib.auth.tokens import PasswordResetTokenGenerator
    from django.utils.http import urlsafe_base64_encode
    from django.utils.encoding import force_bytes
    from django.core.mail import send_mail
    from django.template.loader import render_to_string

    try:
        data = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        data = {}

    email = (data.get("email") or "").strip().lower()
    if not email or "@" not in email:
        return JsonResponse(
            {"errors": [{"detail": "A valid email address is required."}]},
            status=400,
        )

    User = get_user_model()
    user = User.objects.filter(email__iexact=email).first()

    if user:
        token_generator = PasswordResetTokenGenerator()
        token = token_generator.make_token(user)
        uid = urlsafe_base64_encode(force_bytes(user.pk))
        reset_url = f"{settings.FRONTEND_URL}/reset-password?token={token}&uid={uid}"

        body = render_to_string(
            "password_reset_email.txt", {"reset_url": reset_url}
        )
        send_mail(
            subject="Password Reset — Career Caddy",
            message=body,
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=[user.email],
        )

    return JsonResponse(
        {
            "message": (
                "If an account with that email exists, "
                "a password reset link has been sent."
            )
        },
        status=200,
    )


@csrf_exempt
def password_reset_confirm(request):
    """Confirm a password reset with token, uid, and new password."""
    if request.method != "POST":
        return JsonResponse({"error": "method not allowed"}, status=405)

    import json
    from django.contrib.auth.tokens import PasswordResetTokenGenerator
    from django.contrib.auth.password_validation import validate_password
    from django.core.exceptions import ValidationError
    from django.utils.http import urlsafe_base64_decode
    from django.utils.encoding import force_str

    try:
        data = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        data = {}

    token = data.get("token", "")
    uid = data.get("uid", "")
    new_password = data.get("new_password", "")

    if not token or not uid or not new_password:
        return JsonResponse(
            {"errors": [{"detail": "token, uid, and new_password are required."}]},
            status=400,
        )

    User = get_user_model()
    try:
        user_id = force_str(urlsafe_base64_decode(uid))
        user = User.objects.get(pk=user_id)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        return JsonResponse(
            {"errors": [{"detail": "Invalid or expired reset link."}]},
            status=400,
        )

    token_generator = PasswordResetTokenGenerator()
    if not token_generator.check_token(user, token):
        return JsonResponse(
            {"errors": [{"detail": "Invalid or expired reset link."}]},
            status=400,
        )

    try:
        validate_password(new_password, user)
    except ValidationError as e:
        return JsonResponse(
            {"errors": [{"detail": msg} for msg in e.messages]},
            status=400,
        )

    user.set_password(new_password)
    user.save()

    return JsonResponse(
        {"message": "Password has been reset successfully."}, status=200
    )


_INCLUDE_PARAM = OpenApiParameter(
    "include",
    OpenApiTypes.STR,
    OpenApiParameter.QUERY,
    required=False,
    description="Comma-separated relationships to sideload (e.g. resumes,api-keys).",
)

_PAGE_PARAMS = [
    OpenApiParameter(
        "page[number]",
        OpenApiTypes.INT,
        OpenApiParameter.QUERY,
        required=False,
        description="Page number (1-based). Default: 1.",
    ),
    OpenApiParameter(
        "page[size]",
        OpenApiTypes.INT,
        OpenApiParameter.QUERY,
        required=False,
        description="Items per page. Default: 50.",
    ),
    _INCLUDE_PARAM,
]

_SORT_PARAM = OpenApiParameter(
    "sort",
    OpenApiTypes.STR,
    OpenApiParameter.QUERY,
    required=False,
    description="Comma-separated sort fields. Prefix with '-' for descending (e.g., '-created_at' for newest first).",
)

_FILTER_QUERY_PARAM = OpenApiParameter(
    "filter[query]",
    OpenApiTypes.STR,
    OpenApiParameter.QUERY,
    required=False,
    description="Search across title, description, company name, and company display_name (case-insensitive OR).",
)
_FILTER_COMPANY_PARAM = OpenApiParameter(
    "filter[company]",
    OpenApiTypes.STR,
    OpenApiParameter.QUERY,
    required=False,
    description="Filter by company name (case-insensitive contains).",
)
_FILTER_COMPANY_ID_PARAM = OpenApiParameter(
    "filter[company_id]",
    OpenApiTypes.INT,
    OpenApiParameter.QUERY,
    required=False,
    description="Filter by exact company ID.",
)
_FILTER_TITLE_PARAM = OpenApiParameter(
    "filter[title]",
    OpenApiTypes.STR,
    OpenApiParameter.QUERY,
    required=False,
    description="Filter by job post title (case-insensitive contains).",
)

_FILTER_APP_QUERY_PARAM = OpenApiParameter(
    "filter[query]",
    OpenApiTypes.STR,
    OpenApiParameter.QUERY,
    required=False,
    description="Search across job post title, company name, company display_name, status, and notes (case-insensitive OR).",
)
_FILTER_APP_STATUS_PARAM = OpenApiParameter(
    "filter[status]",
    OpenApiTypes.STR,
    OpenApiParameter.QUERY,
    required=False,
    description="Filter by application status (case-insensitive contains).",
)

_JSONAPI_LIST = OpenApiResponse(
    description="JSON:API list",
    response=inline_serializer(
        name="JsonApiList",
        fields={
            "data": drf_serializers.ListField(child=drf_serializers.DictField()),
            "included": drf_serializers.ListField(
                child=drf_serializers.DictField(), required=False
            ),
        },
    ),
)
_JSONAPI_ITEM = OpenApiResponse(
    description="JSON:API resource",
    response=inline_serializer(
        name="JsonApiItem",
        fields={
            "data": drf_serializers.DictField(),
            "included": drf_serializers.ListField(
                child=drf_serializers.DictField(), required=False
            ),
        },
    ),
)
_JSONAPI_WRITE = inline_serializer(
    name="JsonApiWrite",
    fields={
        "data": drf_serializers.DictField(
            help_text="JSON:API resource object with 'type', 'attributes', and optional 'relationships'."
        )
    },
)


class BaseViewSet(viewsets.ViewSet):
    permission_classes = [IsAuthenticated, IsGuestReadOnly]
    parser_classes = [MultiPartParser, VndApiJSONParser, JSONParser]
    model = None
    serializer_class = None

    def get_permissions(self):
        # Always allow OPTIONS for CORS preflight and API metadata
        if getattr(self.request, "method", "").upper() == "OPTIONS":
            return [AllowAny()]
        return super().get_permissions()

    def options(self, request, *args, **kwargs):
        # Explicitly handle CORS preflight to avoid auth and ensure proper headers
        allow_methods = "GET, POST, PUT, PATCH, DELETE, OPTIONS"
        origin = request.META.get("HTTP_ORIGIN")
        requested_headers = request.META.get("HTTP_ACCESS_CONTROL_REQUEST_HEADERS")

        resp = Response(status=200)
        resp["Allow"] = allow_methods
        resp["Access-Control-Allow-Methods"] = allow_methods
        if origin:
            resp["Access-Control-Allow-Origin"] = origin
            resp["Vary"] = "Origin"
            # If frontend sends credentials (cookies/Authorization), allow them
            resp["Access-Control-Allow-Credentials"] = "true"
        else:
            resp["Access-Control-Allow-Origin"] = "*"

        if requested_headers:
            resp["Access-Control-Allow-Headers"] = requested_headers
        else:
            resp["Access-Control-Allow-Headers"] = "Authorization, Content-Type"

        resp["Access-Control-Max-Age"] = "600"
        return resp

    def get_throttles(self):
        # Disable DRF throttling in development-like environments
        try:
            debug = bool(getattr(settings, "DEBUG", False))
        except Exception:
            debug = False
        env_val = os.environ.get("ENV", "").lower()
        disable_flag = str(os.environ.get("DISABLE_THROTTLE", "")).strip().lower() in (
            "1",
            "true",
            "yes",
        )
        settings_disable = bool(getattr(settings, "DISABLE_THROTTLE", False))
        if (
            debug
            or env_val in ("dev", "development", "local")
            or disable_flag
            or settings_disable
        ):
            return []
        return super().get_throttles()

    def _get_obj(self, pk):
        """Fetch a single object by PK."""
        return self.model.objects.filter(pk=int(pk)).first()

    def get_serializer(self, *args, slim=False, **kwargs):
        ser = self.serializer_class()
        ser.slim = slim
        return ser

    def _is_slim_request(self, request) -> bool:
        return bool(request.query_params.get("slim"))

    def pre_save_payload(self, request, attrs: dict, creating: bool) -> dict:
        """Hook for subclasses to adjust/force attributes before persistence."""
        return attrs

    def _parse_include(self, request):
        raw = []
        inc = request.query_params.get("include")
        incs = request.query_params.get("includes")
        if inc:
            raw.append(str(inc))
        if incs and incs != inc:
            raw.append(str(incs))
        if not raw:
            return []
        parts = []
        for chunk in raw:
            parts.extend([s.strip() for s in chunk.split(",") if s and s.strip()])
        # de-duplicate while preserving order
        seen = set()
        out = []
        for p in parts:
            if p not in seen:
                seen.add(p)
                out.append(p)
        return out

    def _normalize_rel_for_serializer(self, name: str, serializer) -> str:
        """
        Normalize relationship names to match serializer relationship keys.
        Frontend convention: dasherized singular model names (e.g., 'job-application')
        Serializer convention: dasherized plural relationship keys (e.g., 'job-applications')
        """
        rels = getattr(serializer, "relationships", {}) or {}
        rel_keys = set(rels.keys())

        # Direct match - return immediately if found
        if name in rel_keys:
            return name

        # Convert to dasherized form if not already
        dasherized = name.replace("_", "-")
        if dasherized in rel_keys:
            return dasherized

        # Static mapping from frontend model names (singular) to backend relationship keys
        MODEL_TO_RELATIONSHIP = {
            # Plural relationships (to-many)
            "answer": "answers",
            "api-key": "api-keys",
            "certification": "certifications",
            "cover-letter": "cover-letters",
            "description": "descriptions",
            "education": "educations",
            "experience": "experiences",
            "job-application": "applications",
            "job-post": "job-posts",
            "question": "questions",
            "resume": "resumes",
            "score": "scores",
            "scrape": "scrapes",
            "skill": "skills",
            "status": "statuses",
            "summary": "summaries",
            "user": "users",
            # Singular relationships (to-one) - map to themselves
            "company": "company",
            "career-data": "career-data",
        }

        # Look up in mapping - this is the primary normalization path
        mapped = MODEL_TO_RELATIONSHIP.get(dasherized)
        if mapped and mapped in rel_keys:
            return mapped

        # Fallback: check if the relationship type matches the requested name
        for rel_key, cfg in rels.items():
            rel_type = (cfg or {}).get("type")
            if rel_type == name or rel_type == dasherized:
                return rel_key

        # No match found - return the mapped value if we have one, otherwise original
        if mapped:
            return mapped
        return name

    def _build_included(
        self, objs, include_rels, request=None, primary_serializer=None
    ):
        included = []
        seen = set()  # (type, id)
        primary_ser = primary_serializer or self.get_serializer()

        def _include_recursive(
            objects,
            path_segments,
            current_serializer,
            parent_type=None,
            parent_id=None,
            parent_rel=None,
        ):
            if not path_segments or not objects:
                return

            segment = path_segments[0]
            remaining_segments = path_segments[1:]

            for obj in objects:
                normalized_rel = self._normalize_rel_for_serializer(
                    segment, current_serializer
                )

                # Get relationship config first
                cfg = getattr(current_serializer, "relationships", {}).get(
                    normalized_rel
                )

                rel_type, targets = current_serializer.get_related(obj, normalized_rel)
                effective_type = rel_type or (cfg and cfg.get("type"))

                # Only continue if we have no effective type and no config
                if not effective_type and not cfg:
                    continue

                # FK fallback for to-one relationships when targets is empty
                if not targets and cfg and not cfg.get("uselist", True):
                    fk_field = getattr(current_serializer, "relationship_fks", {}).get(
                        normalized_rel
                    )
                    if fk_field:
                        rel_id = getattr(obj, fk_field, None)
                        if rel_id is not None:
                            effective_type = effective_type or cfg.get("type")
                            ser_cls = TYPE_TO_SERIALIZER.get(effective_type)
                            if ser_cls:
                                rel_ser = ser_cls()
                                model_cls = rel_ser.model
                                try:
                                    fetched = model_cls.objects.filter(
                                        pk=int(rel_id)
                                    ).first()
                                    if fetched:
                                        targets = [fetched]
                                except (TypeError, ValueError, AttributeError):
                                    pass

                # Recompute effective_type if still None and cfg exists
                effective_type = effective_type or (cfg and cfg.get("type"))

                # Only proceed if we have an effective type and targets
                if not effective_type or not targets:
                    continue

                ser_cls = TYPE_TO_SERIALIZER.get(effective_type)
                if not ser_cls:
                    continue

                rel_ser = ser_cls()
                rel_ser.request = request
                # Provide parent context so serializers can customize included resources
                if hasattr(rel_ser, "set_parent_context"):
                    rel_ser.set_parent_context(
                        current_serializer.type, obj.id, normalized_rel
                    )

                for t in targets:
                    key = (effective_type, str(t.id))
                    already_seen = key in seen

                    if not already_seen:
                        # Filter user-owned resources to only include those owned by authenticated user
                        if (
                            effective_type in ("cover-letter", "score", "summary", "job-application")
                            and request
                            and hasattr(request, "user")
                            and request.user.is_authenticated
                        ):
                            if getattr(t, "user_id", None) != request.user.id:
                                continue

                        seen.add(key)
                        included.append(rel_ser.to_resource(t))

                    # If there are more segments, always recurse (even if this node was already seen)
                    if remaining_segments:
                        _include_recursive([t], remaining_segments, rel_ser)

                    # Auto-include children of experience (descriptions and company) when path ends at experience
                    if effective_type == "experience" and not remaining_segments:
                        exp_child_ser = ExperienceSerializer()
                        if hasattr(exp_child_ser, "set_parent_context"):
                            exp_child_ser.set_parent_context("experience", t.id, None)
                        for child_rel in ("descriptions", "company"):
                            _include_recursive([t], [child_rel], exp_child_ser)

                    # Auto-include descriptions for projects
                    if effective_type == "project" and not remaining_segments:
                        proj_child_ser = ProjectSerializer()
                        _include_recursive([t], ["descriptions"], proj_child_ser)

        # Process each include path
        for include_path in include_rels:
            path_segments = include_path.split(".")
            _include_recursive(objs, path_segments, primary_ser)

        return included

    def _page_params(self):
        """Return (page_number, page_size) parsed from request, supporting both
        page[number]/page[size] (JSON:API) and page/per_page (simple) styles."""
        qp = self.request.query_params
        try:
            page_number = int(qp.get("page[number]") or qp.get("page") or 1)
        except Exception:
            page_number = 1
        try:
            page_size = int(qp.get("page[size]") or qp.get("per_page") or 50)
        except Exception:
            page_size = 50
        page_number = max(1, page_number)
        page_size = max(1, min(page_size, 200))
        return page_number, page_size

    def paginate(self, items):
        page_number, page_size = self._page_params()
        start = (page_number - 1) * page_size
        end = start + page_size
        return items[start:end]

    @extend_schema(
        tags=["API"],
        summary="List resources",
        parameters=_PAGE_PARAMS,
        responses={200: _JSONAPI_LIST},
    )
    def list(self, request):
        items = list(self.model.objects.all())
        items = self.paginate(items)
        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["API"],
        summary="Retrieve a resource",
        parameters=[_INCLUDE_PARAM],
        responses={200: _JSONAPI_ITEM, 404: OpenApiResponse(description="Not found")},
    )
    def retrieve(self, request, pk=None):
        obj = self._get_obj(pk)
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["API"],
        summary="Create a resource",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Validation error"),
        },
    )
    def create(self, request):
        # Handle both JSON:API and plain JSON payloads
        data = request.data if isinstance(request.data, dict) else {}
        ser = self.get_serializer()

        # Check if this is JSON:API format (has "data" wrapper)
        if "data" in data:
            # Use JSON:API parser
            try:
                attrs = ser.parse_payload(request.data)
            except ValueError as e:
                return Response({"errors": [{"detail": str(e)}]}, status=400)
        else:
            # Handle plain JSON format
            attrs = {}
            for key in [
                "username",
                "email",
                "first_name",
                "last_name",
                "password",
                "phone",
            ]:
                if key in data:
                    attrs[key] = data[key]

        attrs = self.pre_save_payload(request, attrs, creating=True)
        # Respect model schema: only set created_by_id if the model supports it
        if hasattr(self.model, "created_by_id"):
            # Ignore any client-supplied created_by_id; set from authenticated user if available
            attrs.pop("created_by_id", None)
            if getattr(request, "user", None) and getattr(
                request.user, "is_authenticated", False
            ):
                attrs["created_by_id"] = request.user.id
        else:
            # Ensure unsupported ownership fields are not passed to the model
            attrs.pop("created_by_id", None)
            attrs.pop("created_by", None)
        obj = self.model.objects.create(**attrs)
        return Response({"data": ser.to_resource(obj)}, status=status.HTTP_201_CREATED)

    @extend_schema(
        tags=["API"],
        summary="Replace a resource (full update)",
        request=_JSONAPI_WRITE,
        responses={
            200: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Validation error"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    def update(self, request, pk=None):
        return self._upsert(request, pk, partial=False)

    @extend_schema(
        tags=["API"],
        summary="Partially update a resource",
        request=_JSONAPI_WRITE,
        responses={
            200: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Validation error"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    def partial_update(self, request, pk=None):
        return self._upsert(request, pk, partial=True)

    def _upsert(self, request, pk, partial=False):
        obj = self._get_obj(pk)
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)
        attrs = self.pre_save_payload(request, attrs, creating=False)
        # Do not allow changing ownership
        attrs.pop("created_by_id", None)
        for k, v in attrs.items():
            setattr(obj, k, v)
        obj.save()
        return Response({"data": ser.to_resource(obj)})

    @extend_schema(
        tags=["API"],
        summary="Delete a resource",
        responses={204: OpenApiResponse(description="Deleted")},
    )
    def destroy(self, request, pk=None):
        self.model.objects.filter(pk=int(pk)).delete()
        return Response(status=204)

    # JSON:API relationships linkage endpoint:
    # GET /<type>/{id}/relationships/<rel-name>
    @extend_schema(
        tags=["API"],
        summary="Get JSON:API relationship linkage data",
        responses={
            200: OpenApiResponse(description="Relationship linkage"),
            404: OpenApiResponse(description="Not found or relationship not found"),
        },
    )
    @action(detail=True, methods=["get"], url_path=r"relationships/(?P<rel>[^/]+)")
    def relationships(self, request, pk=None, rel=None):
        obj = self._get_obj(pk)
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()

        # Normalize the relationship name
        normalized_rel = self._normalize_rel_for_serializer(rel, ser)

        cfg = (
            ser.relationships.get(normalized_rel)
            if ser and hasattr(ser, "relationships")
            else None
        )
        if not cfg:
            return Response(
                {"errors": [{"detail": "Relationship not found"}]}, status=404
            )

        rel_type = cfg["type"]
        uselist = cfg.get("uselist", True)
        target = getattr(obj, cfg["attr"], None)

        if uselist:
            if hasattr(target, "all"):
                target = target.all()
            data = [{"type": rel_type, "id": str(i.id)} for i in (target or [])]
            links = {
                "self": f"{_resource_base_path(ser.type)}/{obj.id}/relationships/{rel}",
            }
        else:
            # Try to get target_id from object or FK fallback
            target_id = None
            if target is not None and getattr(target, "id", None) is not None:
                target_id = target.id
            else:
                # FK fallback: check if we have a foreign key field for this relationship
                fk_field = getattr(ser, "relationship_fks", {}).get(normalized_rel)
                if fk_field:
                    fk_value = getattr(obj, fk_field, None)
                    if fk_value is not None:
                        target_id = fk_value

            data = (
                {"type": rel_type, "id": str(target_id)}
                if target_id is not None
                else None
            )
            links = {
                "self": f"{_resource_base_path(ser.type)}/{obj.id}/relationships/{rel}",
            }
            # Include related link if we have a target_id
            if target_id is not None:
                links["related"] = f"{_resource_base_path(rel_type)}/{target_id}"

        return Response({"data": data, "links": links})


@extend_schema_view(
    list=extend_schema(tags=["Summaries"], summary="List summaries"),
    retrieve=extend_schema(tags=["Summaries"], summary="Retrieve a summary"),
    update=extend_schema(tags=["Summaries"], summary="Update a summary"),
    partial_update=extend_schema(
        tags=["Summaries"], summary="Partially update a summary"
    ),
    destroy=extend_schema(tags=["Summaries"], summary="Delete a summary"),
)
class SummaryViewSet(BaseViewSet):
    model = Summary
    serializer_class = SummarySerializer



    def list(self, request):
        qs = Summary.objects.filter(user_id=request.user.id)

        query_filter = request.query_params.get("filter[query]")
        if query_filter:
            qs = qs.filter(content__icontains=query_filter)

        qs = qs.order_by("-id")
        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs.all()[offset: offset + page_size])

        ser = self.get_serializer()
        payload = {
            "data": [ser.to_resource(obj) for obj in items],
            "meta": {"total": total, "page": page_number, "per_page": page_size, "total_pages": total_pages},
        }
        if page_number < total_pages:
            base = request.build_absolute_uri(request.path)
            qp = request.query_params.dict()
            qp["page"] = page_number + 1
            qp["per_page"] = page_size
            payload["links"] = {"next": base + "?" + "&".join(f"{k}={v}" for k, v in qp.items())}
        else:
            payload["links"] = {"next": None}
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = Summary.objects.filter(pk=pk).first()
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        return Response({"data": ser.to_resource(obj)})

    def update(self, request, pk=None):
        return self._update_summary(request, pk)

    def partial_update(self, request, pk=None):
        return self._update_summary(request, pk)

    def _update_summary(self, request, pk):
        obj = Summary.objects.filter(pk=pk).first()
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs = node.get("attributes") or {}
        relationships = node.get("relationships") or {}

        # Update Summary fields
        if "content" in attrs:
            obj.content = attrs["content"]
        if "status" in attrs:
            obj.status = attrs["status"]
        obj.save()

        # Update ResumeSummary.active if provided
        if "active" in attrs:
            active_val = bool(attrs["active"])

            # Resolve resume_id: from relationship or flat attribute
            resume_id = attrs.get("resume_id")
            if resume_id is None:
                resume_rel = relationships.get("resume") or relationships.get("resumes")
                if isinstance(resume_rel, dict):
                    rel_data = resume_rel.get("data")
                    if isinstance(rel_data, dict):
                        resume_id = rel_data.get("id")

            if resume_id is not None:
                try:
                    resume_id = int(resume_id)
                except (TypeError, ValueError):
                    resume_id = None

            # Fall back to the single linked resume if unambiguous
            if resume_id is None:
                linked = list(ResumeSummary.objects.filter(summary_id=obj.id).values_list("resume_id", flat=True))
                if len(linked) == 1:
                    resume_id = linked[0]

            if resume_id is not None:
                ResumeSummary.objects.get_or_create(resume_id=resume_id, summary_id=obj.id)
                if active_val:
                    ResumeSummary.objects.filter(resume_id=resume_id).update(active=False)
                    ResumeSummary.objects.filter(resume_id=resume_id, summary_id=obj.id).update(active=True)
                else:
                    ResumeSummary.objects.filter(resume_id=resume_id, summary_id=obj.id).update(active=False)

        ser = self.get_serializer()
        return Response({"data": ser.to_resource(obj)})

    def destroy(self, request, pk=None):
        obj = Summary.objects.filter(pk=pk).first()
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        obj.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

    @extend_schema(
        tags=["Summaries"],
        summary="Create a summary (or AI-generate if content omitted and job-post provided)",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Missing resume or invalid IDs"),
            503: OpenApiResponse(description="AI client not configured"),
        },
    )
    def create(self, request):
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs = node.get("attributes") or {}
        relationships = node.get("relationships") or {}

        def _first_id(node):
            if isinstance(node, dict):
                d = node.get("data")
            else:
                d = None
            if isinstance(d, dict) and "id" in d:
                return d["id"]
            if isinstance(d, list) and d:
                first = d[0]
                if isinstance(first, dict) and "id" in first:
                    return first["id"]
            return None

        def _rel_id(*keys):
            for k in keys:
                v = relationships.get(k)
                if v is not None:
                    rid = _first_id(v)
                    if rid is not None:
                        return rid
            return None

        # Accept both hyphenated and underscored keys; allow nullable job_post
        resume_id = _rel_id("resume", "resumes")
        job_post_id = _rel_id(
            "job-post", "job_post", "jobPost", "job-posts", "jobPosts"
        )
        user_id = _rel_id("user", "users")

        # Resume is optional — omitting it or passing id=0 falls back to career-data
        resume = None
        if resume_id is not None:
            try:
                rid = int(resume_id)
            except (TypeError, ValueError):
                rid = None
            if rid:  # 0 → treated as "no resume" → career-data fallback
                resume = Resume.objects.filter(pk=rid).first()
                if not resume:
                    return Response(
                        {"errors": [{"detail": "Invalid resume ID"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

        job_post = None
        if job_post_id is not None:
            try:
                job_post = JobPost.objects.filter(pk=int(job_post_id)).first()
            except (TypeError, ValueError):
                pass
            if not job_post:
                return Response(
                    {"errors": [{"detail": "Invalid job-post ID"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

        user_id = request.user.id

        content = attrs.get("content")

        if content:
            summary = Summary.objects.create(
                job_post_id=job_post.id if job_post else None,
                user_id=user_id,
                content=content,
            )
        else:
            if not job_post:
                return Response(
                    {
                        "errors": [
                            {
                                "detail": "Provide 'attributes.content' or a job-post relationship to generate content"
                            }
                        ]
                    },
                    status=status.HTTP_400_BAD_REQUEST,
                )

            client = get_client(required=False)
            if client is None:
                return Response(
                    {
                        "errors": [
                            {"detail": "AI client not configured. Set OPENAI_API_KEY."}
                        ]
                    },
                    status=503,
                )

            if resume is None:
                # No resume — use the user's full career data
                career_data = CareerData.for_user(user_id)
                prompt_builder = ApplicationPromptBuilder(max_section_chars=60000)
                career_markdown = prompt_builder.build_from_career_data(career_data)
                if not career_markdown.strip():
                    return Response(
                        {"errors": [{"detail": "No career data found for this user. Add favorite resumes or provide a resume relationship."}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )
                summary_service = SummaryService(client, job=job_post, resume_markdown=career_markdown, user_id=user_id)
            else:
                summary_service = SummaryService(client, job=job_post, resume=resume)

            summary = summary_service.generate_summary()

        # Link to resume when one was provided
        if resume is not None:
            ResumeSummary.objects.filter(resume_id=resume.id).update(active=False)
            ResumeSummary.objects.get_or_create(
                resume_id=resume.id, summary_id=summary.id, defaults={"active": True}
            )
            ResumeSummary.objects.filter(resume_id=resume.id, summary_id=summary.id).update(
                active=True
            )
            ResumeSummary.ensure_single_active_for_resume(resume.id)

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(summary)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([summary], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)


_USER_RESOURCE_RESPONSE = OpenApiResponse(
    description="JSON:API user resource",
    response=inline_serializer(
        name="UserResource",
        fields={"data": drf_serializers.DictField()},
    ),
)

_USER_LIST_RESPONSE = OpenApiResponse(
    description="JSON:API list of user resources",
    response=inline_serializer(
        name="UserListResource",
        fields={
            "data": drf_serializers.ListField(child=drf_serializers.DictField()),
            "included": drf_serializers.ListField(
                child=drf_serializers.DictField(), required=False
            ),
        },
    ),
)

_USER_WRITE_REQUEST = inline_serializer(
    name="UserWriteRequest",
    fields={
        "username": drf_serializers.CharField(required=False),
        "email": drf_serializers.EmailField(required=False),
        "password": drf_serializers.CharField(required=False),
        "first_name": drf_serializers.CharField(required=False),
        "last_name": drf_serializers.CharField(required=False),
        "phone": drf_serializers.CharField(required=False),
    },
)


class DjangoUserViewSet(viewsets.ViewSet):
    permission_classes = [IsAuthenticated]
    parser_classes = [VndApiJSONParser, JSONParser]

    def get_permissions(self):
        """Allow unauthenticated access for create and bootstrap_superuser actions."""
        if self.action in ["create", "bootstrap_superuser"]:
            return [AllowAny()]
        return super().get_permissions()

    def options(self, request, *args, **kwargs):
        # Explicitly handle CORS preflight to avoid auth and ensure proper headers
        allow_methods = "GET, POST, PUT, PATCH, DELETE, OPTIONS"
        origin = request.META.get("HTTP_ORIGIN")
        requested_headers = request.META.get("HTTP_ACCESS_CONTROL_REQUEST_HEADERS")

        resp = Response(status=200)
        resp["Allow"] = allow_methods
        resp["Access-Control-Allow-Methods"] = allow_methods
        if origin:
            resp["Access-Control-Allow-Origin"] = origin
            resp["Vary"] = "Origin"
            resp["Access-Control-Allow-Credentials"] = "true"
        else:
            resp["Access-Control-Allow-Origin"] = "*"

        if requested_headers:
            resp["Access-Control-Allow-Headers"] = requested_headers
        else:
            resp["Access-Control-Allow-Headers"] = "Authorization, Content-Type"

        resp["Access-Control-Max-Age"] = "600"
        return resp

    def get_throttles(self):
        # Disable DRF throttling in development-like environments
        try:
            debug = bool(getattr(settings, "DEBUG", False))
        except Exception:
            debug = False
        env_val = os.environ.get("ENV", "").lower()
        disable_flag = str(os.environ.get("DISABLE_THROTTLE", "")).strip().lower() in (
            "1",
            "true",
            "yes",
        )
        settings_disable = bool(getattr(settings, "DISABLE_THROTTLE", False))
        if (
            debug
            or env_val in ("dev", "development", "local")
            or disable_flag
            or settings_disable
        ):
            return []
        return super().get_throttles()

    def get_serializer(self, *args, **kwargs):
        return DjangoUserSerializer()

    def _parse_include(self, request):
        raw = []
        inc = request.query_params.get("include")
        incs = request.query_params.get("includes")
        if inc:
            raw.append(str(inc))
        if incs and incs != inc:
            raw.append(str(incs))
        if not raw:
            return []
        parts = []
        for chunk in raw:
            parts.extend([s.strip() for s in chunk.split(",") if s and s.strip()])
        # de-duplicate while preserving order
        seen = set()
        out = []
        for p in parts:
            if p not in seen:
                seen.add(p)
                out.append(p)
        return out

    def _build_included(self, objs, include_rels):
        included = []
        seen = set()  # (type, id)
        primary_ser = self.get_serializer()

        for obj in objs:
            for rel in include_rels:
                rel_type, targets = primary_ser.get_related(obj, rel)
                if not rel_type:
                    continue
                ser_cls = TYPE_TO_SERIALIZER.get(rel_type)
                if not ser_cls:
                    continue
                rel_ser = ser_cls()
                for t in targets:
                    key = (rel_type, str(t.id))
                    if key in seen:
                        continue
                    seen.add(key)
                    included.append(rel_ser.to_resource(t))
        return included

    @extend_schema(
        tags=["Users"],
        summary="List users (staff sees all; others see only themselves)",
        parameters=[_INCLUDE_PARAM],
        responses={200: _USER_LIST_RESPONSE},
    )
    def list(self, request):
        User = get_user_model()

        # Restrict list to staff users or return only current user
        if request.user.is_staff:
            users = User.objects.all()
        else:
            users = [request.user]

        ser = self.get_serializer()
        data = [ser.to_resource(u) for u in users]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(users, include_rels)
        return Response(payload)

    @extend_schema(
        tags=["Users"],
        summary="Retrieve a user by ID",
        parameters=[_INCLUDE_PARAM],
        responses={
            200: _USER_RESOURCE_RESPONSE,
            403: OpenApiResponse(description="Forbidden"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    def retrieve(self, request, pk=None):
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
        except (User.DoesNotExist, ValueError):
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Only allow staff to retrieve other users
        if not request.user.is_staff and user.id != request.user.id:
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(user)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([user], include_rels)
        return Response(payload)

    @extend_schema(
        tags=["Users"],
        summary="Store the OpenAI API key (superuser only)",
        request=inline_serializer(
            name="SetOpenAIKeyRequest",
            fields={"openai_api_key": drf_serializers.CharField()},
        ),
        responses={
            201: OpenApiResponse(
                description="Key saved",
                response=inline_serializer(
                    name="SetOpenAIKeyResponse",
                    fields={"meta": drf_serializers.DictField()},
                ),
            ),
            400: OpenApiResponse(description="Missing or invalid key"),
            403: OpenApiResponse(description="Forbidden — superuser only"),
        },
    )
    @action(
        detail=False,
        methods=["post"],
        url_path="set-openai-api-key",
        permission_classes=[IsAuthenticated],
    )
    def set_openai_api_key(self, request):
        user = request.user
        if not getattr(user, "is_superuser", False):
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)

        # Accept API key via header or JSON/JSON:API body
        api_key = (request.META.get("HTTP_X_OPENAI_API_KEY", "") or "").strip()
        if not api_key:
            data = request.data if isinstance(request.data, dict) else {}
            attrs = {}
            if isinstance(data.get("data"), dict):
                attrs = data["data"].get("attributes") or {}
            else:
                attrs = data or {}
            api_key = (
                attrs.get("openai_api_key")
                or attrs.get("OPENAI_API_KEY")
                or attrs.get("openaiApiKey")
                or ""
            )
            api_key = str(api_key).strip()

        if not api_key:
            return Response(
                {"errors": [{"detail": "Missing openai_api_key"}]}, status=400
            )

        try:
            set_api_key(api_key)
        except Exception as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)

        return Response({"meta": {"openai_api_key_saved": True}}, status=201)

    @extend_schema(
        tags=["Auth"],
        summary="Register a new user",
        auth=[],
        request=_USER_WRITE_REQUEST,
        responses={
            201: _USER_RESOURCE_RESPONSE,
            400: OpenApiResponse(
                description="Validation error (missing fields, duplicate username/email)"
            ),
        },
    )
    def create(self, request):
        # Staff users can always create accounts; public registration requires REGISTRATION_OPEN
        is_staff = request.user and request.user.is_authenticated and request.user.is_staff
        if not is_staff and not settings.REGISTRATION_OPEN:
            return Response(
                {"errors": [{"detail": "Registration is currently closed."}]},
                status=403,
            )

        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)

        # Extract profile fields before user creation
        profile_fields = {}
        for pf in ("phone", "linkedin", "github", "address", "links"):
            if pf in attrs:
                profile_fields[pf] = attrs.pop(pf)

        username = attrs.get("username")
        password = attrs.get("password")
        email = attrs.get("email", "")
        first_name = attrs.get("first_name", "")
        last_name = attrs.get("last_name", "")

        if not username:
            return Response(
                {"errors": [{"detail": "Username is required"}]}, status=400
            )
        if not password:
            return Response(
                {"errors": [{"detail": "Password is required"}]}, status=400
            )

        User = get_user_model()

        # Check uniqueness
        if User.objects.filter(username=username).exists():
            return Response(
                {"errors": [{"detail": "Username already exists"}]}, status=400
            )
        if email and User.objects.filter(email=email).exists():
            return Response(
                {"errors": [{"detail": "Email already exists"}]}, status=400
            )

        from django.contrib.auth.password_validation import validate_password
        from django.core.exceptions import ValidationError
        try:
            validate_password(password)
        except ValidationError as e:
            return Response(
                {"errors": [{"detail": msg} for msg in e.messages]}, status=400
            )

        user = User(
            username=username, email=email, first_name=first_name, last_name=last_name
        )
        user.set_password(password)
        user.save()

        # Handle profile fields via Django Profile if any provided
        if profile_fields:
            from job_hunting.models import Profile

            prof, _ = Profile.objects.get_or_create(user_id=user.id)
            if "phone" in profile_fields:
                val = str(profile_fields["phone"] or "").strip()
                prof.phone = (val[:50] or None) if val else None
            if "linkedin" in profile_fields:
                val = str(profile_fields["linkedin"] or "").strip()
                prof.linkedin = (val[:255] or None) if val else None
            if "github" in profile_fields:
                val = str(profile_fields["github"] or "").strip()
                prof.github = (val[:255] or None) if val else None
            if "address" in profile_fields:
                prof.address = (str(profile_fields["address"] or "").strip() or None)
            if "links" in profile_fields:
                prof.links = profile_fields["links"] if isinstance(profile_fields["links"], dict) else {}
            prof.save()

        # Send welcome email (non-blocking — don't fail registration on email error)
        if email:
            try:
                from django.core.mail import send_mail
                from django.template.loader import render_to_string

                login_url = f"{settings.FRONTEND_URL}/login"
                body = render_to_string(
                    "welcome_email.txt",
                    {"username": username, "login_url": login_url},
                )
                send_mail(
                    subject="Welcome to Career Caddy",
                    message=body,
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[email],
                )
            except Exception:
                logger.warning("Failed to send welcome email to %s", email)

        return Response({"data": ser.to_resource(user)}, status=status.HTTP_201_CREATED)

    @extend_schema(
        tags=["Users"],
        summary="Replace a user (full update)",
        request=_USER_WRITE_REQUEST,
        responses={
            200: _USER_RESOURCE_RESPONSE,
            400: OpenApiResponse(description="Validation error"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    def update(self, request, pk=None):
        return self._upsert(request, pk, partial=False)

    @extend_schema(
        tags=["Users"],
        summary="Partially update a user",
        request=_USER_WRITE_REQUEST,
        responses={
            200: _USER_RESOURCE_RESPONSE,
            400: OpenApiResponse(description="Validation error"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    def partial_update(self, request, pk=None):
        return self._upsert(request, pk, partial=True)

    def _upsert(self, request, pk, partial=False):
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
        except (User.DoesNotExist, ValueError):
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)

        # Extract profile fields before user updates
        profile_fields = {}
        for pf in ("phone", "linkedin", "github", "address", "links"):
            if pf in attrs:
                profile_fields[pf] = attrs.pop(pf)

        # Update allowed fields
        if "email" in attrs:
            user.email = attrs["email"]
        if "first_name" in attrs:
            user.first_name = attrs["first_name"]
        if "last_name" in attrs:
            user.last_name = attrs["last_name"]
        if "password" in attrs:
            user.set_password(attrs["password"])

        # Staff-only fields
        if "is_staff" in attrs:
            if not request.user.is_staff:
                return Response({"errors": [{"detail": "Forbidden"}]}, status=403)
            user.is_staff = bool(attrs["is_staff"])
        if "is_active" in attrs:
            if not request.user.is_staff:
                return Response({"errors": [{"detail": "Forbidden"}]}, status=403)
            user.is_active = bool(attrs["is_active"])

        user.save()

        # Handle profile fields via Django Profile if any provided
        if profile_fields:
            from job_hunting.models import Profile

            prof, _ = Profile.objects.get_or_create(user_id=user.id)
            if "phone" in profile_fields:
                val = str(profile_fields["phone"] or "").strip()
                prof.phone = (val[:50] or None) if val else None
            if "linkedin" in profile_fields:
                val = str(profile_fields["linkedin"] or "").strip()
                prof.linkedin = (val[:255] or None) if val else None
            if "github" in profile_fields:
                val = str(profile_fields["github"] or "").strip()
                prof.github = (val[:255] or None) if val else None
            if "address" in profile_fields:
                prof.address = (str(profile_fields["address"] or "").strip() or None)
            if "links" in profile_fields:
                prof.links = profile_fields["links"] if isinstance(profile_fields["links"], dict) else {}
            prof.save()

        return Response({"data": ser.to_resource(user)})

    @extend_schema(
        tags=["Users"],
        summary="Delete a user (staff only)",
        responses={
            204: OpenApiResponse(description="Deleted"),
            403: OpenApiResponse(description="Forbidden — staff only"),
        },
    )
    def destroy(self, request, pk=None):
        if not request.user.is_staff:
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
            user.delete()
        except (User.DoesNotExist, ValueError):
            pass
        return Response(status=204)

    @extend_schema(
        tags=["Auth"],
        summary="[Deprecated] Bootstrap superuser — use /api/v1/initialize/ instead",
        auth=[],
        deprecated=True,
        responses={410: OpenApiResponse(description="Gone — endpoint removed")},
    )
    @action(
        detail=False,
        methods=["post"],
        url_path="bootstrap-superuser",
        permission_classes=[AllowAny],
    )
    def bootstrap_superuser(self, request):
        """Deprecated: Use /api/v1/initialize/ instead."""
        return Response(
            {
                "errors": [
                    {
                        "detail": "This endpoint is deprecated. Use /api/v1/initialize/ instead."
                    }
                ]
            },
            status=410,  # Gone
        )

    @extend_schema(
        tags=["Users"],
        summary="List resumes for a user",
        responses={
            200: OpenApiResponse(description="JSON:API list of resume resources")
        },
    )
    @action(detail=True, methods=["get"])
    def resumes(self, request, pk=None):
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
        except (User.DoesNotExist, ValueError):
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        resumes = list(Resume.objects.filter(user_id=user.id))
        data = [ResumeSerializer().to_resource(r) for r in resumes]
        return Response({"data": data})

    @extend_schema(
        tags=["Users"],
        summary="List scores for a user",
        responses={
            200: OpenApiResponse(description="JSON:API list of score resources")
        },
    )
    @action(detail=True, methods=["get"])
    def scores(self, request, pk=None):
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
        except (User.DoesNotExist, ValueError):
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        scores = list(Score.objects.filter(user_id=user.id))
        data = [ScoreSerializer().to_resource(s) for s in scores]
        return Response({"data": data})

    @extend_schema(
        tags=["Users"],
        summary="List cover letters for a user",
        responses={
            200: OpenApiResponse(description="JSON:API list of cover-letter resources")
        },
    )
    @action(detail=True, methods=["get"], url_path="cover-letters")
    def cover_letters(self, request, pk=None):
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
        except (User.DoesNotExist, ValueError):
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        cover_letters = list(CoverLetter.objects.filter(user_id=user.id))
        data = [CoverLetterSerializer().to_resource(c) for c in cover_letters]
        return Response({"data": data})

    @extend_schema(
        tags=["Users"],
        summary="List job applications for a user",
        responses={
            200: OpenApiResponse(description="JSON:API list of application resources")
        },
    )
    @action(detail=True, methods=["get"], url_path="job-applications")
    def applications(self, request, pk=None):
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
        except (User.DoesNotExist, ValueError):
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        applications = list(JobApplication.objects.filter(user_id=user.id))
        data = [JobApplicationSerializer().to_resource(a) for a in applications]
        return Response({"data": data})

    @extend_schema(
        tags=["Users"],
        summary="List summaries for a user",
        responses={
            200: OpenApiResponse(description="JSON:API list of summary resources")
        },
    )
    @action(detail=True, methods=["get"])
    def summaries(self, request, pk=None):
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
        except (User.DoesNotExist, ValueError):
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        summaries = list(Summary.objects.filter(user_id=user.id))
        data = [SummarySerializer().to_resource(s) for s in summaries]
        return Response({"data": data})

    @extend_schema(
        tags=["Users"],
        summary="List API keys for a user (own keys, or any if staff)",
        responses={
            200: OpenApiResponse(description="JSON:API list of api-key resources"),
            403: OpenApiResponse(description="Forbidden"),
        },
    )
    @action(detail=True, methods=["get"], url_path="api-keys")
    def api_keys(self, request, pk=None):
        """Get API keys for a user"""
        User = get_user_model()
        try:
            user = User.objects.get(id=int(pk))
        except (User.DoesNotExist, ValueError):
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Only allow users to see their own API keys or staff to see any
        if not request.user.is_staff and user.id != request.user.id:
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)

        api_keys = ApiKey.objects.filter(user_id=user.id)
        data = [ApiKeySerializer().to_resource(k) for k in api_keys]
        return Response({"data": data})

    @extend_schema(
        tags=["Auth"],
        summary="Get the authenticated user's own record",
        parameters=[_INCLUDE_PARAM],
        responses={200: _USER_RESOURCE_RESPONSE},
    )
    @action(
        detail=False,
        methods=["get"],
        url_path="me",
        permission_classes=[IsAuthenticated],
    )
    def me(self, request):
        user = request.user
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(user)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([user], include_rels)
        return Response(payload)


@extend_schema_view(
    create=extend_schema(tags=["Resumes"], summary="Create a resume"),
    update=extend_schema(
        tags=["Resumes"],
        summary="Update a resume (supports nested experiences/educations/skills reconciliation)",
    ),
    partial_update=extend_schema(
        tags=["Resumes"],
        summary="Partially update a resume (supports nested experiences/educations/skills reconciliation)",
    ),
    destroy=extend_schema(tags=["Resumes"], summary="Delete a resume"),
)
class ResumeViewSet(BaseViewSet):
    model = Resume
    serializer_class = ResumeSerializer

    @extend_schema(
        tags=["Resumes"],
        summary="List resumes (auto-includes all relationships)",
        parameters=_PAGE_PARAMS,
        responses={200: _JSONAPI_LIST},
    )
    def list(self, request):
        items = list(self.model.objects.filter(user_id=request.user.id))
        items = self.paginate(items)
        slim = self._is_slim_request(request)
        ser = self.get_serializer(slim=slim)
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        if not slim:
            include_rels = self._parse_include(request)
            if include_rels:
                payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["Resumes"],
        summary="Retrieve a resume (auto-includes all relationships)",
        parameters=[_INCLUDE_PARAM],
        responses={200: _JSONAPI_ITEM, 404: OpenApiResponse(description="Not found")},
    )
    def retrieve(self, request, pk=None):
        obj = self.model.objects.filter(pk=int(pk)).first()
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def _upsert(self, request, pk, partial=False):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)

        # Update scalar attributes
        for k, v in attrs.items():
            setattr(obj, k, v)
        obj.save()

        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs_node = node.get("attributes") or {}
        # Merge relationships from both the standard location and the nested "data" block
        # (some clients send a nested data.data.relationships for sub-relationships)
        nested_rels = (node.get("data") or {}).get("relationships") or {}
        rels_node = {**(node.get("relationships") or {}), **nested_rels}

        # Optional: update active summary content if attributes.summary is provided
        incoming_summary = attrs_node.get("summary")
        if isinstance(incoming_summary, str):
            new_content = incoming_summary.strip()
            # Find active link
            active_link = ResumeSummary.objects.filter(
                resume_id=obj.id, active=True
            ).first()
            if active_link:
                sm = Summary.objects.filter(pk=active_link.summary_id).first()
                if sm and (sm.content or "") != new_content:
                    sm.content = new_content
                    sm.save()
                # Ensure only this one is active
                ResumeSummary.objects.filter(resume_id=obj.id).exclude(
                    pk=active_link.id
                ).update(active=False)
                active_link.active = True
                active_link.save()
            else:
                # No active summary; create one and activate it
                sm = Summary.objects.create(
                    job_post_id=None,
                    user_id=getattr(obj, "user_id", None),
                    content=new_content,
                )
                ResumeSummary.objects.filter(resume_id=obj.id).update(active=False)
                ResumeSummary.objects.create(
                    resume_id=obj.id, summary_id=sm.id, active=True
                )

        # Helpers for relationship reconciliation
        def _int_or_none(v):
            try:
                return int(v) if v is not None else None
            except (TypeError, ValueError):
                return None

        def _dp(val):
            if not val:
                return None
            try:
                dt = dateparser.parse(str(val))
                return dt.date() if dt else None
            except Exception:
                return None

        # Experiences: if present, PATCH to match provided set and update Experience attributes/links
        experiences_in = node.get("experiences") or data.get("experiences")
        if experiences_in is not None:
            desired_exp_ids = []
            for wrapper in experiences_in or []:
                exp_node = (wrapper or {}).get("data") or wrapper or {}
                exp_id = _int_or_none(exp_node.get("id"))
                if not exp_id:
                    return Response(
                        {"errors": [{"detail": "Experience id is required in PATCH"}]},
                        status=400,
                    )
                exp = Experience.objects.filter(pk=exp_id).first()
                if not exp:
                    return Response(
                        {"errors": [{"detail": f"Invalid experience id: {exp_id}"}]},
                        status=400,
                    )

                # Update experience attributes
                exp_attrs = exp_node.get("attributes") or {}
                if "title" in exp_attrs:
                    exp.title = exp_attrs.get("title")
                if "location" in exp_attrs:
                    exp.location = exp_attrs.get("location")
                if "summary" in exp_attrs:
                    exp.summary = exp_attrs.get("summary")
                if "start_date" in exp_attrs:
                    exp.start_date = _dp(exp_attrs.get("start_date"))
                if "end_date" in exp_attrs:
                    exp.end_date = _dp(exp_attrs.get("end_date"))

                # Update company relation (optional)
                exp_rels = exp_node.get("relationships") or {}
                comp_rel = (exp_rels.get("company") or {}).get("data") or {}
                comp_id = _int_or_none(comp_rel.get("id"))
                if comp_rel:
                    exp.company_id = comp_id

                exp.save()

                # Ensure resume <-> experience link exists
                ResumeExperience.objects.get_or_create(
                    resume_id=obj.id, experience_id=exp.id
                )

                # Reconcile descriptions for this experience (keep order as provided)
                desc_nodes = (exp_rels.get("descriptions") or {}).get("data") or []
                desired_desc_ids_ordered = []
                invalid_desc_ids = []
                for d in desc_nodes:
                    did = _int_or_none((d or {}).get("id"))
                    if did is None:
                        continue
                    if not Description.objects.filter(pk=did).exists():
                        invalid_desc_ids.append(did)
                        continue
                    desired_desc_ids_ordered.append(did)
                if invalid_desc_ids:
                    return Response(
                        {
                            "errors": [
                                {
                                    "detail": f"Invalid description ID(s): {', '.join(map(str, invalid_desc_ids))}"
                                }
                            ]
                        },
                        status=400,
                    )

                if desc_nodes is not None:
                    existing_links = list(
                        ExperienceDescription.objects.filter(experience_id=exp.id)
                    )
                    existing_by_desc = {
                        lnk.description_id: lnk for lnk in existing_links
                    }
                    desired_set = set(desired_desc_ids_ordered)
                    existing_set = set(existing_by_desc.keys())

                    # Remove links not desired
                    to_remove = existing_set - desired_set
                    if to_remove:
                        ExperienceDescription.objects.filter(
                            experience_id=exp.id,
                            description_id__in=list(to_remove),
                        ).delete()

                    # Add/update desired links with order
                    for order_idx, did in enumerate(desired_desc_ids_ordered):
                        link = existing_by_desc.get(did)
                        if not link:
                            ExperienceDescription.objects.create(
                                experience_id=exp.id,
                                description_id=did,
                                order=order_idx,
                            )
                        else:
                            link.order = order_idx
                            link.save()

                desired_exp_ids.append(exp.id)

            # Reconcile resume_experience set to match provided experiences
            existing_links = list(ResumeExperience.objects.filter(resume_id=obj.id))
            existing_ids = {lnk.experience_id for lnk in existing_links}
            desired_ids = set(desired_exp_ids)
            to_add = desired_ids - existing_ids
            to_remove = existing_ids - desired_ids

            for eid in to_add:
                ResumeExperience.objects.get_or_create(
                    resume_id=obj.id, experience_id=eid
                )
            if to_remove:
                ResumeExperience.objects.filter(
                    resume_id=obj.id,
                    experience_id__in=list(to_remove),
                ).delete()

        # Educations: reconcile join set if provided
        educations_in = node.get("educations") or data.get("educations")
        if educations_in is not None:
            desired_ids = set()
            invalid = []
            for item in educations_in or []:
                ed_node = (item or {}).get("data") or item or {}
                eid = _int_or_none(ed_node.get("id"))
                if eid is None:
                    continue
                if not Education.objects.filter(pk=eid).exists():
                    invalid.append(eid)
                else:
                    desired_ids.add(eid)
            if invalid:
                return Response(
                    {
                        "errors": [
                            {
                                "detail": f"Invalid education ID(s): {', '.join(map(str, invalid))}"
                            }
                        ]
                    },
                    status=400,
                )
            existing_ids = set(
                ResumeEducation.objects.filter(resume_id=obj.id).values_list(
                    "education_id", flat=True
                )
            )
            to_add = desired_ids - existing_ids
            to_remove = existing_ids - desired_ids
            for eid in to_add:
                ResumeEducation.objects.create(resume_id=obj.id, education_id=eid)
            if to_remove:
                ResumeEducation.objects.filter(
                    resume_id=obj.id,
                    education_id__in=list(to_remove),
                ).delete()

        # Certifications: reconcile join set if provided
        certifications_in = node.get("certifications") or data.get("certifications")
        if certifications_in is not None:
            desired_ids = set()
            invalid = []
            for item in certifications_in or []:
                c_node = (item or {}).get("data") or item or {}
                cid = _int_or_none(c_node.get("id"))
                if cid is None:
                    continue
                if not Certification.objects.filter(pk=cid).exists():
                    invalid.append(cid)
                else:
                    desired_ids.add(cid)
            if invalid:
                return Response(
                    {
                        "errors": [
                            {
                                "detail": f"Invalid certification ID(s): {', '.join(map(str, invalid))}"
                            }
                        ]
                    },
                    status=400,
                )
            existing_ids = set(
                ResumeCertification.objects.filter(resume_id=obj.id).values_list(
                    "certification_id", flat=True
                )
            )
            to_add = desired_ids - existing_ids
            to_remove = existing_ids - desired_ids
            for cid in to_add:
                ResumeCertification.objects.create(
                    resume_id=obj.id, certification_id=cid
                )
            if to_remove:
                ResumeCertification.objects.filter(
                    resume_id=obj.id,
                    certification_id__in=list(to_remove),
                ).delete()

        # Skills: reconcile join set if provided
        skills_in = node.get("skills") or data.get("skills")
        if skills_in is not None:
            desired_active_by_id = {}
            invalid = []
            for item in skills_in or []:
                if not isinstance(item, dict):
                    continue
                s_node = (item.get("data") or item) or {}
                sid = _int_or_none(s_node.get("id"))
                skill = None
                if sid is not None:
                    skill = Skill.objects.filter(pk=sid).first()
                    if not skill:
                        invalid.append(sid)
                        continue
                else:
                    s_attrs = s_node.get("attributes") or {}
                    text = s_attrs.get("text") or s_node.get("text")
                    if not text:
                        continue
                    # Create or find by text
                    skill, _ = Skill.objects.get_or_create(text=str(text).strip())

                # Determine desired active flag (default True)
                active_val = (s_node.get("attributes") or {}).get("active")
                active_val = bool(active_val) if active_val is not None else True
                desired_active_by_id[skill.id] = active_val

            if invalid:
                return Response(
                    {
                        "errors": [
                            {
                                "detail": f"Invalid skill ID(s): {', '.join(map(str, invalid))}"
                            }
                        ]
                    },
                    status=400,
                )

            existing_links = list(ResumeSkill.objects.filter(resume_id=obj.id))
            existing_ids = {lnk.skill_id for lnk in existing_links}
            desired_ids = set(desired_active_by_id.keys())

            to_add = desired_ids - existing_ids
            to_remove = existing_ids - desired_ids
            to_update = desired_ids & existing_ids

            # Remove undesired links
            if to_remove:
                ResumeSkill.objects.filter(
                    resume_id=obj.id,
                    skill_id__in=list(to_remove),
                ).delete()

            # Add missing links
            for sid in to_add:
                ResumeSkill.objects.get_or_create(
                    resume_id=obj.id,
                    skill_id=sid,
                    defaults={"active": desired_active_by_id[sid]},
                )

            # Update 'active' where needed
            for link in existing_links:
                if link.skill_id in to_update:
                    desired_active = desired_active_by_id[link.skill_id]
                    if bool(link.active) != bool(desired_active):
                        link.active = bool(desired_active)
                        link.save()

        # Summaries: reconcile join set if provided
        # Accept from node["summaries"], data["summaries"], or relationships.summaries.data
        _rel_summaries = (rels_node.get("summaries") or {}).get("data")
        summaries_in = node.get("summaries") or data.get("summaries") or _rel_summaries
        if summaries_in is not None:
            desired_ids_ordered = []  # preserve order — last item becomes active
            invalid = []
            desired_active_sid = None  # explicit active=true flag wins
            for item in summaries_in or []:
                s_node = (item or {}).get("data") or item or {}
                sid = _int_or_none(s_node.get("id"))
                if sid is None:
                    continue
                if not Summary.objects.filter(pk=sid).exists():
                    invalid.append(sid)
                else:
                    if sid not in desired_ids_ordered:
                        desired_ids_ordered.append(sid)
                    active_flag = (s_node.get("attributes") or {}).get("active")
                    if active_flag:
                        desired_active_sid = sid
            if invalid:
                return Response(
                    {"errors": [{"detail": f"Invalid summary ID(s): {', '.join(map(str, invalid))}"}]},
                    status=400,
                )

            desired_ids = set(desired_ids_ordered)
            existing_links = list(ResumeSummary.objects.filter(resume_id=obj.id))
            existing_ids = {lnk.summary_id for lnk in existing_links}

            to_add = desired_ids - existing_ids
            to_remove = existing_ids - desired_ids

            for sid in to_add:
                ResumeSummary.objects.create(resume_id=obj.id, summary_id=sid, active=False)
            if to_remove:
                ResumeSummary.objects.filter(
                    resume_id=obj.id, summary_id__in=list(to_remove)
                ).delete()

            # Determine which summary should be active:
            # 1. explicit active=true flag on an item, else
            # 2. last item in the provided list
            active_sid = desired_active_sid or (desired_ids_ordered[-1] if desired_ids_ordered else None)
            if active_sid:
                ResumeSummary.objects.filter(resume_id=obj.id).update(active=False)
                ResumeSummary.objects.filter(
                    resume_id=obj.id, summary_id=active_sid
                ).update(active=True)

        # Always enforce exactly one active summary
        ResumeSummary.ensure_single_active_for_resume(obj.id)

        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def create(self, request):
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        # Inject user_id from relationships ("user" or "users") into attrs for creation
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        relationships = node.get("relationships") or {}

        # Extract user_id from multiple possible shapes:
        # - data.relationships.user(.data.id)
        # - data.user / data.users (object or RI)
        # - data.attributes.user_id | user-id | userId
        # - top-level user_id | user-id | userId
        def _first_id(node_val):
            if isinstance(node_val, dict):
                d = node_val.get("data", node_val)
                if isinstance(d, dict) and "id" in d:
                    return d.get("id")
                if isinstance(d, list) and d:
                    first = d[0]
                    if isinstance(first, dict) and "id" in first:
                        return first["id"]
            elif node_val is not None:
                return node_val
            return None

        user_id_val = None

        # relationships.user|users
        rel_user = relationships.get("user") or relationships.get("users")
        if isinstance(rel_user, dict):
            user_id_val = _first_id(rel_user)

        # node-level user|users
        if user_id_val is None:
            user_id_val = _first_id(node.get("user") or node.get("users"))

        # attributes user_id variants
        if user_id_val is None:
            attrs_node = node.get("attributes") or {}
            user_id_val = (
                attrs_node.get("user_id")
                or attrs_node.get("user-id")
                or attrs_node.get("userId")
            )

        # top-level convenience keys
        if user_id_val is None:
            user_id_val = (
                data.get("user_id") or data.get("user-id") or data.get("userId")
            )

        # Default to authenticated user if no user_id provided
        if user_id_val is None and request.user.is_authenticated:
            user_id_val = request.user.id

        if user_id_val is not None:
            try:
                user_id_int = int(user_id_val)
            except (TypeError, ValueError):
                return Response(
                    {"errors": [{"detail": "Invalid user ID"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            User = get_user_model()
            if not User.objects.filter(id=user_id_int).exists():
                return Response(
                    {"errors": [{"detail": "Invalid user ID"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            attrs["user_id"] = user_id_int

        # Create the resume
        resume = Resume.objects.create(**attrs)

        # Extract potential child arrays from JSON:API attributes or top-level convenience keys
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs_node = node.get("attributes") or {}
        experiences_in = node.get("experiences") or data.get("experiences") or []
        educations_in = node.get("educations") or data.get("educations") or []
        certifications_in = (
            node.get("certifications") or data.get("certifications") or []
        )
        summaries_in = node.get("summaries") or data.get("summaries") or []

        # Helpers
        def _int_or_none(v):
            try:
                return int(v) if v is not None else None
            except (TypeError, ValueError):
                return None

        def _dp(val):
            if not val:
                return None
            try:
                dt = dateparser.parse(str(val))
                return dt.date() if dt else None
            except Exception:
                return None

        # Upsert Experiences and link to resume; also upsert/link nested descriptions
        for item in experiences_in or []:
            if not isinstance(item, dict):
                continue
            exp = None
            # Support explicit id if provided
            eid = (item.get("data") or {}).get("id") or item.get("id")
            if eid:
                exp = Experience.objects.filter(pk=eid).first()

            # Helper to normalize company_id
            rels = item.get("relationships") or {}
            company_rel = rels.get("company") or {}
            company_id = item.get("company_id") or (company_rel.get("data") or {}).get(
                "id"
            )
            company_id = int(company_id) if company_id is not None else None

            # Helper to collect incoming description lines in order
            incoming_lines = []
            if isinstance(item.get("description_lines"), list):
                incoming_lines = [
                    str(s).strip() for s in item["description_lines"] if str(s).strip()
                ]
            elif isinstance(item.get("descriptions"), list):
                # Accept either content or reference by id
                for d in item["descriptions"]:
                    if not isinstance(d, dict):
                        continue
                    if "content" in d and d["content"]:
                        incoming_lines.append(str(d["content"]).strip())
                    elif "id" in d and d["id"]:
                        dd = Description.objects.filter(pk=int(d["id"])).first()
                        if dd and getattr(dd, "content", None):
                            incoming_lines.append(dd.content.strip())

            # Parse dates via dateparser
            s_date = _dp(item.get("start_date"))
            e_date = _dp(item.get("end_date"))

            if exp is None:
                # Try to find an existing experience with matching scalars and identical description list
                candidates = list(
                    Experience.objects.filter(
                        company_id=company_id,
                        title=item.get("title"),
                        location=item.get("location"),
                        summary=(item.get("summary") or ""),
                        start_date=s_date,
                        end_date=e_date,
                    )
                )
                match = None
                for cand in candidates:
                    try:
                        existing_lines = [
                            d.content.strip()
                            for d in (cand.descriptions or [])
                            if getattr(d, "content", None)
                        ]
                    except Exception:
                        existing_lines = []
                    if existing_lines == incoming_lines:
                        match = cand
                        break

                if match:
                    exp = match
                else:
                    # Create new experience
                    exp = Experience.objects.create(
                        company_id=company_id,
                        title=item.get("title"),
                        location=item.get("location"),
                        summary=(item.get("summary") or ""),
                        start_date=s_date,
                        end_date=e_date,
                        content=item.get("content"),
                    )
                    # Link descriptions in order, creating Description rows as needed
                    ExperienceDescription.objects.filter(experience_id=exp.id).delete()
                    for idx, line in enumerate(incoming_lines or []):
                        if not line:
                            continue
                        desc, _ = Description.objects.get_or_create(content=line)
                        ExperienceDescription.objects.create(
                            experience_id=exp.id,
                            description_id=desc.id,
                            order=idx,
                        )

            # Join resume_experience (avoid duplicates)
            ResumeExperience.objects.get_or_create(
                resume_id=resume.id,
                experience_id=exp.id,
            )

            # Nested descriptions for this experience
            for d in item.get("descriptions") or []:
                if not isinstance(d, dict):
                    continue
                desc = None
                did = _int_or_none(d.get("id"))
                if did:
                    desc = Description.objects.filter(pk=did).first()
                if desc is None:
                    content = d.get("content")
                    if not content:
                        continue
                    desc, _ = Description.objects.get_or_create(content=content)
                # Link with optional order
                order = d.get("order")
                if order is None and isinstance(d.get("meta"), dict):
                    order = d["meta"].get("order")
                try:
                    order = int(order) if order is not None else 0
                except (TypeError, ValueError):
                    order = 0
                ExperienceDescription.objects.get_or_create(
                    experience_id=exp.id,
                    description_id=desc.id,
                    defaults={"order": order},
                )

        # Upsert Educations and link
        for item in educations_in or []:
            if not isinstance(item, dict):
                continue
            edu = None
            eid = _int_or_none(item.get("id"))
            if eid:
                edu = Education.objects.filter(pk=eid).first()
            if edu is None:
                lookup = {
                    "institution": item.get("institution"),
                    "degree": item.get("degree"),
                    "major": item.get("major"),
                    "minor": item.get("minor"),
                    "issue_date": _parse_date(item.get("issue_date")),
                }
                lookup = {k: v for k, v in lookup.items() if v is not None}
                edu = Education.objects.filter(**lookup).first()
                if not edu:
                    create_attrs = {
                        "institution": item.get("institution"),
                        "degree": item.get("degree"),
                        "major": item.get("major"),
                        "minor": item.get("minor"),
                        "issue_date": _parse_date(item.get("issue_date")),
                    }
                    create_attrs = {
                        k: v for k, v in create_attrs.items() if v is not None
                    }
                    edu = Education.objects.create(**create_attrs)
            ResumeEducation.objects.get_or_create(
                resume_id=resume.id, education_id=edu.id
            )

        # Upsert Certifications and link
        for item in certifications_in or []:
            if not isinstance(item, dict):
                continue
            cert = None
            cid = _int_or_none(item.get("id"))
            if cid:
                cert = Certification.objects.filter(pk=cid).first()
            if cert is None:
                lookup = {
                    "issuer": item.get("issuer"),
                    "title": item.get("title"),
                    "issue_date": _parse_date(item.get("issue_date")),
                }
                lookup = {k: v for k, v in lookup.items() if v is not None}
                cert = Certification.objects.filter(**lookup).first()
                if not cert:
                    create_attrs = {
                        "issuer": item.get("issuer"),
                        "title": item.get("title"),
                        "issue_date": _parse_date(item.get("issue_date")),
                        "content": item.get("content"),
                    }
                    create_attrs = {
                        k: v for k, v in create_attrs.items() if v is not None
                    }
                    cert = Certification.objects.create(**create_attrs)
            ResumeCertification.objects.get_or_create(
                resume_id=resume.id, certification_id=cert.id
            )

        # Upsert Skills and link
        skills_in = node.get("skills") or data.get("skills") or []
        for item in skills_in or []:
            if not isinstance(item, dict):
                continue
            s_node = (item.get("data") or item) or {}
            # Resolve or create skill
            skill = None
            sid = _int_or_none(s_node.get("id"))
            if sid:
                skill = Skill.objects.filter(pk=sid).first()
            if skill is None:
                s_attrs = s_node.get("attributes") or {}
                text = s_attrs.get("text") or s_node.get("text")
                if not text:
                    continue  # ignore invalid entries
                skill, _ = Skill.objects.get_or_create(text=str(text).strip())
            # Determine 'active' (default True)
            active_val = (s_node.get("attributes") or {}).get("active")
            active_val = bool(active_val) if active_val is not None else True
            ResumeSkill.objects.get_or_create(
                resume_id=resume.id,
                skill_id=skill.id,
                defaults={"active": active_val},
            )

        # Upsert Summaries and link
        active_set = False
        for item in summaries_in or []:
            if not isinstance(item, dict):
                continue
            s_node = (item.get("data") or item) or {}
            summary = None
            sid = s_node.get("id")
            if sid is not None:
                try:
                    summary = Summary.objects.filter(pk=int(sid)).first()
                except (TypeError, ValueError):
                    summary = None
                if summary is None:
                    return Response(
                        {"errors": [{"detail": f"Invalid summary id: {sid}"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )
            else:
                s_attrs = s_node.get("attributes") or {}
                s_rels = s_node.get("relationships") or {}
                content = s_attrs.get("content")
                if not content:
                    return Response(
                        {
                            "errors": [
                                {
                                    "detail": "Summary content is required when no id is provided"
                                }
                            ]
                        },
                        status=status.HTTP_400_BAD_REQUEST,
                    )

                # Resolve job_post if provided (accept hyphenated/underscored keys)
                jp_rel = (
                    s_rels.get("job-post")
                    or s_rels.get("job_post")
                    or s_rels.get("jobPost")
                    or s_rels.get("job-posts")
                    or s_rels.get("jobPosts")
                    or {}
                )
                jp_id = None
                if isinstance(jp_rel, dict):
                    d = jp_rel.get("data")
                    if isinstance(d, dict):
                        jp_id = d.get("id")
                try:
                    jp_id = int(jp_id) if jp_id is not None else None
                except (TypeError, ValueError):
                    return Response(
                        {
                            "errors": [
                                {
                                    "detail": "Invalid job-post ID in summary.relationships"
                                }
                            ]
                        },
                        status=status.HTTP_400_BAD_REQUEST,
                    )

                # Resolve user_id (fallback to resume.user_id)
                user_rel = s_rels.get("user") or s_rels.get("users") or {}
                u_id = None
                if isinstance(user_rel, dict):
                    d = user_rel.get("data")
                    if isinstance(d, dict):
                        u_id = d.get("id")
                try:
                    u_id = (
                        int(u_id)
                        if u_id is not None
                        else getattr(resume, "user_id", None)
                    )
                except (TypeError, ValueError):
                    u_id = getattr(resume, "user_id", None)

                summary = Summary.objects.create(
                    job_post_id=jp_id, user_id=u_id, content=content
                )

            # Link summary to resume; mark the first one as active
            ResumeSummary.objects.get_or_create(
                resume_id=resume.id,
                summary_id=summary.id,
                defaults={"active": (not active_set)},
            )
            active_set = True
        ResumeSummary.ensure_single_active_for_resume(resume.id)

        payload = {"data": ser.to_resource(resume)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([resume], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=["get"])
    def scores(self, request, pk=None):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        data = [ScoreSerializer().to_resource(s) for s in obj.scores.all()]
        return Response({"data": data})

    @action(
        detail=True,
        methods=["get"],
        url_path="cover-letters",
        permission_classes=[IsAuthenticated],
    )
    def cover_letters(self, request, pk=None):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        cover_letters = list(
            CoverLetter.objects.filter(resume_id=obj.id, user_id=request.user.id)
        )
        data = [CoverLetterSerializer().to_resource(c) for c in cover_letters]
        return Response({"data": data})

    @action(detail=True, methods=["get"], url_path="job-applications")
    def applications(self, request, pk=None):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        data = [
            JobApplicationSerializer().to_resource(a) for a in obj.applications.all()
        ]
        return Response({"data": data})

    @extend_schema(
        methods=["GET"],
        tags=["Resumes"],
        summary="List summaries for a resume",
        responses={200: _JSONAPI_LIST},
    )
    @extend_schema(
        methods=["POST"],
        tags=["Resumes"],
        summary="Create/AI-generate a summary for a resume",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Missing job-post or invalid IDs"),
            503: OpenApiResponse(description="AI client not configured"),
        },
    )
    @action(detail=True, methods=["get", "post"])
    def summaries(self, request, pk=None):
        if request.method.lower() == "post":
            obj = Resume.objects.filter(pk=int(pk)).first()  # obj is the Resume
            if not obj:
                return Response({"errors": [{"detail": "Not found"}]}, status=404)

            data = request.data if isinstance(request.data, dict) else {}
            node = data.get("data") or {}
            attrs = node.get("attributes") or {}
            relationships = node.get("relationships") or {}

            # Accept hyphenated or underscored job-post relationship; allow null if content is provided
            job_post_rel = (
                relationships.get("job-post")
                or relationships.get("job_post")
                or relationships.get("jobPost")
                or relationships.get("job-posts")
                or relationships.get("jobPosts")
                or {}
            )
            job_post_id = None
            if isinstance(job_post_rel, dict):
                d = job_post_rel.get("data")
                if isinstance(d, dict):
                    job_post_id = d.get("id")

            job_post = None
            if job_post_id is not None:
                try:
                    job_post = JobPost.objects.filter(pk=int(job_post_id)).first()
                except (TypeError, ValueError):
                    job_post = None
                if not job_post:
                    return Response(
                        {"errors": [{"detail": "Invalid job-post ID"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

            content = attrs.get("content")
            if content:
                summary = Summary(
                    job_post_id=job_post.id if job_post else None,
                    user_id=getattr(obj, "user_id", None),
                    content=content,
                )
                summary.save()
            else:
                if not job_post:
                    return Response(
                        {
                            "errors": [
                                {
                                    "detail": "Provide 'attributes.content' or a job-post relationship to generate content"
                                }
                            ]
                        },
                        status=status.HTTP_400_BAD_REQUEST,
                    )

                client = get_client(required=False)
                if client is None:
                    return Response(
                        {
                            "errors": [
                                {
                                    "detail": "AI client not configured. Set OPENAI_API_KEY."
                                }
                            ]
                        },
                        status=503,
                    )

                summary_service = SummaryService(client, job=job_post, resume=obj)
                summary = summary_service.generate_summary()

            ResumeSummary.objects.filter(resume_id=obj.id).update(active=False)
            ResumeSummary.objects.get_or_create(
                resume_id=obj.id, summary_id=summary.id, defaults={"active": True}
            )
            ResumeSummary.objects.filter(
                resume_id=obj.id, summary_id=summary.id
            ).update(active=True)
            ResumeSummary.ensure_single_active_for_resume(obj.id)

            ser = SummarySerializer()
            payload = {"data": ser.to_resource(summary)}
            include_rels = self._parse_include(request)
            if include_rels:
                payload["included"] = self._build_included(
                    [summary], include_rels, request
                )
            return Response(payload, status=status.HTTP_201_CREATED)

        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        ser = SummarySerializer()
        # Parent context available if serializers want to customize behavior
        if hasattr(ser, "set_parent_context"):
            ser.set_parent_context("resume", obj.id, "summaries")

        _summary_ids = list(
            ResumeSummary.objects.filter(resume_id=obj.id).values_list(
                "summary_id", flat=True
            )
        )
        items = list(Summary.objects.filter(pk__in=_summary_ids))
        data = [ser.to_resource(s) for s in items]

        # Build included only when ?include=... is provided
        include_rels = self._parse_include(request)

        payload = {"data": data}
        if include_rels:
            payload["included"] = self._build_included(
                items, include_rels, request, primary_serializer=ser
            )
        return Response(payload)

    @extend_schema(
        tags=["Resumes"],
        summary="Retrieve a specific summary linked to a resume",
        responses={200: _JSONAPI_ITEM, 404: OpenApiResponse(description="Not found")},
    )
    @action(detail=True, methods=["get"], url_path=r"summaries/(?P<summary_id>\d+)")
    def summary(self, request, pk=None, summary_id=None):
        resume = Resume.objects.filter(pk=int(pk)).first()
        if not resume:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        try:
            sid = int(summary_id)
        except (TypeError, ValueError):
            return Response({"errors": [{"detail": "Invalid summary id"}]}, status=400)
        summary = Summary.objects.filter(pk=sid).first()
        if not summary:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        # Ensure the summary is associated with this resume via the link table
        link = ResumeSummary.objects.filter(
            resume_id=resume.id, summary_id=summary.id
        ).first()
        if not link:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = SummarySerializer()
        if hasattr(ser, "set_parent_context"):
            ser.set_parent_context("resume", resume.id, "summaries")

        # Primary data
        items = [summary]
        data = ser.to_resource(summary)

        # Build included only when ?include=... is provided, using Summary serializer
        include_rels = self._parse_include(request)

        payload = {"data": data}
        if include_rels:
            payload["included"] = self._build_included(
                items, include_rels, request, primary_serializer=ser
            )
        return Response(payload)

    @extend_schema(
        tags=["Resumes"],
        summary="List experiences for a resume",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def experiences(self, request, pk=None):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = ExperienceSerializer()
        ser.set_parent_context("resume", obj.id, "experiences")
        _exp_ids = list(
            ResumeExperience.objects.filter(resume_id=obj.id)
            .order_by("order")
            .values_list("experience_id", flat=True)
        )
        _exp_map = {e.id: e for e in Experience.objects.filter(pk__in=_exp_ids)}
        items = [_exp_map[eid] for eid in _exp_ids if eid in _exp_map]
        data = [ser.to_resource(e) for e in items]

        # Build included only when ?include=... is provided
        include_rels = self._parse_include(request)

        payload = {"data": data}
        if include_rels:
            payload["included"] = self._build_included(
                items, include_rels, request, primary_serializer=ser
            )
        return Response(payload)

    @extend_schema(
        tags=["Resumes"],
        summary="List educations for a resume",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def educations(self, request, pk=None):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = EducationSerializer()
        ser.set_parent_context("resume", obj.id, "educations")
        _edu_ids = list(
            ResumeEducation.objects.filter(resume_id=obj.id).values_list(
                "education_id", flat=True
            )
        )
        data = [ser.to_resource(e) for e in Education.objects.filter(pk__in=_edu_ids)]
        return Response({"data": data})

    @extend_schema(
        tags=["Resumes"],
        summary="List skills for a resume",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def skills(self, request, pk=None):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = SkillSerializer()
        ser.set_parent_context("resume", obj.id, "skills")
        skill_ids = list(
            ResumeSkill.objects.filter(resume_id=obj.id).values_list(
                "skill_id", flat=True
            )
        )
        items = list(Skill.objects.filter(pk__in=skill_ids))
        data = [ser.to_resource(s) for s in items]

        # Build included only when ?include=... is provided
        include_rels = self._parse_include(request)

        payload = {"data": data}
        if include_rels:
            payload["included"] = self._build_included(
                items, include_rels, request, primary_serializer=ser
            )
        return Response(payload)

    @extend_schema(
        tags=["Resumes"],
        summary="Export a resume as DOCX or Markdown",
        parameters=[
            OpenApiParameter(
                "format",
                OpenApiTypes.STR,
                OpenApiParameter.QUERY,
                required=False,
                description="'docx' (default) or 'md'",
            ),
            OpenApiParameter(
                "template_path",
                OpenApiTypes.STR,
                OpenApiParameter.QUERY,
                required=False,
                description="Path to a custom DOCX template (docx only)",
            ),
        ],
        responses={
            200: OpenApiResponse(
                description="File download (application/vnd.openxmlformats-officedocument.wordprocessingml.document or text/markdown)"
            ),
            404: OpenApiResponse(description="Not found"),
        },
    )
    @action(detail=True, methods=["get"], url_path="export")
    def export(self, request, pk=None):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        format_param = request.query_params.get("format", "docx").lower()
        template_path_param = request.query_params.get("template_path")

        if format_param == "md":
            # Export as markdown
            exporter = DbExportService()
            markdown_content = exporter.resume_markdown_export(obj)

            # Generate filename
            import re

            filename_parts = ["resume", str(obj.id)]
            try:
                if getattr(obj, "user", None) and getattr(obj.user, "name", None):
                    name = str(obj.user.name)
                    sanitized_name = re.sub(r"[^A-Za-z0-9_-]+", "-", name)
                    if sanitized_name:
                        filename_parts.append(sanitized_name)
                if getattr(obj, "title", None):
                    title = str(obj.title)
                    sanitized_title = re.sub(r"[^A-Za-z0-9_-]+", "-", title)
                    if sanitized_title:
                        filename_parts.append(sanitized_title)
            except Exception:
                pass
            filename = "-".join([p for p in filename_parts if p]) + ".md"

            response = HttpResponse(
                markdown_content.encode("utf-8"),
                content_type="text/markdown; charset=utf-8",
            )
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            return response

        else:
            # Export as DOCX (default)
            try:
                svc = ResumeExportService()
                data = svc.render_docx(obj, template_path=template_path_param)
            except ImportError:
                return Response(
                    {
                        "errors": [
                            {"detail": "DOCX export requires 'docxtpl' to be installed"}
                        ]
                    },
                    status=status.HTTP_501_NOT_IMPLEMENTED,
                )
            except Exception as e:
                return Response(
                    {"errors": [{"detail": str(e)}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            # Generate filename
            import re

            filename_parts = ["resume", str(obj.id)]
            try:
                if getattr(obj, "user", None) and getattr(obj.user, "name", None):
                    name = str(obj.user.name)
                    sanitized_name = re.sub(r"[^A-Za-z0-9_-]+", "-", name)
                    if sanitized_name:
                        filename_parts.append(sanitized_name)
                if getattr(obj, "title", None):
                    title = str(obj.title)
                    sanitized_title = re.sub(r"[^A-Za-z0-9_-]+", "-", title)
                    if sanitized_title:
                        filename_parts.append(sanitized_title)
            except Exception:
                pass
            filename = "-".join([p for p in filename_parts if p]) + ".docx"

            response = HttpResponse(
                data,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            return response

    @extend_schema(
        tags=["Resumes"],
        summary="Clone a resume and all its relationships",
        responses={
            201: _JSONAPI_ITEM,
            404: OpenApiResponse(description="Not found"),
        },
    )
    @action(detail=True, methods=["post"], url_path="clone")
    def clone(self, request, pk=None):
        obj = Resume.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        new_resume = Resume.objects.create(
            user_id=obj.user_id,
            file_path=obj.file_path,
            title=f"{obj.title} (clone)" if obj.title else None,
            name=f"{obj.name} (clone)" if obj.name else None,
            notes=obj.notes,
            favorite=obj.favorite,
        )

        for rs in ResumeSkill.objects.filter(resume_id=obj.id):
            ResumeSkill.objects.create(resume_id=new_resume.id, skill_id=rs.skill_id, active=rs.active)

        for re in ResumeExperience.objects.filter(resume_id=obj.id):
            ResumeExperience.objects.create(resume_id=new_resume.id, experience_id=re.experience_id, order=re.order)

        for rp in ResumeProject.objects.filter(resume_id=obj.id):
            ResumeProject.objects.create(resume_id=new_resume.id, project_id=rp.project_id, order=rp.order)

        for rc in ResumeCertification.objects.filter(resume_id=obj.id):
            ResumeCertification.objects.create(
                resume_id=new_resume.id,
                certification_id=rc.certification_id,
                issuer=rc.issuer,
                title=rc.title,
                issue_date=rc.issue_date,
                content=rc.content,
            )

        for red in ResumeEducation.objects.filter(resume_id=obj.id):
            ResumeEducation.objects.create(
                resume_id=new_resume.id,
                education_id=red.education_id,
                institution=red.institution,
                degree=red.degree,
                issue_date=red.issue_date,
                content=red.content,
            )

        for rsm in ResumeSummary.objects.filter(resume_id=obj.id):
            ResumeSummary.objects.create(resume_id=new_resume.id, summary_id=rsm.summary_id, active=rsm.active)

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(new_resume)}
        ser_rels = list(getattr(ser, "relationships", {}).keys())
        include_rels = self._parse_include(request) or ser_rels
        if include_rels:
            payload["included"] = self._build_included([new_resume], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)

    @extend_schema(
        tags=["Resumes"],
        summary="Ingest a resume from an uploaded DOCX file",
        request=inline_serializer(
            name="IngestResumeRequest",
            fields={
                "file": drf_serializers.FileField(
                    help_text="DOCX resume file (multipart/form-data)"
                )
            },
        ),
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="No file provided or unsupported format"),
        },
    )
    @action(detail=False, methods=["post"], url_path="ingest")
    def ingest(self, request):
        """
        Ingest a resume from an uploaded docx file and create a new resume.

        Expects a multipart/form-data request with a 'file' field containing the docx.
        """
        # Check if file was uploaded
        if "file" not in request.FILES:
            return Response(
                {
                    "errors": [
                        {
                            "detail": "No file uploaded. Expected 'file' field with docx content."
                        }
                    ]
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        uploaded_file = request.FILES["file"]

        # Validate file type
        if not uploaded_file.name.lower().endswith(".docx"):
            return Response(
                {"errors": [{"detail": "Only .docx files are supported"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            # Read file content as blob
            file_blob = uploaded_file.read()

            # Create a new resume for the authenticated user
            # resume = Resume(user_id=request.user.id, file_path=uploaded_file.name)

            # Create IngestResume service with the blob
            resume_name = uploaded_file.name
            ingest_service = IngestResume(
                user=request.user,
                resume=file_blob,  # Pass blob instead of path
                resume_name=resume_name,
                agent=None,  # Will use default agent
                # agent=ollama_model,  # Will use default agent
            )

            # Process the resume
            result = ingest_service.process()
            resume = result or ingest_service.db_resume

            # Guard against None resume from failed ingestion
            if resume is None:
                return Response(
                    {"errors": [{"detail": "Ingest failed: no resume created"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            # Set resume title from filename if not already set
            if hasattr(resume, "title") and not getattr(resume, "title", None):
                # Derive name from uploaded filename
                base_name = uploaded_file.name
                if base_name.lower().endswith(".docx"):
                    base_name = base_name[:-5]  # Remove .docx extension
                derived_name = base_name.strip()[:100]  # Trim to reasonable length
                if derived_name:
                    resume.title = derived_name
                    resume.save()
            elif hasattr(resume, "name") and not getattr(resume, "name", None):
                # Fallback to 'name' field if 'title' doesn't exist
                base_name = uploaded_file.name
                if base_name.lower().endswith(".docx"):
                    base_name = base_name[:-5]  # Remove .docx extension
                derived_name = base_name.strip()[:100]  # Trim to reasonable length
                if derived_name:
                    resume.name = derived_name
                    resume.save()

            # Return the created resume with all relationships
            ser = self.get_serializer()
            payload = {"data": ser.to_resource(resume)}

            # Include all relationships by default for ingest response
            ser_rels = list(getattr(ser, "relationships", {}).keys())
            include_rels = self._parse_include(request) or ser_rels
            if include_rels:
                payload["included"] = self._build_included(
                    [resume], include_rels, request
                )

            return Response(payload, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response(
                {"errors": [{"detail": f"Failed to process resume: {str(e)}"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )


@extend_schema_view(
    list=extend_schema(tags=["Scores"], summary="List scores"),
    retrieve=extend_schema(tags=["Scores"], summary="Retrieve a score"),
    update=extend_schema(tags=["Scores"], summary="Update a score"),
    partial_update=extend_schema(tags=["Scores"], summary="Partially update a score"),
    destroy=extend_schema(tags=["Scores"], summary="Delete a score"),
)
class ScoreViewSet(BaseViewSet):
    model = Score
    serializer_class = ScoreSerializer

    @extend_schema(
        tags=["Scores"],
        summary="AI-score a job post against a resume — returns immediately with status=pending; poll for status=completed",
        request=_JSONAPI_WRITE,
        responses={
            202: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Missing required relationships or no career data"),
            503: OpenApiResponse(description="AI client not configured"),
        },
    )
    def create(self, request):
        data = request.data if isinstance(request.data, dict) else {}

        client = get_client(required=False)
        if client is None:
            return Response(
                {
                    "errors": [
                        {"detail": "AI client not configured. Set OPENAI_API_KEY."}
                    ]
                },
                status=503,
            )

        myJobScorer = JobScorer(client)

        relationships = (data.get("data") or {}).get("relationships") or {}

        def _first_id(node):
            if isinstance(node, dict):
                data = node.get("data")
            else:
                data = None
            if isinstance(data, dict) and "id" in data:
                return data["id"]
            if isinstance(data, list) and data:
                first = data[0]
                if isinstance(first, dict) and "id" in first:
                    return first["id"]
            return None

        def _rel_id(*keys):
            for k in keys:
                val = relationships.get(k)
                if val is not None:
                    rid = _first_id(val)
                    if rid is not None:
                        return rid
            return None

        job_post_id = _rel_id(
            "job-post", "job_post", "jobPost", "job-posts", "jobPosts"
        )
        user_id = _rel_id("user", "users")
        resume_id = _rel_id("resume", "resumes")

        if job_post_id is None:
            return Response(
                {"errors": [{"detail": "Missing required relationship: job-post"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

        job_post_id = int(job_post_id)
        # Infer user from auth token; relationship is optional
        user_id = int(user_id) if user_id is not None else request.user.id
        # Missing or null resume defaults to career-data scoring (equivalent to resume_id=0)
        resume_id = int(resume_id) if resume_id is not None else 0

        jp = JobPost.objects.filter(pk=job_post_id).first()
        if not jp:
            return Response(
                {"errors": [{"detail": "Job post not found"}]},
                status=status.HTTP_404_NOT_FOUND,
            )
        if not jp.description or not jp.description.strip():
            return Response(
                {"errors": [{"detail": "Job post has no description to score against"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        exporter = DbExportService()

        if resume_id == 0:
            # Score against the user's full career data (all favorite resumes, cover letters, answers)
            career_data = CareerData.for_user(user_id)
            prompt_builder = ApplicationPromptBuilder(max_section_chars=60000)
            resume_markdown = prompt_builder.build_from_career_data(career_data)
            if not resume_markdown.strip():
                return Response(
                    {"errors": [{"detail": "No career data found for this user to score against"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            score_resume_id = None
        else:
            resume = Resume.objects.filter(pk=resume_id).first()
            if not resume:
                return Response(
                    {"errors": [{"detail": "Resume not found"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            resume_markdown = exporter.resume_markdown_export(resume)
            score_resume_id = resume_id

        myScore = Score.objects.filter(
            job_post_id=job_post_id, resume_id=score_resume_id, user_id=user_id
        ).first()
        if myScore:
            myScore.status = "pending"
            myScore.score = None
            myScore.explanation = None
            myScore.save()
        else:
            myScore = Score.objects.create(
                job_post_id=job_post_id,
                resume_id=score_resume_id,
                user_id=user_id,
                status="pending",
            )

        score_id = myScore.id
        captured_description = jp.description
        captured_resume_markdown = resume_markdown

        captured_user_id = request.user.id

        def _score():
            import django
            django.db.close_old_connections()
            try:
                result = myJobScorer.score_job_match(captured_description, captured_resume_markdown)
                Score.objects.filter(pk=score_id).update(
                    score=result.score,
                    explanation=result.evaluation,
                    status="completed",
                )
                # Record AI usage for scoring
                usage = getattr(result, "_usage", None)
                model_name = getattr(result, "_model_name", "unknown")
                if usage:
                    AiUsage.objects.create(
                        user_id=captured_user_id,
                        agent_name="job_scorer",
                        model_name=model_name,
                        trigger="score",
                        request_tokens=usage.request_tokens or 0,
                        response_tokens=usage.response_tokens or 0,
                        total_tokens=usage.total_tokens or 0,
                        request_count=usage.requests or 1,
                    )
            except Exception:
                Score.objects.filter(pk=score_id).update(status="failed")

        import threading
        threading.Thread(target=_score, daemon=True).start()

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(myScore)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([myScore], include_rels, request)
        return Response(payload, status=status.HTTP_202_ACCEPTED)

    def _parse_eval(self, e):
        # Expect a structured result from JobScorer: dict or JSON string
        data = None
        if isinstance(e, dict):
            data = e
        else:
            try:
                import json as _json

                data = _json.loads(str(e))
            except Exception:
                return None, str(e)

        s = data.get("score")
        explanation = (
            data.get("explanation") or data.get("evaluation") or data.get("explination")
        )
        try:
            s_int = int(s) if s is not None else None
        except (TypeError, ValueError):
            s_int = None

        if s_int is None or not (1 <= s_int <= 100):
            return None, str(explanation or "").strip() or str(e)

        return s_int, str(explanation or "").strip()

    def list(self, request):
        qs = Score.objects.filter(user_id=request.user.id).order_by("-created_at", "-id")
        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs[offset: offset + page_size])
        ser = self.get_serializer()
        return Response({
            "data": [ser.to_resource(o) for o in items],
            "meta": {"total": total, "page": page_number, "per_page": page_size, "total_pages": total_pages},
        })

    def retrieve(self, request, pk=None):
        obj = Score.objects.filter(pk=pk).first()
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        return Response({"data": ser.to_resource(obj)})

    def destroy(self, request, pk=None):
        obj = Score.objects.filter(pk=pk).first()
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        obj.delete()
        return Response(status=204)


@extend_schema_view(
    list=extend_schema(
        tags=["Job Posts"],
        summary="List job posts",
        parameters=_PAGE_PARAMS + [
            _SORT_PARAM,
            _FILTER_QUERY_PARAM,
            _FILTER_COMPANY_PARAM,
            _FILTER_COMPANY_ID_PARAM,
            _FILTER_TITLE_PARAM,
        ],
    ),
    retrieve=extend_schema(tags=["Job Posts"], summary="Retrieve a job post"),
    create=extend_schema(
        tags=["Job Posts"],
        summary="Create a job post (created_by set to authenticated user)",
    ),
    update=extend_schema(tags=["Job Posts"], summary="Update a job post"),
    partial_update=extend_schema(
        tags=["Job Posts"], summary="Partially update a job post"
    ),
    destroy=extend_schema(tags=["Job Posts"], summary="Delete a job post"),
)
class JobPostViewSet(BaseViewSet):
    model = JobPost
    serializer_class = JobPostSerializer



    def pre_save_payload(self, request, attrs, creating):
        # Remove any client-supplied ownership fields so they can't be spoofed
        attrs.pop("created_by", None)
        attrs.pop("created_by_id", None)  # defensive

        # If creating, set created_by_id to the authenticated user
        if creating:
            attrs["created_by_id"] = request.user.id

        return attrs

    @staticmethod
    def _parse_date_attrs(attrs):
        """Parse posted_date and extraction_date from ISO strings to date objects."""
        from dateutil import parser as dateutil_parser
        from datetime import date as date_type

        errors = {}
        for field in ("posted_date", "extraction_date"):
            if field not in attrs or attrs[field] is None:
                continue
            val = attrs[field]
            if isinstance(val, date_type):
                continue
            try:
                attrs[field] = dateutil_parser.parse(str(val)).date()
            except (ValueError, TypeError):
                errors[field] = (
                    f"Invalid {field}: {val!r}. Expected a date (e.g. '2025-01-15')."
                )
        return errors

    def list(self, request):
        qs = JobPost.objects.filter(
            Q(created_by_id=request.user.id) |
            Q(applications__user_id=request.user.id) |
            Q(scores__user_id=request.user.id)
        ).distinct()
        link_filter = request.query_params.get("filter[link]")
        if link_filter is not None:
            qs = qs.filter(link=link_filter)

        company_id_filter = request.query_params.get("filter[company_id]")
        if company_id_filter is not None:
            qs = qs.filter(company_id=company_id_filter)

        company_filter = request.query_params.get("filter[company]")
        if company_filter is not None:
            qs = qs.filter(company__name__icontains=company_filter)

        title_filter = request.query_params.get("filter[title]")
        if title_filter is not None:
            qs = qs.filter(title__icontains=title_filter)

        query_filter = request.query_params.get("filter[query]")
        if query_filter is not None:
            qs = qs.filter(
                Q(title__icontains=query_filter)
                | Q(description__icontains=query_filter)
                | Q(company__name__icontains=query_filter)
                | Q(company__display_name__icontains=query_filter)
            ).distinct()

        sort_param = request.query_params.get("sort")
        if sort_param:
            sort_fields = []
            for field in sort_param.split(","):
                field = field.strip()
                if field.startswith("-"):
                    sort_fields.append(f"-{field[1:]}")
                else:
                    sort_fields.append(field)
            if sort_fields:
                qs = qs.order_by(*sort_fields)

        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs.all()[offset: offset + page_size])

        # Attach the highest Score to each job post in one query
        if items:
            job_post_ids = [jp.id for jp in items]
            all_scores = list(
                Score.objects.filter(job_post_id__in=job_post_ids, user_id=request.user.id).order_by("job_post_id", "-score")
            )
            top_score_map = {}
            for s in all_scores:
                if s.job_post_id not in top_score_map:
                    top_score_map[s.job_post_id] = s
            for jp in items:
                jp._top_score = top_score_map.get(jp.id)

        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {
            "data": data,
            "meta": {
                "total": total,
                "page": page_number,
                "per_page": page_size,
                "total_pages": total_pages,
            },
        }
        if page_number < total_pages:
            base = request.build_absolute_uri(request.path)
            # Preserve existing query params, overriding page
            qp = request.query_params.dict()
            qp["page"] = page_number + 1
            qp["per_page"] = page_size
            next_url = base + "?" + "&".join(f"{k}={v}" for k, v in qp.items())
            payload["links"] = {"next": next_url}
        else:
            payload["links"] = {"next": None}

        include_rels = self._parse_include(request) or ["top-score"]
        payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = JobPost.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        has_access = (
            obj.created_by_id == request.user.id or
            obj.applications.filter(user_id=request.user.id).exists() or
            obj.scores.filter(user_id=request.user.id).exists()
        )
        if not has_access:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def create(self, request):
        data = request.data if isinstance(request.data, dict) else {}
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data) if "data" in data else {}
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)
        attrs = self.pre_save_payload(request, attrs, creating=True)
        date_errors = self._parse_date_attrs(attrs)
        if date_errors:
            return Response(
                {"errors": [{"detail": v} for v in date_errors.values()]}, status=400
            )
        obj = JobPost(**attrs)
        obj.save()
        if not obj.posted_date:
            obj.posted_date = obj.created_at.date()
            obj.save(update_fields=["posted_date"])
        return Response({"data": ser.to_resource(obj)}, status=status.HTTP_201_CREATED)

    def update(self, request, pk=None):
        return self._upsert_django(request, pk, partial=False)

    def partial_update(self, request, pk=None):
        return self._upsert_django(request, pk, partial=True)

    def _upsert_django(self, request, pk, partial=False):
        obj = JobPost.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        if obj.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)
        attrs = self.pre_save_payload(request, attrs, creating=False)
        attrs.pop("created_by_id", None)
        attrs.pop("created_at", None)  # never allow overriding auto timestamp
        date_errors = self._parse_date_attrs(attrs)
        if date_errors:
            return Response(
                {"errors": [{"detail": v} for v in date_errors.values()]}, status=400
            )
        for k, v in attrs.items():
            setattr(obj, k, v)
        obj.save()
        return Response({"data": ser.to_resource(obj)})

    def destroy(self, request, pk=None):
        obj = JobPost.objects.filter(pk=pk).first()
        if not obj:
            return Response(status=204)
        if obj.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)
        if obj.applications.exclude(user_id=request.user.id).exists():
            return Response(
                {"errors": [{"detail": "Cannot delete: other users have applications on this post"}]},
                status=409,
            )
        obj.delete()
        return Response(status=204)

    @extend_schema(
        tags=["Job Posts"],
        summary="List scores for a job post",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def scores(self, request, pk=None):
        if not JobPost.objects.filter(pk=pk).exists():
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        scores = list(Score.objects.filter(job_post_id=int(pk), user_id=request.user.id))
        data = [ScoreSerializer().to_resource(s) for s in scores]
        return Response({"data": data})

    @extend_schema(
        tags=["Job Posts"],
        summary="List scrapes for a job post",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def scrapes(self, request, pk=None):
        if not JobPost.objects.filter(pk=pk).exists():
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        scrapes = list(Scrape.objects.filter(job_post_id=int(pk)))
        data = [ScrapeSerializer().to_resource(s) for s in scrapes]
        return Response({"data": data})

    @extend_schema(
        tags=["Job Posts"],
        summary="List cover letters for a job post (authenticated user's only)",
        responses={200: _JSONAPI_LIST},
    )
    @action(
        detail=True,
        methods=["get"],
        url_path="cover-letters",
        permission_classes=[IsAuthenticated],
    )
    def cover_letters(self, request, pk=None):
        if not JobPost.objects.filter(pk=pk).exists():
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        cover_letters = list(
            CoverLetter.objects.filter(job_post_id=int(pk), user_id=request.user.id)
        )
        data = [CoverLetterSerializer().to_resource(c) for c in cover_letters]
        return Response({"data": data})

    @extend_schema(
        tags=["Job Posts"],
        summary="List job applications for a job post",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"], url_path="job-applications")
    def applications(self, request, pk=None):
        if not JobPost.objects.filter(pk=pk).exists():
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        apps = list(JobApplication.objects.filter(job_post_id=int(pk), user_id=request.user.id))
        data = [JobApplicationSerializer().to_resource(a) for a in apps]
        return Response({"data": data})

    @extend_schema(
        methods=["GET"],
        tags=["Job Posts"],
        summary="List questions for a job post",
        responses={200: _JSONAPI_LIST},
    )
    @extend_schema(
        methods=["POST"],
        tags=["Job Posts"],
        summary="Create a question for a job post (company auto-set from job post)",
        request=_JSONAPI_WRITE,
        responses={201: _JSONAPI_ITEM, 404: OpenApiResponse(description="Not found")},
    )
    @action(detail=True, methods=["get", "post"])
    def questions(self, request, pk=None):
        job_post = JobPost.objects.filter(pk=int(pk)).first()
        if not job_post:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        ser = QuestionSerializer()

        if request.method.lower() == "post":
            try:
                attrs = ser.parse_payload(request.data)
            except ValueError as e:
                return Response({"errors": [{"detail": str(e)}]}, status=400)
            attrs["created_by_id"] = request.user.id
            attrs.setdefault("company_id", job_post.company_id)
            attrs.setdefault("job_post_id", job_post.id)
            safe_attrs = {
                k: v
                for k, v in attrs.items()
                if k in ("content", "favorite", "application_id", "company_id", "created_by_id")
            }
            obj = Question.objects.create(**safe_attrs)
            include_rels = self._parse_include(request)
            payload = {"data": ser.to_resource(obj)}
            if include_rels:
                payload["included"] = self._build_included([obj], include_rels, request, primary_serializer=ser)
            return Response(payload, status=status.HTTP_201_CREATED)

        items = list(Question.objects.filter(application__job_post_id=int(pk), created_by_id=request.user.id))
        include_rels = self._parse_include(request)
        payload = {"data": [ser.to_resource(i) for i in items]}
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request, primary_serializer=ser)
        return Response(payload)

    @extend_schema(
        methods=["GET"],
        tags=["Job Posts"],
        summary="List summaries for a job post",
        responses={200: _JSONAPI_LIST},
    )
    @extend_schema(
        methods=["POST"],
        tags=["Job Posts"],
        summary="Create/AI-generate a summary for a job post",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Missing resume"),
            503: OpenApiResponse(description="AI client not configured"),
        },
    )
    @action(detail=True, methods=["get", "post"])
    def summaries(self, request, pk=None):
        if request.method.lower() == "post":
            obj = JobPost.objects.filter(pk=pk).first()
            if not obj:
                return Response({"errors": [{"detail": "Not found"}]}, status=404)

            data = request.data if isinstance(request.data, dict) else {}
            node = data.get("data") or {}
            attrs = node.get("attributes") or {}
            relationships = node.get("relationships") or {}

            # Accept "resume"/"resumes" for resume relationship
            resume_rel = (
                relationships.get("resumes") or relationships.get("resume") or {}
            )
            resume_id = None
            if isinstance(resume_rel, dict):
                d = resume_rel.get("data")
                if isinstance(d, dict):
                    resume_id = d.get("id")

            if not resume_id:
                return Response(
                    {"errors": [{"detail": "Missing required relationship: resume"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            try:
                resume = Resume.objects.filter(pk=int(resume_id)).first()
            except (TypeError, ValueError):
                resume = None

            if not resume:
                return Response(
                    {"errors": [{"detail": "Invalid resume ID"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            content = attrs.get("content")
            if content:
                summary = Summary(
                    job_post_id=obj.id,
                    user_id=getattr(resume, "user_id", None),
                    content=content,
                )
                summary.save()
            else:
                client = get_client(required=False)
                if client is None:
                    return Response(
                        {
                            "errors": [
                                {
                                    "detail": "AI client not configured. Set OPENAI_API_KEY."
                                }
                            ]
                        },
                        status=503,
                    )

                summary_service = SummaryService(client, job=obj, resume=resume)
                summary = summary_service.generate_summary()

            ResumeSummary.objects.filter(resume_id=resume.id).update(active=False)
            ResumeSummary.objects.get_or_create(
                resume_id=resume.id, summary_id=summary.id, defaults={"active": True}
            )
            ResumeSummary.objects.filter(
                resume_id=resume.id, summary_id=summary.id
            ).update(active=True)
            ResumeSummary.ensure_single_active_for_resume(resume.id)

            ser = SummarySerializer()
            payload = {"data": ser.to_resource(summary)}
            include_rels = self._parse_include(request)
            if include_rels:
                payload["included"] = self._build_included(
                    [summary], include_rels, request
                )
            return Response(payload, status=status.HTTP_201_CREATED)

        obj = JobPost.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        summaries = list(Summary.objects.filter(job_post_id=obj.id, user_id=request.user.id))
        data = [SummarySerializer().to_resource(s) for s in summaries]
        return Response({"data": data})


@extend_schema_view(
    list=extend_schema(
        tags=["Scrapes"],
        summary="List scrapes",
        parameters=_PAGE_PARAMS + [_SORT_PARAM],
    ),
    update=extend_schema(tags=["Scrapes"], summary="Update a scrape"),
    partial_update=extend_schema(tags=["Scrapes"], summary="Partially update a scrape"),
    destroy=extend_schema(tags=["Scrapes"], summary="Delete a scrape"),
)
class ScrapeViewSet(BaseViewSet):
    model = Scrape
    serializer_class = ScrapeSerializer

    def list(self, request):
        qs = Scrape.objects.filter(
            Q(created_by=request.user)
            | Q(job_post__created_by_id=request.user.id)
            | Q(job_post__isnull=True, created_by__isnull=True)
        )

        # Sorting
        sort_param = request.query_params.get("sort")
        if sort_param:
            sort_fields = []
            for field in sort_param.split(","):
                field = field.strip()
                if field.startswith("-"):
                    sort_fields.append(f"-{field[1:]}")
                else:
                    sort_fields.append(field)
            if sort_fields:
                qs = qs.order_by(*sort_fields)
        else:
            # Default: latest first
            try:
                Scrape._meta.get_field("created_at")
                qs = qs.order_by("-created_at")
            except Exception:
                qs = qs.order_by("-id")

        # Status filter
        status_filter = request.query_params.get("filter[status]")
        if status_filter:
            qs = qs.filter(status=status_filter)

        # Pagination
        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs.all()[offset : offset + page_size])

        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {
            "data": data,
            "meta": {
                "total": total,
                "page": page_number,
                "per_page": page_size,
                "total_pages": total_pages,
            },
        }

        if page_number < total_pages:
            base = request.build_absolute_uri(request.path)
            qp = request.query_params.dict()
            qp["page"] = page_number + 1
            qp["per_page"] = page_size
            payload["links"] = {"next": base + "?" + "&".join(f"{k}={v}" for k, v in qp.items())}
        else:
            payload["links"] = {"next": None}

        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def pre_save_payload(self, request, attrs, creating=False):
        attrs = super().pre_save_payload(request, attrs, creating=creating)
        if "html" in attrs and attrs["html"] and not attrs.get("job_content"):
            from job_hunting.lib.scrapers.html_cleaner import clean_html_to_markdown
            attrs["job_content"] = clean_html_to_markdown(attrs["html"])
        if attrs.get("job_content"):
            from job_hunting.lib.scrapers.html_cleaner import strip_agent_chat
            attrs["job_content"] = strip_agent_chat(attrs["job_content"])
        if attrs.get("status") == "completed" and not attrs.get("scraped_at"):
            from django.utils import timezone
            attrs["scraped_at"] = timezone.now()
        return attrs

    def _sync_associations(self, pk):
        """After an update, ensure company_id mirrors the job post's company."""
        scrape = Scrape.objects.filter(pk=int(pk)).first()
        if not scrape:
            return
        if scrape.job_post_id and not scrape.company_id:
            jp = JobPost.objects.filter(pk=scrape.job_post_id).first()
            if jp and jp.company_id:
                scrape.company_id = jp.company_id
                scrape.save(update_fields=["company_id"])

    def _maybe_trigger_extraction(self, pk):
        from job_hunting.lib.scraper import _maybe_caddy_extract
        scrape = Scrape.objects.filter(pk=int(pk)).first()
        if scrape:
            _maybe_caddy_extract(scrape)

    def update(self, request, pk=None):
        response = super().update(request, pk=pk)
        self._sync_associations(pk)
        self._maybe_trigger_extraction(pk)
        return response

    def partial_update(self, request, pk=None):
        response = super().partial_update(request, pk=pk)
        self._sync_associations(pk)
        self._maybe_trigger_extraction(pk)
        return response

    @extend_schema(
        tags=["Scrapes"],
        summary="Get the current status of a scrape",
        responses={200: _JSONAPI_ITEM, 404: OpenApiResponse(description="Not found")},
    )
    def retrieve(self, request, pk=None):
        """Get the current status of a scrape"""
        obj = Scrape.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        if obj.created_by_id and obj.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        if obj.job_post_id and obj.job_post.created_by_id and obj.job_post.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["Scrapes"],
        summary="Initiate a URL scrape (async — returns 202 Accepted). Returns existing scrape if URL already processed.",
        request=inline_serializer(
            name="ScrapeCreateRequest",
            fields={"url": drf_serializers.URLField(help_text="URL to scrape")},
        ),
        responses={
            202: OpenApiResponse(description="Scrape started"),
            200: OpenApiResponse(description="Existing scrape returned"),
            400: OpenApiResponse(description="URL missing"),
            501: OpenApiResponse(description="Scraping disabled"),
        },
    )
    def create(self, request):

        # Detect a "url" key in either a plain JSON body or JSON:API attributes
        data = request.data if isinstance(request.data, dict) else {}
        url = data.get("url")
        attrs = {}
        if isinstance(data.get("data"), dict):
            attrs = data["data"].get("attributes") or {}
            if url is None:
                url = attrs.get("url")

        # "hold" status: create the scrape record without dispatching the scraper.
        # Used by MCP agents to queue URLs for later processing.
        req_status = attrs.get("status") or data.get("status")
        if req_status == "hold":
            if not url:
                return Response(
                    {"errors": [{"detail": "URL is required"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            scrape = Scrape.objects.create(url=url, status="hold", created_by=request.user)
            # Link to existing job post if URL matches
            existing_jp = JobPost.objects.filter(link=url).first()
            if existing_jp:
                scrape.job_post = existing_jp
                scrape.company_id = existing_jp.company_id
                scrape.save()
            logger.info("ScrapeViewSet.create: hold scrape id=%s url=%s", scrape.id, url)
            scr_ser = self.get_serializer()
            return Response(
                {"data": scr_ser.to_resource(scrape)},
                status=status.HTTP_201_CREATED,
            )

        # Check if scraping is enabled
        if not getattr(settings, "SCRAPING_ENABLED", False):
            logger.warning("ScrapeViewSet.create: SCRAPING_ENABLED=False, rejecting request")
            return Response(
                {"errors": [{"detail": "Scraping functionality is disabled"}]},
                status=status.HTTP_501_NOT_IMPLEMENTED,
            )

        if not url:
            logger.warning("ScrapeViewSet.create: missing url in request body")
            return Response(
                {"errors": [{"detail": "URL is required"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

        logger.info("ScrapeViewSet.create url=%s", url)

        # Check for existing scrape with the same URL
        existing_scrape = Scrape.objects.filter(url=url).first()

        # If there's an existing scrape that's pending or completed, return it
        if existing_scrape:
            logger.info(
                "ScrapeViewSet.create: existing scrape found id=%s status=%s",
                existing_scrape.id,
                existing_scrape.status,
            )
            if existing_scrape.status in ("pending", "processing", "running"):
                # Return the existing pending/processing/running scrape
                scr_ser = self.get_serializer()
                scrape_resource = scr_ser.to_resource(existing_scrape)
                return Response(
                    {
                        "data": scrape_resource,
                        "meta": {"message": "Scrape already in progress for this URL"},
                    },
                    status=status.HTTP_200_OK,
                )
            elif existing_scrape.status == "completed":
                # Return the existing completed scrape
                scr_ser = self.get_serializer()
                scrape_resource = scr_ser.to_resource(existing_scrape)
                return Response(
                    {
                        "data": scrape_resource,
                        "meta": {
                            "message": "Scrape already completed for this URL. Use the redo action to re-scrape."
                        },
                    },
                    status=status.HTTP_200_OK,
                )
            # If failed, we'll create a new scrape below
            logger.info("ScrapeViewSet.create: existing scrape status=%s, creating new scrape", existing_scrape.status)

        scrape = Scrape.objects.create(url=url, status="pending", created_by=request.user)
        logger.info("ScrapeViewSet.create: created scrape id=%s", scrape.id)

        # Associate with an existing job post (and its company) if the URL matches
        existing_jp = JobPost.objects.filter(link=url).first()
        if existing_jp:
            scrape.job_post = existing_jp
            scrape.company_id = existing_jp.company_id
            scrape.save()
            logger.info("ScrapeViewSet.create: linked scrape id=%s to job_post id=%s company_id=%s", scrape.id, existing_jp.id, existing_jp.company_id)

        browser_service_url = getattr(settings, "BROWSER_SERVICE_URL", "http://localhost:3012")
        logger.info("ScrapeViewSet.create: dispatching scraper browser_service_url=%s scrape_id=%s", browser_service_url, scrape.id)
        Scraper(browser_service_url, url, scrape_id=scrape.id).dispatch()

        scr_ser = self.get_serializer()
        scrape_resource = scr_ser.to_resource(scrape)
        return Response({"data": scrape_resource}, status=status.HTTP_202_ACCEPTED)

    @extend_schema(
        tags=["Scrapes"],
        summary="Re-scrape a URL (async — resets to pending)",
        responses={
            202: OpenApiResponse(description="Scrape restarted"),
            400: OpenApiResponse(description="Already pending/processing"),
            501: OpenApiResponse(description="Scraping disabled"),
        },
    )
    @action(detail=True, methods=["post"])
    def redo(self, request, pk=None):
        """Redo a scrape - resets status to pending and starts a new scrape process"""

        # Check if scraping is enabled
        if not getattr(settings, "SCRAPING_ENABLED", False):
            logger.warning("ScrapeViewSet.redo: SCRAPING_ENABLED=False, rejecting request")
            return Response(
                {"errors": [{"detail": "Scraping functionality is disabled"}]},
                status=status.HTTP_501_NOT_IMPLEMENTED,
            )

        obj = Scrape.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        logger.info("ScrapeViewSet.redo: id=%s previous_status=%s url=%s", obj.id, obj.status, obj.url)

        # Don't allow redo if already pending or processing
        # if obj.status in ("pending", "processing"):
        #     return Response(
        #         {"errors": [{"detail": f"Scrape is already {obj.status}"}]},
        #         status=status.HTTP_400_BAD_REQUEST,
        #     )

        obj.status = "pending"
        obj.save()

        browser_service_url = getattr(settings, "BROWSER_SERVICE_URL", "http://localhost:3012")
        logger.info("ScrapeViewSet.redo: dispatching browser_service_url=%s scrape_id=%s", browser_service_url, obj.id)
        Scraper(browser_service_url, obj.url, scrape_id=obj.id).dispatch()

        # Return the updated scrape
        scr_ser = self.get_serializer()
        scrape_resource = scr_ser.to_resource(obj)
        return Response(
            {"data": scrape_resource, "meta": {"message": "Scrape restarted"}},
            status=status.HTTP_202_ACCEPTED,
        )

    @extend_schema(
        tags=["Scrapes"],
        summary="Parse scrape content into a JobPost and Company",
        responses={
            200: OpenApiResponse(description="Parsed successfully"),
            404: OpenApiResponse(description="Not found"),
            422: OpenApiResponse(description="No content to parse"),
        },
    )
    @action(detail=True, methods=["post"])
    def parse(self, request, pk=None):
        """Parse a completed scrape's content and create/update the JobPost and Company."""
        obj = Scrape.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        if not (obj.job_content or obj.html):
            return Response(
                {"errors": [{"detail": "Scrape has no content to parse"}]},
                status=status.HTTP_422_UNPROCESSABLE_ENTITY,
            )

        logger.info("ScrapeViewSet.parse: id=%s", obj.id)

        from job_hunting.lib.parsers.generic_parser import GenericParser
        try:
            GenericParser().parse(obj, user=request.user)
        except Exception:
            logger.exception("ScrapeViewSet.parse: failed id=%s", obj.id)
            return Response(
                {"errors": [{"detail": "Parsing failed"}]},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        # Reload to pick up updated job_post_id / company_id
        obj = Scrape.objects.filter(pk=int(pk)).first()
        scr_ser = self.get_serializer()
        scrape_resource = scr_ser.to_resource(obj)
        return Response({"data": scrape_resource})

    @action(detail=True, methods=["get"], url_path="screenshots")
    def screenshots(self, request, pk=None):
        """List screenshot filenames for a scrape. Staff only."""
        if not request.user.is_staff:
            return Response(
                {"errors": [{"detail": "Staff access required"}]},
                status=status.HTTP_403_FORBIDDEN,
            )
        from job_hunting.lib.screenshot_store import ScreenshotStore
        store = ScreenshotStore(settings.SCREENSHOT_DIR)
        files = store.list_for_scrape(int(pk))
        return Response({"data": files})

    @action(
        detail=True,
        methods=["get"],
        url_path="screenshots/(?P<filename>[^/]+)",
        url_name="screenshot-file",
    )
    def screenshot_file(self, request, pk=None, filename=None):
        """Serve a screenshot PNG. Staff only."""
        if not request.user.is_staff:
            return Response(
                {"errors": [{"detail": "Staff access required"}]},
                status=status.HTTP_403_FORBIDDEN,
            )
        from django.http import FileResponse
        from job_hunting.lib.screenshot_store import ScreenshotStore
        store = ScreenshotStore(settings.SCREENSHOT_DIR)
        path = store.read(filename)
        if not path:
            return Response(
                {"errors": [{"detail": "Screenshot not found"}]},
                status=status.HTTP_404_NOT_FOUND,
            )
        return FileResponse(open(path, "rb"), content_type="image/png")


@extend_schema_view(
    list=extend_schema(tags=["Companies"], summary="List companies"),
    retrieve=extend_schema(tags=["Companies"], summary="Retrieve a company"),
    create=extend_schema(tags=["Companies"], summary="Create a company"),
    update=extend_schema(tags=["Companies"], summary="Update a company"),
    partial_update=extend_schema(
        tags=["Companies"], summary="Partially update a company"
    ),
    destroy=extend_schema(tags=["Companies"], summary="Delete a company"),
)
class CompanyViewSet(BaseViewSet):
    model = Company
    serializer_class = CompanySerializer



    def list(self, request):
        qs = Company.objects

        query_filter = request.query_params.get("filter[query]")
        if query_filter is not None:
            qs = qs.filter(
                Q(name__icontains=query_filter) | Q(display_name__icontains=query_filter)
            ).distinct()

        sort_param = request.query_params.get("sort", "relevant")
        if sort_param in ("relevant", "-relevant"):
            qs = qs.annotate(
                latest_job_post=Max(
                    "job_posts__created_at",
                    filter=Q(job_posts__created_by_id=request.user.id),
                )
            ).order_by("-latest_job_post")
        elif sort_param:
            sort_fields = []
            for field in sort_param.split(","):
                field = field.strip()
                if field.startswith("-"):
                    sort_fields.append(f"-{field[1:]}")
                else:
                    sort_fields.append(field)
            if sort_fields:
                qs = qs.order_by(*sort_fields)

        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs.all()[offset: offset + page_size])

        ser = self.get_serializer()
        payload = {
            "data": [ser.to_resource(obj) for obj in items],
            "meta": {
                "total": total,
                "page": page_number,
                "per_page": page_size,
                "total_pages": total_pages,
            },
        }
        if page_number < total_pages:
            base = request.build_absolute_uri(request.path)
            qp = request.query_params.dict()
            qp["page"] = page_number + 1
            qp["per_page"] = page_size
            payload["links"] = {"next": base + "?" + "&".join(f"{k}={v}" for k, v in qp.items())}
        else:
            payload["links"] = {"next": None}

        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = Company.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def create(self, request):
        data = request.data.get("data", {})
        attrs = data.get("attributes", {})
        obj = Company.objects.create(
            name=attrs.get("name", ""),
            display_name=attrs.get("display_name"),
            notes=attrs.get("notes"),
        )
        ser = self.get_serializer()
        return Response({"data": ser.to_resource(obj)}, status=status.HTTP_201_CREATED)

    def update(self, request, pk=None):
        obj = Company.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        data = request.data.get("data", {})
        attrs = data.get("attributes", {})
        for field in ("name", "display_name", "notes"):
            if field in attrs:
                setattr(obj, field, attrs[field])
        obj.save()
        ser = self.get_serializer()
        return Response({"data": ser.to_resource(obj)})

    def partial_update(self, request, pk=None):
        return self.update(request, pk=pk)

    def destroy(self, request, pk=None):
        obj = Company.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        obj.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

    @extend_schema(
        tags=["Companies"],
        summary="List job posts for a company",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"], url_path="job-posts")
    def job_posts(self, request, pk=None):
        if not Company.objects.filter(pk=pk).exists():
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        posts = list(
            JobPost.objects.filter(company_id=int(pk)).filter(
                Q(created_by_id=request.user.id) |
                Q(applications__user_id=request.user.id) |
                Q(scores__user_id=request.user.id)
            ).distinct()
        )
        data = [JobPostSerializer().to_resource(j) for j in posts]
        return Response({"data": data})

    @extend_schema(
        tags=["Companies"],
        summary="List job applications for a company",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"], url_path="job-applications")
    def applications(self, request, pk=None):
        if not Company.objects.filter(pk=pk).exists():
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        apps = list(JobApplication.objects.filter(company_id=int(pk), user_id=request.user.id))
        data = [JobApplicationSerializer().to_resource(a) for a in apps]
        return Response({"data": data})

    @extend_schema(
        tags=["Companies"],
        summary="List scrapes for a company",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def scrapes(self, request, pk=None):
        if not Company.objects.filter(pk=pk).exists():
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        scrapes_list = list(Scrape.objects.filter(company_id=int(pk)))
        data = [ScrapeSerializer().to_resource(s) for s in scrapes_list]
        return Response({"data": data})


@extend_schema_view(
    list=extend_schema(
        tags=["Cover Letters"], summary="List cover letters (authenticated user's only)"
    ),
    retrieve=extend_schema(
        tags=["Cover Letters"],
        summary="Retrieve a cover letter (owner only)",
        responses={
            200: _JSONAPI_ITEM,
            403: OpenApiResponse(description="Forbidden"),
            404: OpenApiResponse(description="Not found"),
        },
    ),
    update=extend_schema(
        tags=["Cover Letters"], summary="Update a cover letter (owner only)"
    ),
    partial_update=extend_schema(
        tags=["Cover Letters"], summary="Partially update a cover letter (owner only)"
    ),
    destroy=extend_schema(
        tags=["Cover Letters"], summary="Delete a cover letter (owner only)"
    ),
)
class CoverLetterViewSet(BaseViewSet):
    model = CoverLetter
    serializer_class = CoverLetterSerializer
    permission_classes = [IsAuthenticated]

    def list(self, request):
        items = list(CoverLetter.objects.filter(user_id=request.user.id))
        items = self.paginate(items)
        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = CoverLetter.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Verify ownership
        if obj.user_id != request.user.id:
            return Response(
                {"errors": [{"detail": "Forbidden"}]},
                status=status.HTTP_403_FORBIDDEN,
            )

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def _upsert(self, request, pk, partial=False):
        obj = CoverLetter.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Verify ownership
        if obj.user_id != request.user.id:
            return Response(
                {"errors": [{"detail": "Forbidden"}]},
                status=status.HTTP_403_FORBIDDEN,
            )

        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)

        # Remove any user_id from attrs to prevent ownership changes
        attrs.pop("user_id", None)

        # If resume_id is being changed, validate new resume ownership
        if attrs.get("resume_id", None) is not None:
            new_resume = Resume.objects.filter(pk=attrs["resume_id"]).first()
            if not new_resume or new_resume.user_id != request.user.id:
                return Response(
                    {"errors": [{"detail": "Forbidden"}]},
                    status=status.HTTP_403_FORBIDDEN,
                )

        for k, v in attrs.items():
            setattr(obj, k, v)
        obj.save()
        return Response({"data": ser.to_resource(obj)})

    def destroy(self, request, pk=None):
        obj = CoverLetter.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response(status=204)

        # Verify ownership
        if obj.user_id != request.user.id:
            return Response(
                {"errors": [{"detail": "Forbidden"}]},
                status=status.HTTP_403_FORBIDDEN,
            )

        obj.delete()
        return Response(status=204)

    @extend_schema(
        tags=["Cover Letters"],
        summary="Create a cover letter (or AI-generate if content omitted — requires resume + job-post relationships)",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            202: OpenApiResponse(description="AI generation started — poll the returned resource for state changes"),
            400: OpenApiResponse(description="Missing/invalid resume or job-post"),
            403: OpenApiResponse(description="Resume not owned by user"),
            503: OpenApiResponse(description="AI client not configured"),
        },
    )
    def create(self, request):
        data = request.data if isinstance(request.data, dict) else {}
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(data)
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        node = data.get("data") or {}
        relationships = node.get("relationships") or {}
        attrs_node = node.get("attributes") or {}

        def _first_id(n):
            if isinstance(n, dict):
                d = n.get("data", n)
                if isinstance(d, dict) and "id" in d:
                    return d.get("id")
            return None

        def _rel_id(*keys):
            for k in keys:
                rel = relationships.get(k)
                if rel is not None:
                    rid = _first_id(rel)
                    if rid is not None:
                        return rid
            return None

        # Resolve IDs from attrs (serializer) or relationships fallbacks
        resume_id = attrs.get("resume_id") or _rel_id("resume", "resumes")
        job_post_id = attrs.get("job_post_id") or _rel_id(
            "job-post", "job_post", "jobPost", "job-posts", "jobPosts"
        )
        company_id = attrs.get("company_id") or _rel_id("company", "companies")

        # Additional fallbacks: accept convenience keys at attributes/top-level
        if job_post_id is None:
            job_post_id = (
                attrs_node.get("job_post_id")
                or attrs_node.get("jobPostId")
                or attrs_node.get("job-post-id")
                or node.get("job_post_id")
                or node.get("jobPostId")
                or node.get("job-post-id")
                or data.get("job_post_id")
                or data.get("jobPostId")
                or data.get("job-post-id")
            )
        if resume_id is None:
            resume_id = (
                attrs_node.get("resume_id")
                or attrs_node.get("resumeId")
                or attrs_node.get("resume-id")
                or node.get("resume_id")
                or node.get("resumeId")
                or node.get("resume-id")
                or data.get("resume_id")
                or data.get("resumeId")
                or data.get("resume-id")
            )
        if company_id is None:
            company_id = (
                attrs_node.get("company_id")
                or attrs_node.get("companyId")
                or attrs_node.get("company-id")
                or node.get("company_id")
                or node.get("companyId")
                or node.get("company-id")
                or data.get("company_id")
                or data.get("companyId")
                or data.get("company-id")
            )

        try:
            resume_id = int(resume_id) if resume_id is not None else None
            job_post_id = int(job_post_id) if job_post_id is not None else None
            company_id = int(company_id) if company_id is not None else None
        except (TypeError, ValueError):
            return Response(
                {"errors": [{"detail": "Invalid resume, job-post, or company ID"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        # id=0 means "no resume" → career-data fallback
        if resume_id == 0:
            resume_id = None

        resume = Resume.get(resume_id) if resume_id is not None else None
        job_post = (
            JobPost.objects.filter(pk=job_post_id).first()
            if job_post_id is not None
            else None
        )
        company = Company.get(company_id) if company_id is not None else None

        if job_post_id is not None and not job_post:
            return Response(
                {"errors": [{"detail": "Invalid job-post ID"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        if resume_id is not None and not resume:
            return Response(
                {"errors": [{"detail": "Invalid resume ID"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        if company_id is not None and not company:
            return Response(
                {"errors": [{"detail": "Invalid company ID"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Validate that the referenced resume belongs to the authenticated user
        if resume and resume.user_id != request.user.id:
            return Response(
                {
                    "errors": [
                        {
                            "detail": "Forbidden: resume does not belong to the authenticated user"
                        }
                    ]
                },
                status=status.HTTP_403_FORBIDDEN,
            )

        # Force ownership to authenticated user
        user_id = request.user.id

        # Accept provided content; otherwise, generate via AI
        content = attrs_node.get("content")
        if content:
            cover_letter = CoverLetter(
                content=content,
                user_id=user_id,
                resume_id=(resume.id if resume else None),
                job_post_id=(job_post.id if job_post else None),
                company_id=(company.id if company else None),
            )
            cover_letter.save()
            payload = {"data": ser.to_resource(cover_letter)}
            include_rels = self._parse_include(request)
            if include_rels:
                payload["included"] = self._build_included(
                    [cover_letter], include_rels, request
                )
            return Response(payload, status=status.HTTP_201_CREATED)

        # AI generation path — create a pending record and dispatch async
        if job_post is None:
            return Response(
                {
                    "errors": [
                        {
                            "detail": "Provide 'relationships.job-post' when generating content without providing attributes.content"
                        }
                    ]
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Build career markdown when no resume provided (mirrors score career-data path)
        career_markdown = None
        if resume is None:
            career_data = CareerData.for_user(user_id)
            prompt_builder = ApplicationPromptBuilder(max_section_chars=60000)
            career_markdown = prompt_builder.build_from_career_data(career_data)
            if not career_markdown.strip():
                return Response(
                    {"errors": [{"detail": "No career data found for this user. Add favorite resumes or provide a resume relationship."}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

        client = get_client(required=False)
        if client is None:
            return Response(
                {"errors": [{"detail": "AI client not configured. Set OPENAI_API_KEY."}]},
                status=503,
            )

        cover_letter = CoverLetter.objects.create(
            user_id=user_id,
            resume_id=(resume.id if resume else None),
            job_post_id=job_post.id,
            company_id=(company.id if company else None),
            status="pending",
        )
        cl_id = cover_letter.id

        def _generate():
            import django
            django.db.close_old_connections()
            try:
                cl_service = CoverLetterService(
                    client, job_post, resume=resume, resume_markdown=career_markdown, user_id=user_id
                )
                result = cl_service.generate_cover_letter()
                CoverLetter.objects.filter(pk=cl_id).update(
                    content=result.content, status="completed"
                )
            except Exception:
                CoverLetter.objects.filter(pk=cl_id).update(status="failed")

        import threading
        threading.Thread(target=_generate, daemon=True).start()

        return Response({"data": ser.to_resource(cover_letter)}, status=status.HTTP_202_ACCEPTED)

    @extend_schema(
        tags=["Cover Letters"],
        summary="Export a cover letter as DOCX",
        responses={
            200: OpenApiResponse(description="DOCX file download"),
            403: OpenApiResponse(description="Forbidden"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    @action(
        detail=True,
        methods=["get"],
        url_path="export",
        permission_classes=[IsAuthenticated],
    )
    def export_docx(self, request, pk=None):
        cl = self.model.get(int(pk))
        if not cl:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Verify ownership
        if cl.user_id != request.user.id:
            return Response(
                {"errors": [{"detail": "Forbidden"}]},
                status=status.HTTP_403_FORBIDDEN,
            )

        try:
            from docx import Document  # python-docx
        except Exception:
            return Response(
                {
                    "errors": [
                        {"detail": "DOCX export requires 'python-docx' to be installed"}
                    ]
                },
                status=status.HTTP_501_NOT_IMPLEMENTED,
            )

        from io import BytesIO

        doc = Document()

        # Title
        title_parts = ["Cover Letter"]
        try:
            if getattr(cl, "job_post", None) and getattr(cl.job_post, "title", None):
                title_parts.append(str(cl.job_post.title))
            if getattr(cl, "job_post", None) and getattr(cl.job_post, "company", None):
                if getattr(cl.job_post.company, "name", None):
                    title_parts.append(str(cl.job_post.company.name))
        except Exception:
            pass
        if title_parts:
            doc.add_heading(" - ".join(title_parts), level=1)

        # Body
        content = (cl.content or "").strip()
        if content:
            import re as _re

            for para in [p for p in _re.split(r"\n\s*\n", content) if p.strip()]:
                p = doc.add_paragraph()
                for line in para.splitlines():
                    if p.text:
                        p.add_run("\n")
                    p.add_run(line)

        # Footer/meta
        try:
            meta_bits = []
            if getattr(cl, "created_at", None):
                meta_bits.append(f"Created: {cl.created_at}")
            if getattr(cl, "user", None) and getattr(cl.user, "name", None):
                meta_bits.append(f"Author: {cl.user.name}")
            if meta_bits:
                doc.add_paragraph("\n".join(meta_bits))
        except Exception:
            pass

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        import re as _re2

        filename_parts = ["cover-letter", str(cl.id)]
        try:
            if getattr(cl, "job_post", None) and getattr(cl.job_post, "company", None):
                nm = getattr(cl.job_post.company, "name", "") or ""
                if nm:
                    filename_parts.append(_re2.sub(r"[^A-Za-z0-9_-]+", "-", nm))
        except Exception:
            pass
        filename = "-".join([p for p in filename_parts if p]) + ".docx"

        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp


@extend_schema_view(
    list=extend_schema(
        tags=["Job Applications"],
        summary="List job applications",
        parameters=_PAGE_PARAMS + [
            _SORT_PARAM,
            _FILTER_APP_QUERY_PARAM,
            _FILTER_APP_STATUS_PARAM,
            _FILTER_COMPANY_PARAM,
            _FILTER_COMPANY_ID_PARAM,
        ],
    ),
    retrieve=extend_schema(
        tags=["Job Applications"], summary="Retrieve a job application"
    ),
    create=extend_schema(
        tags=["Job Applications"],
        summary="Create a job application (user_id auto-set from authenticated user)",
    ),
    update=extend_schema(tags=["Job Applications"], summary="Update a job application"),
    partial_update=extend_schema(
        tags=["Job Applications"], summary="Partially update a job application"
    ),
    destroy=extend_schema(
        tags=["Job Applications"], summary="Delete a job application"
    ),
)
class JobApplicationViewSet(BaseViewSet):
    model = JobApplication
    serializer_class = JobApplicationSerializer

    def list(self, request):
        qs = JobApplication.objects.filter(user_id=request.user.id)

        company_id_filter = request.query_params.get("filter[company_id]")
        if company_id_filter is not None:
            qs = qs.filter(company_id=company_id_filter)

        company_filter = request.query_params.get("filter[company]")
        if company_filter is not None:
            qs = qs.filter(company__name__icontains=company_filter)

        status_filter = request.query_params.get("filter[status]")
        if status_filter is not None:
            qs = qs.filter(status__icontains=status_filter)

        query_filter = request.query_params.get("filter[query]")
        if query_filter is not None:
            qs = qs.filter(
                Q(job_post__title__icontains=query_filter)
                | Q(company__name__icontains=query_filter)
                | Q(company__display_name__icontains=query_filter)
                | Q(status__icontains=query_filter)
                | Q(notes__icontains=query_filter)
            ).distinct()

        # Handle sorting
        sort_param = request.query_params.get("sort")
        if sort_param:
            sort_fields = []
            for field in sort_param.split(","):
                field = field.strip()
                if field.startswith("-"):
                    # Descending order
                    sort_fields.append(f"-{field[1:]}")
                else:
                    # Ascending order
                    sort_fields.append(field)
            if sort_fields:
                qs = qs.order_by(*sort_fields)

        items = list(qs.all())
        items = self.paginate(items)
        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = self._get_obj(pk)
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        # Always include statuses; merge with any additional ?include= rels
        include_rels = list({*self._parse_include(request), "application-statuses"})
        payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def create(self, request):
        response = super().create(request)
        # Auto-create the initial JobApplicationStatus so every application
        # has at least one status entry from the moment it is created.
        if response.status_code == 201:
            app_id = (response.data.get("data") or {}).get("id")
            if app_id:
                app = JobApplication.objects.filter(pk=int(app_id)).first()
                if app and not JobApplicationStatus.objects.filter(application_id=app.id).exists():
                    status_label = app.status or "Unvetted"
                    status_obj, _ = Status.objects.get_or_create(
                        status=status_label,
                        defaults={"status_type": "application"},
                    )
                    from django.utils import timezone
                    JobApplicationStatus.objects.create(
                        application=app,
                        status=status_obj,
                        logged_at=timezone.now(),
                    )
        return response

    def _upsert(self, request, pk, partial=False):
        obj = self._get_obj(pk)
        if not obj or obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        old_status = obj.status
        response = super()._upsert(request, pk, partial=partial)
        if response.status_code == 200:
            obj.refresh_from_db()
            if obj.status and obj.status != old_status:
                from django.utils import timezone
                status_obj, _ = Status.objects.get_or_create(
                    status=obj.status,
                    defaults={"status_type": "application"},
                )
                JobApplicationStatus.objects.create(
                    application=obj,
                    status=status_obj,
                    logged_at=timezone.now(),
                )
        return response

    def pre_save_payload(self, request, attrs, creating):
        """Automatically set user_id and company_id when creating applications"""
        if creating:
            # Set user_id from authenticated user
            attrs["user_id"] = request.user.id

            # Set company_id from job_post if job_post_id is provided
            job_post_id = attrs.get("job_post_id")
            if job_post_id:
                job_post = JobPost.objects.filter(pk=job_post_id).first()
                if job_post and hasattr(job_post, "company_id") and job_post.company_id:
                    attrs["company_id"] = job_post.company_id

        return attrs

    @extend_schema(
        tags=["Job Applications"],
        summary="List application statuses for a job application",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"], url_path="application-statuses")
    def application_statuses(self, request, pk=None):
        app = JobApplication.objects.filter(pk=int(pk)).first()
        if not app or app.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = JobApplicationStatusSerializer()
        items = list(JobApplicationStatus.objects.filter(application_id=int(pk)))
        data = [ser.to_resource(i) for i in items]

        # Build included only when ?include=... is provided
        include_rels = self._parse_include(request)

        payload = {"data": data}
        if include_rels:
            payload["included"] = self._build_included(
                items, include_rels, request, primary_serializer=ser
            )
        return Response(payload)

    @extend_schema(
        tags=["Job Applications"],
        summary="List questions for a job application",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def questions(self, request, pk=None):
        app = JobApplication.objects.filter(pk=int(pk)).first()
        if not app or app.user_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = QuestionSerializer()
        items = list(Question.objects.filter(application_id=int(pk)))
        data = [ser.to_resource(i) for i in items]

        # Build included only when ?include=... is provided
        include_rels = self._parse_include(request)

        payload = {"data": data}
        if include_rels:
            payload["included"] = self._build_included(
                items, include_rels, request, primary_serializer=ser
            )
        return Response(payload)


@extend_schema_view(
    list=extend_schema(tags=["Statuses"], summary="List statuses"),
    retrieve=extend_schema(tags=["Statuses"], summary="Retrieve a status"),
    create=extend_schema(tags=["Statuses"], summary="Create a status"),
    update=extend_schema(tags=["Statuses"], summary="Update a status"),
    partial_update=extend_schema(
        tags=["Statuses"], summary="Partially update a status"
    ),
    destroy=extend_schema(tags=["Statuses"], summary="Delete a status"),
)
class StatusViewSet(viewsets.ModelViewSet):
    queryset = Status.objects.all()
    serializer_class = StatusSerializer

    def create(self, request):
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs = node.get("attributes") or {}
        relationships = node.get("relationships") or {}

        # If a job_application relationship is present, create a JobApplicationStatus
        app_rel = relationships.get("job_application") or relationships.get("job-application")
        app_rel_data = (app_rel or {}).get("data") or {}
        application_id = app_rel_data.get("id")

        if application_id is not None:
            try:
                application_id = int(application_id)
            except (TypeError, ValueError):
                return Response(
                    {"errors": [{"detail": "Invalid job_application id"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            application = JobApplication.objects.filter(pk=application_id).first()
            if not application:
                return Response(
                    {"errors": [{"detail": "Job application not found"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            status_label = attrs.get("status", "").strip()
            if not status_label:
                return Response(
                    {"errors": [{"detail": "attributes.status is required"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )

            status_obj, _ = Status.objects.get_or_create(
                status=status_label,
                defaults={"status_type": "application"},
            )

            from django.utils import timezone
            note = attrs.get("note")
            logged_at_raw = attrs.get("logged_at")
            if logged_at_raw:
                try:
                    from dateutil import parser as dateutil_parser
                    logged_at = dateutil_parser.parse(str(logged_at_raw))
                except (ValueError, TypeError):
                    logged_at = timezone.now()
            else:
                logged_at = timezone.now()

            app_status = JobApplicationStatus.objects.create(
                application=application,
                status=status_obj,
                note=note,
                logged_at=logged_at,
            )

            ser = JobApplicationStatusSerializer()
            return Response(
                {"data": ser.to_resource(app_status)},
                status=status.HTTP_201_CREATED,
            )

        # No job_application relationship — create a plain Status lookup record
        return super().create(request)


@extend_schema_view(
    list=extend_schema(
        tags=["Job Application Statuses"], summary="List job application statuses"
    ),
    retrieve=extend_schema(
        tags=["Job Application Statuses"], summary="Retrieve a job application status"
    ),
    create=extend_schema(
        tags=["Job Application Statuses"], summary="Create a job application status"
    ),
    update=extend_schema(
        tags=["Job Application Statuses"], summary="Update a job application status"
    ),
    partial_update=extend_schema(
        tags=["Job Application Statuses"],
        summary="Partially update a job application status",
    ),
    destroy=extend_schema(
        tags=["Job Application Statuses"], summary="Delete a job application status"
    ),
)
class JobApplicationStatusViewSet(BaseViewSet):
    model = JobApplicationStatus
    serializer_class = JobApplicationStatusSerializer

    def create(self, request):
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs = node.get("attributes") or {}
        relationships = node.get("relationships") or {}

        # Resolve application FK
        app_rel = (
            relationships.get("application")
            or relationships.get("job_application")
            or relationships.get("job-application")
        )
        app_rel_data = (app_rel or {}).get("data") or {}
        application_id = app_rel_data.get("id")
        if application_id is None:
            return Response(
                {"errors": [{"detail": "relationships.application is required"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        try:
            application_id = int(application_id)
        except (TypeError, ValueError):
            return Response(
                {"errors": [{"detail": "Invalid application id"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        application = JobApplication.objects.filter(pk=application_id).first()
        if not application:
            return Response(
                {"errors": [{"detail": "Job application not found"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Resolve status: prefer relationship FK, fall back to text label in attributes
        status_rel = relationships.get("status") or {}
        status_rel_data = (status_rel or {}).get("data") or {}
        status_rel_id = status_rel_data.get("id")

        if status_rel_id is not None:
            status_obj = Status.objects.filter(pk=int(status_rel_id)).first()
        else:
            status_label = (attrs.get("status") or "").strip()
            if not status_label:
                return Response(
                    {"errors": [{"detail": "attributes.status or relationships.status is required"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            status_obj, _ = Status.objects.get_or_create(
                status=status_label,
                defaults={"status_type": "application"},
            )

        from django.utils import timezone
        note = attrs.get("note")
        logged_at_raw = attrs.get("logged_at")
        if logged_at_raw:
            try:
                from dateutil import parser as dateutil_parser
                logged_at = dateutil_parser.parse(str(logged_at_raw))
            except (ValueError, TypeError):
                logged_at = timezone.now()
        else:
            logged_at = timezone.now()

        app_status = JobApplicationStatus.objects.create(
            application=application,
            status=status_obj,
            note=note,
            logged_at=logged_at,
        )

        # Keep the parent application's status field in sync
        application.status = status_obj.status
        application.save(update_fields=["status"])

        ser = JobApplicationStatusSerializer()
        return Response(
            {"data": ser.to_resource(app_status)},
            status=status.HTTP_201_CREATED,
        )


@extend_schema_view(
    update=extend_schema(tags=["Questions"], summary="Update a question"),
    partial_update=extend_schema(
        tags=["Questions"],
        summary="Partially update a question (also appends an answer if attributes.answer provided)",
    ),
    destroy=extend_schema(tags=["Questions"], summary="Delete a question"),
)
class QuestionViewSet(BaseViewSet):
    model = Question
    serializer_class = QuestionSerializer



    @extend_schema(
        tags=["Questions"],
        summary="List questions (auto-includes company)",
        parameters=_PAGE_PARAMS,
        responses={200: _JSONAPI_LIST},
    )
    def list(self, request):
        qs = Question.objects.filter(created_by_id=request.user.id)

        query_filter = request.query_params.get("filter[query]")
        if query_filter:
            qs = qs.filter(content__icontains=query_filter)

        job_post_filter = request.query_params.get("filter[job_post_id]")
        if job_post_filter:
            qs = qs.filter(job_post_id=job_post_filter)

        application_filter = request.query_params.get("filter[application_id]")
        if application_filter:
            qs = qs.filter(application_id=application_filter)

        qs = qs.order_by("-id")
        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size

        include_rels = self._parse_include(request) or ["company"]
        # Prefetch answers in-bulk to avoid N+1 when include=answers is requested
        if "answers" in include_rels or "answer" in include_rels:
            qs = qs.prefetch_related("answers")

        items = list(qs.all()[offset: offset + page_size])

        ser = self.get_serializer()
        payload = {
            "data": [ser.to_resource(o) for o in items],
            "meta": {"total": total, "page": page_number, "per_page": page_size, "total_pages": total_pages},
        }
        if page_number < total_pages:
            base = request.build_absolute_uri(request.path)
            qp = request.query_params.dict()
            qp["page"] = page_number + 1
            qp["per_page"] = page_size
            payload["links"] = {"next": base + "?" + "&".join(f"{k}={v}" for k, v in qp.items())}
        else:
            payload["links"] = {"next": None}

        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["Questions"],
        summary="Retrieve a question (auto-includes company)",
        parameters=[_INCLUDE_PARAM],
        responses={200: _JSONAPI_ITEM, 404: OpenApiResponse(description="Not found")},
    )
    def retrieve(self, request, pk=None):
        include_rels = self._parse_include(request) or ["company"]
        qs = Question.objects
        if "answers" in include_rels or "answer" in include_rels:
            qs = qs.prefetch_related("answers")
        obj = qs.filter(pk=pk).first()
        if not obj or obj.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["Questions"],
        summary="Create a question (optionally include attributes.answer to auto-create an Answer child)",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Validation error"),
        },
    )
    def create(self, request):
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
            attrs["created_by_id"] = request.user.id
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)

        attrs = self.pre_save_payload(request, attrs, creating=True)
        # Backfill company_id and job_post_id from the application if not supplied
        if attrs.get("application_id") and (
            not attrs.get("company_id") or not attrs.get("job_post_id")
        ):
            app = JobApplication.objects.filter(pk=attrs["application_id"]).first()
            if app:
                attrs.setdefault("company_id", app.company_id)
                attrs.setdefault("job_post_id", app.job_post_id)
        # Remove SA-incompatible attrs; keep only model field names
        safe_attrs = {
            k: v
            for k, v in attrs.items()
            if k
            in ("content", "application_id", "company_id", "created_by_id", "job_post_id")
        }
        obj = Question.objects.create(**safe_attrs)

        # Back-compat write path: accept attributes.answer and create a child Answer
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs_node = node.get("attributes") or {}
        ans_val = attrs_node.get("answer")
        ans_str = ans_val.strip() if isinstance(ans_val, str) else None
        if ans_str:
            try:
                Answer.objects.create(question_id=obj.id, content=ans_str)
            except Exception:
                pass

        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)

    def _upsert(self, request, pk, partial=False):
        obj = Question.objects.filter(pk=pk).first()
        if not obj or obj.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)
        attrs = self.pre_save_payload(request, attrs, creating=False)
        for k, v in attrs.items():
            if k in (
                "content",
                "application_id",
                "company_id",
                "created_by_id",
            ):
                setattr(obj, k, v)
        obj.save()

        # Back-compat write path on update: append a new child Answer if provided
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs_node = node.get("attributes") or {}
        ans_val = attrs_node.get("answer")
        ans_str = ans_val.strip() if isinstance(ans_val, str) else None
        if ans_str:
            try:
                Answer.objects.create(question_id=obj.id, content=ans_str)
            except Exception:
                pass

        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def destroy(self, request, pk=None):
        obj = Question.objects.filter(pk=pk).first()
        if not obj or obj.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        obj.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

    @extend_schema(
        tags=["Questions"],
        summary="List answers for a question",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def answers(self, request, pk=None):
        obj = Question.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        items = list(Answer.objects.filter(question_id=obj.id).order_by("created_at"))
        ser = AnswerSerializer()
        data = [ser.to_resource(i) for i in items]

        include_rels = self._parse_include(request)
        payload = {"data": data}
        if include_rels:
            payload["included"] = self._build_included(
                items, include_rels, request, primary_serializer=ser
            )
        return Response(payload)


@extend_schema_view(
    list=extend_schema(tags=["Answers"], summary="List answers"),
    retrieve=extend_schema(tags=["Answers"], summary="Retrieve an answer"),
    update=extend_schema(tags=["Answers"], summary="Update an answer"),
    partial_update=extend_schema(
        tags=["Answers"], summary="Partially update an answer"
    ),
    destroy=extend_schema(tags=["Answers"], summary="Delete an answer"),
)
class AnswerViewSet(BaseViewSet):
    model = Answer
    serializer_class = AnswerSerializer

    @extend_schema(
        tags=["Answers"],
        summary="List answers",
        parameters=_PAGE_PARAMS,
        responses={200: _JSONAPI_LIST},
    )
    def list(self, request):
        qs = Answer.objects.filter(question__created_by_id=request.user.id)

        query_filter = request.query_params.get("filter[query]")
        if query_filter:
            qs = qs.filter(
                Q(content__icontains=query_filter) |
                Q(question__content__icontains=query_filter)
            ).distinct()

        question_filter = request.query_params.get("filter[question_id]")
        if question_filter:
            qs = qs.filter(question_id=question_filter)

        qs = qs.order_by("-id")
        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs.all()[offset: offset + page_size])

        ser = self.get_serializer()
        payload = {
            "data": [ser.to_resource(o) for o in items],
            "meta": {"total": total, "page": page_number, "per_page": page_size, "total_pages": total_pages},
        }
        if page_number < total_pages:
            base = request.build_absolute_uri(request.path)
            qp = request.query_params.dict()
            qp["page"] = page_number + 1
            qp["per_page"] = page_size
            payload["links"] = {"next": base + "?" + "&".join(f"{k}={v}" for k, v in qp.items())}
        else:
            payload["links"] = {"next": None}

        include_rels = self._parse_include(request) or ["question"]
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["Answers"],
        summary="Create an answer (set ai_assist=true to auto-generate content via AI)",
        request=inline_serializer(
            name="AnswerCreateRequest",
            fields={
                "question_id": drf_serializers.IntegerField(
                    help_text="Required. ID of the parent Question."
                ),
                "content": drf_serializers.CharField(
                    required=False,
                    help_text="Answer text. Required unless ai_assist=true.",
                ),
                "ai_assist": drf_serializers.BooleanField(
                    required=False,
                    help_text="If true and content is empty, AI generates the answer.",
                ),
                "injected_prompt": drf_serializers.CharField(
                    required=False,
                    help_text="Optional custom prompt injected into AI generation.",
                ),
            },
        ),
        responses={
            201: _JSONAPI_ITEM,
            202: OpenApiResponse(description="AI generation started — poll the returned resource for state changes"),
            400: OpenApiResponse(description="Missing content or invalid question"),
            503: OpenApiResponse(description="AI client not configured"),
        },
    )
    def create(self, request):
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)

        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs_node = node.get("attributes") or {}
        relationships = node.get("relationships") or {}

        def _first_id(n):
            if isinstance(n, dict):
                d = n.get("data", n)
                if isinstance(d, dict) and "id" in d:
                    return d.get("id")
            return None

        # Resolve question_id from attrs, relationships, or convenience keys
        qid = attrs.get("question_id")
        if qid is None:
            qid = _first_id(
                relationships.get("question") or relationships.get("questions")
            )
        if qid is None:
            qid = (
                attrs_node.get("question_id")
                or attrs_node.get("questionId")
                or attrs_node.get("question-id")
                or node.get("question_id")
                or node.get("questionId")
                or node.get("question-id")
                or data.get("question_id")
                or data.get("questionId")
                or data.get("question-id")
            )
        try:
            qid = int(qid) if qid is not None else None
        except (TypeError, ValueError):
            return Response({"errors": [{"detail": "Invalid question ID"}]}, status=400)

        question = Question.objects.filter(pk=qid).first() if qid is not None else None
        if question is None:
            return Response(
                {"errors": [{"detail": "Missing or invalid question relationship"}]},
                status=400,
            )

        # Determine content, ai_assist flag, and injected_prompt
        content = attrs.get("content")
        if isinstance(content, str):
            content = content.strip()

        ai_flag_raw = (
            attrs_node.get("ai_assist")
            or node.get("ai_assist")
            or data.get("ai_assist")
            or attrs.get("ai_assist")
        )

        # Extract injected prompt for AI assistance
        injected_prompt = (
            attrs_node.get("injected_prompt")
            or attrs_node.get("prompt")
            or node.get("injected_prompt")
            or node.get("prompt")
            or data.get("injected_prompt")
            or data.get("prompt")
            or attrs.get("injected_prompt")
            or attrs.get("prompt")
        )
        if isinstance(injected_prompt, str):
            injected_prompt = injected_prompt.strip()
        else:
            injected_prompt = None

        def _to_bool(v):
            if isinstance(v, bool):
                return v
            if v is None:
                return False
            s = str(v).strip().lower()
            return s in ("1", "true", "yes", "y", "on")

        ai_assist = _to_bool(ai_flag_raw)

        # Resolve optional resume_id — 0 or absent means use career-data
        resume_id_raw = (
            attrs.get("resume_id")
            or _first_id(relationships.get("resume") or relationships.get("resumes"))
            or attrs_node.get("resume_id") or attrs_node.get("resumeId")
            or node.get("resume_id") or data.get("resume_id")
        )
        try:
            resume_id_int = int(resume_id_raw) if resume_id_raw is not None else 0
        except (TypeError, ValueError):
            resume_id_int = 0

        answer_career_markdown = None
        if resume_id_int:
            answer_resume = Resume.objects.filter(pk=resume_id_int).first()
            if not answer_resume:
                return Response({"errors": [{"detail": "Resume not found"}]}, status=400)
            answer_career_markdown = DbExportService().resume_markdown_export(answer_resume)
        else:
            career_data = CareerData.for_user(request.user.id)
            prompt_builder = ApplicationPromptBuilder(max_section_chars=60000)
            answer_career_markdown = prompt_builder.build_from_career_data(career_data) or None

        if not content and ai_assist:
            # AI generation — create a pending record and dispatch async
            client = get_client(required=False)
            if client is None:
                return Response(
                    {"errors": [{"detail": "AI client not configured. Set OPENAI_API_KEY."}]},
                    status=503,
                )

            obj = Answer.objects.create(question_id=question.id, status="pending")
            ans_id = obj.id
            captured_prompt = injected_prompt
            captured_career_markdown = answer_career_markdown

            def _generate():
                import django
                django.db.close_old_connections()
                logger.info("AnswerViewSet._generate: start ans_id=%s question_id=%s", ans_id, question.id)
                try:
                    svc = AnswerService(client)
                    logger.info("AnswerViewSet._generate: calling generate_answer ans_id=%s injected_prompt=%r", ans_id, captured_prompt)
                    result = svc.generate_answer(
                        question=question,
                        save=False,
                        injected_prompt=captured_prompt,
                        career_markdown=captured_career_markdown,
                    )
                    logger.info("AnswerViewSet._generate: got result type=%s ans_id=%s", type(result).__name__, ans_id)
                    generated_content = result.content if isinstance(result, Answer) else str(result or "")
                    logger.info("AnswerViewSet._generate: saving content len=%s ans_id=%s", len(generated_content) if generated_content else 0, ans_id)
                    Answer.objects.filter(pk=ans_id).update(
                        content=generated_content, status="completed"
                    )
                    logger.info("AnswerViewSet._generate: completed ans_id=%s", ans_id)
                except Exception:
                    logger.exception("AnswerViewSet._generate: failed ans_id=%s", ans_id)
                    Answer.objects.filter(pk=ans_id).update(status="failed")

            import threading
            threading.Thread(target=_generate, daemon=True).start()

            return Response({"data": ser.to_resource(obj)}, status=status.HTTP_202_ACCEPTED)

        # Synchronous path — content provided directly
        if not content:
            return Response(
                {
                    "errors": [
                        {
                            "detail": "content is required when ai_assist is not true"
                        }
                    ]
                },
                status=400,
            )
        obj = Answer.objects.create(question_id=question.id, content=content)

        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)

    def retrieve(self, request, pk=None):
        obj = Answer.objects.filter(pk=pk).select_related("question").first()
        if not obj or obj.question.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        return Response({"data": ser.to_resource(obj)})

    def _upsert(self, request, pk, partial=False):
        obj = Answer.objects.filter(pk=pk).select_related("question").first()
        if not obj or obj.question.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response({"errors": [{"detail": str(e)}]}, status=400)
        for k, v in attrs.items():
            if k in ("content", "favorite", "status"):
                setattr(obj, k, v)
        obj.save()
        self._sync_question_favorite(obj.question_id)
        return Response({"data": ser.to_resource(obj)})

    def update(self, request, pk=None):
        return self._upsert(request, pk)

    def partial_update(self, request, pk=None):
        return self._upsert(request, pk, partial=True)

    def destroy(self, request, pk=None):
        obj = Answer.objects.filter(pk=pk).select_related("question").first()
        if not obj or obj.question.created_by_id != request.user.id:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        question_id = obj.question_id
        obj.delete()
        self._sync_question_favorite(question_id)
        return Response(status=204)

    @staticmethod
    def _sync_question_favorite(question_id):
        has_fav = Answer.objects.filter(question_id=question_id, favorite=True).exists()
        Question.objects.filter(pk=question_id).update(favorite=has_fav)


@extend_schema_view(
    list=extend_schema(tags=["Experiences"], summary="List experiences"),
    retrieve=extend_schema(tags=["Experiences"], summary="Retrieve an experience"),
    update=extend_schema(
        tags=["Experiences"],
        summary="Update an experience (also adds resume join links if relationships.resumes provided)",
    ),
    partial_update=extend_schema(
        tags=["Experiences"],
        summary="Partially update an experience (also adds resume join links if relationships.resumes provided)",
    ),
    destroy=extend_schema(tags=["Experiences"], summary="Delete an experience"),
)
class ExperienceViewSet(BaseViewSet):
    model = Experience
    serializer_class = ExperienceSerializer

    @extend_schema(
        tags=["Experiences"],
        summary="List descriptions for an experience",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def descriptions(self, request, pk=None):
        obj = Experience.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = DescriptionSerializer()
        ser.set_parent_context("experience", obj.id, "descriptions")
        _desc_ids = list(
            ExperienceDescription.objects.filter(experience_id=obj.id)
            .order_by("order")
            .values_list("description_id", flat=True)
        )
        _desc_map = {d.id: d for d in Description.objects.filter(pk__in=_desc_ids)}
        items = [_desc_map[did] for did in _desc_ids if did in _desc_map]
        data = [ser.to_resource(d) for d in items]
        return Response({"data": data})

    @extend_schema(
        tags=["Experiences"],
        summary="Create an experience (optionally link to one or more resumes via relationships.resumes)",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Invalid resume IDs"),
        },
    )
    def create(self, request):
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        relationships = node.get("relationships") or {}

        # Accept both "resumes" (list) and "resume" (single)
        res_rel = relationships.get("resumes") or relationships.get("resume") or {}
        resume_ids = []
        if isinstance(res_rel, dict):
            d = res_rel.get("data")
            items = d if isinstance(d, list) else ([d] if isinstance(d, dict) else [])
            for it in items:
                if not isinstance(it, dict):
                    continue
                rid = it.get("id")
                if rid is None:
                    continue
                try:
                    resume_ids.append(int(rid))
                except (TypeError, ValueError):
                    return Response(
                        {"errors": [{"detail": f"Invalid resume id: {rid}"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

        # Validate referenced resumes (if provided)
        invalid = [
            rid for rid in resume_ids if not Resume.objects.filter(pk=rid).exists()
        ]
        if invalid:
            return Response(
                {
                    "errors": [
                        {
                            "detail": f"Invalid resume ID(s): {', '.join(map(str, invalid))}"
                        }
                    ]
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Use existing Experience if client supplies an id; otherwise create a new one
        provided_id = node.get("id")
        exp = None
        if provided_id is not None:
            try:
                exp = Experience.objects.filter(pk=int(provided_id)).first()
            except (TypeError, ValueError):
                exp = None
            if not exp:
                return Response(
                    {"errors": [{"detail": f"Invalid experience ID: {provided_id}"}]},
                    status=status.HTTP_400_BAD_REQUEST,
                )
        else:
            exp = Experience.objects.create(**attrs)

        # Populate join table (avoid duplicates)
        for rid in resume_ids:
            ResumeExperience.objects.get_or_create(resume_id=rid, experience_id=exp.id)

        payload = {"data": ser.to_resource(exp)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([exp], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)

    def _upsert(self, request, pk, partial=False):
        exp = Experience.objects.filter(pk=int(pk)).first()
        if not exp:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        # Update scalar attributes (including company via relationship_fks)
        for k, v in attrs.items():
            setattr(exp, k, v)
        exp.save()

        # Parse resume relationship to add join(s)
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        relationships = node.get("relationships") or {}

        res_rel = relationships.get("resumes") or relationships.get("resume") or {}
        resume_ids = []
        if isinstance(res_rel, dict):
            d = res_rel.get("data")
            items = d if isinstance(d, list) else ([d] if isinstance(d, dict) else [])
            for it in items:
                if not isinstance(it, dict):
                    continue
                rid = it.get("id")
                if rid is None:
                    continue
                try:
                    resume_ids.append(int(rid))
                except (TypeError, ValueError):
                    return Response(
                        {"errors": [{"detail": f"Invalid resume id: {rid}"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

        # Validate resumes and create missing links
        invalid = [
            rid for rid in resume_ids if not Resume.objects.filter(pk=rid).exists()
        ]
        if invalid:
            return Response(
                {
                    "errors": [
                        {
                            "detail": f"Invalid resume ID(s): {', '.join(map(str, invalid))}"
                        }
                    ]
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        for rid in resume_ids:
            ResumeExperience.objects.get_or_create(resume_id=rid, experience_id=exp.id)

        payload = {"data": ser.to_resource(exp)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([exp], include_rels, request)
        return Response(payload)


@extend_schema_view(
    list=extend_schema(tags=["Educations"], summary="List educations"),
    retrieve=extend_schema(tags=["Educations"], summary="Retrieve an education"),
    update=extend_schema(tags=["Educations"], summary="Update an education"),
    partial_update=extend_schema(
        tags=["Educations"], summary="Partially update an education"
    ),
    destroy=extend_schema(tags=["Educations"], summary="Delete an education"),
)
class EducationViewSet(BaseViewSet):
    model = Education
    serializer_class = EducationSerializer



    def list(self, request):
        items = list(Education.objects.all())
        items = self.paginate(items)
        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = Education.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def destroy(self, request, pk=None):
        Education.objects.filter(pk=int(pk)).delete()
        return Response(status=204)

    @extend_schema(
        tags=["Educations"],
        summary="Create an education (optionally link to resumes via relationships.resumes)",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Invalid resume IDs"),
        },
    )
    def create(self, request):
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        relationships = node.get("relationships") or {}

        res_rel = relationships.get("resumes") or relationships.get("resume") or {}
        resume_ids = []
        if isinstance(res_rel, dict):
            d = res_rel.get("data")
            items = d if isinstance(d, list) else ([d] if isinstance(d, dict) else [])
            for it in items:
                if not isinstance(it, dict):
                    continue
                rid = it.get("id")
                if rid is None:
                    continue
                try:
                    resume_ids.append(int(rid))
                except (TypeError, ValueError):
                    return Response(
                        {"errors": [{"detail": f"Invalid resume id: {rid}"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

        invalid = [
            rid for rid in resume_ids if not Resume.objects.filter(pk=rid).exists()
        ]
        if invalid:
            return Response(
                {
                    "errors": [
                        {
                            "detail": f"Invalid resume ID(s): {', '.join(map(str, invalid))}"
                        }
                    ]
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        edu = Education.objects.create(**attrs)

        for rid in resume_ids:
            ResumeEducation.objects.get_or_create(resume_id=rid, education_id=edu.id)

        payload = {"data": ser.to_resource(edu)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([edu], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)


@extend_schema_view(
    list=extend_schema(tags=["Certifications"], summary="List certifications"),
    retrieve=extend_schema(tags=["Certifications"], summary="Retrieve a certification"),
    update=extend_schema(tags=["Certifications"], summary="Update a certification"),
    partial_update=extend_schema(
        tags=["Certifications"], summary="Partially update a certification"
    ),
    destroy=extend_schema(tags=["Certifications"], summary="Delete a certification"),
)
class CertificationViewSet(BaseViewSet):
    model = Certification
    serializer_class = CertificationSerializer



    def list(self, request):
        items = list(Certification.objects.all())
        items = self.paginate(items)
        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = Certification.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def destroy(self, request, pk=None):
        Certification.objects.filter(pk=int(pk)).delete()
        return Response(status=204)

    @extend_schema(
        tags=["Certifications"],
        summary="Create a certification (optionally link to resumes via relationships.resumes)",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Invalid resume IDs"),
        },
    )
    def create(self, request):
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        relationships = node.get("relationships") or {}

        # Accept both "resumes" (list) and "resume" (single)
        res_rel = relationships.get("resumes") or relationships.get("resume") or {}
        resume_ids = []
        if isinstance(res_rel, dict):
            d = res_rel.get("data")
            items = d if isinstance(d, list) else ([d] if isinstance(d, dict) else [])
            for it in items:
                if not isinstance(it, dict):
                    continue
                rid = it.get("id")
                if rid is None:
                    continue
                try:
                    resume_ids.append(int(rid))
                except (TypeError, ValueError):
                    return Response(
                        {"errors": [{"detail": f"Invalid resume id: {rid}"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

        # Validate referenced resumes (if provided)
        invalid = [
            rid for rid in resume_ids if not Resume.objects.filter(pk=rid).exists()
        ]
        if invalid:
            return Response(
                {
                    "errors": [
                        {
                            "detail": f"Invalid resume ID(s): {', '.join(map(str, invalid))}"
                        }
                    ]
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Create Certification
        cert = Certification.objects.create(**attrs)

        # Populate join table
        for rid in resume_ids:
            ResumeCertification.objects.get_or_create(
                resume_id=rid, certification_id=cert.id
            )

        payload = {"data": ser.to_resource(cert)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([cert], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)


@extend_schema_view(
    list=extend_schema(tags=["Descriptions"], summary="List descriptions"),
    retrieve=extend_schema(tags=["Descriptions"], summary="Retrieve a description"),
    update=extend_schema(
        tags=["Descriptions"],
        summary="Update a description (also upserts experience join links with ordering)",
    ),
    partial_update=extend_schema(
        tags=["Descriptions"],
        summary="Partially update a description (also upserts experience join links with ordering)",
    ),
    destroy=extend_schema(tags=["Descriptions"], summary="Delete a description"),
)
class DescriptionViewSet(BaseViewSet):
    model = Description
    serializer_class = DescriptionSerializer



    def list(self, request):
        items = list(Description.objects.all())
        items = self.paginate(items)
        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    def retrieve(self, request, pk=None):
        obj = Description.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    def destroy(self, request, pk=None):
        deleted, _ = Description.objects.filter(pk=int(pk)).delete()
        return Response(status=204)

    @extend_schema(
        tags=["Descriptions"],
        summary="Create a description (optionally link to experiences via relationships.experiences; support per-link order via meta.order)",
        request=_JSONAPI_WRITE,
        responses={
            201: _JSONAPI_ITEM,
            400: OpenApiResponse(description="Invalid experience IDs"),
        },
    )
    def create(self, request):
        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        relationships = node.get("relationships") or {}
        attrs_node = node.get("attributes") or {}
        global_order = attrs_node.get("order")

        exp_rel = (
            relationships.get("experiences") or relationships.get("experience") or {}
        )
        exp_items = []  # list of tuples (experience_id, order)
        if isinstance(exp_rel, dict):
            d = exp_rel.get("data")
            # Accept list or single object
            items = d if isinstance(d, list) else ([d] if isinstance(d, dict) else [])
            for it in items:
                if not isinstance(it, dict):
                    continue
                rid = it.get("id")
                if rid is None:
                    continue
                try:
                    eid = int(rid)
                except (TypeError, ValueError):
                    return Response(
                        {"errors": [{"detail": f"Invalid experience id: {rid}"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )
                order = None
                meta = it.get("meta")
                if isinstance(meta, dict) and "order" in meta:
                    try:
                        order = int(meta.get("order"))
                    except (TypeError, ValueError):
                        order = None
                if order is None and global_order is not None:
                    try:
                        order = int(global_order)
                    except (TypeError, ValueError):
                        order = None
                exp_items.append((eid, order))

        # Validate referenced experiences before creating description
        invalid_ids = [
            eid
            for eid, _ in exp_items
            if not Experience.objects.filter(pk=eid).exists()
        ]
        if invalid_ids:
            return Response(
                {
                    "errors": [
                        {
                            "detail": f"Invalid experience ID(s): {', '.join(map(str, invalid_ids))}"
                        }
                    ]
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Create description
        desc = Description.objects.create(**attrs)

        # Populate join table with optional per-link order
        for eid, order in exp_items:
            ExperienceDescription.objects.get_or_create(
                experience_id=eid,
                description_id=desc.id,
                defaults={"order": (order if order is not None else 0)},
            )

        payload = {"data": ser.to_resource(desc)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([desc], include_rels, request)
        return Response(payload, status=status.HTTP_201_CREATED)

    def _upsert(self, request, pk, partial=False):
        desc = Description.objects.filter(pk=int(pk)).first()
        if not desc:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        ser = self.get_serializer()
        try:
            attrs = ser.parse_payload(request.data)  # handles 'content'
        except ValueError as e:
            return Response(
                {"errors": [{"detail": str(e)}]}, status=status.HTTP_400_BAD_REQUEST
            )

        # Update scalar attributes
        for k, v in attrs.items():
            setattr(desc, k, v)
        desc.save()

        # Handle experiences relationship and per-link 'order'
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        relationships = node.get("relationships") or {}
        attrs_node = node.get("attributes") or {}
        global_order = attrs_node.get("order")

        exp_rel = (
            relationships.get("experiences") or relationships.get("experience") or {}
        )
        if isinstance(exp_rel, dict):
            d = exp_rel.get("data")
            items = d if isinstance(d, list) else ([d] if isinstance(d, dict) else [])
            for it in items:
                if not isinstance(it, dict):
                    continue
                rid = it.get("id")
                if rid is None:
                    continue
                try:
                    eid = int(rid)
                except (TypeError, ValueError):
                    return Response(
                        {"errors": [{"detail": f"Invalid experience id: {rid}"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )
                if not Experience.objects.filter(pk=eid).exists():
                    return Response(
                        {"errors": [{"detail": f"Invalid experience ID: {eid}"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

                # Determine order: meta.order takes precedence; fallback to attributes.order
                order_val = None
                meta = it.get("meta")
                if isinstance(meta, dict) and "order" in meta:
                    try:
                        order_val = int(meta.get("order"))
                    except (TypeError, ValueError):
                        order_val = None
                if order_val is None and global_order is not None:
                    try:
                        order_val = int(global_order)
                    except (TypeError, ValueError):
                        order_val = None

                link, created = ExperienceDescription.objects.get_or_create(
                    experience_id=eid,
                    description_id=desc.id,
                    defaults={"order": (order_val if order_val is not None else 0)},
                )
                if not created and order_val is not None:
                    link.order = order_val
                    link.save()

        payload = {"data": ser.to_resource(desc)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([desc], include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["Descriptions"],
        summary="List experiences linked to a description",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def experiences(self, request, pk=None):
        obj = Description.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        exp_ids = list(
            ExperienceDescription.objects.filter(description_id=obj.id).values_list(
                "experience_id", flat=True
            )
        )
        experiences = list(Experience.objects.filter(pk__in=exp_ids))
        data = [ExperienceSerializer().to_resource(e) for e in experiences]
        return Response({"data": data})


class ApiKeyViewSet(BaseViewSet):
    model = ApiKey
    serializer_class = ApiKeySerializer
    permission_classes = [IsAuthenticated]

    @extend_schema(
        tags=["API Keys"],
        summary="List API keys (staff sees all; others see only their own)",
        parameters=_PAGE_PARAMS,
        responses={200: _JSONAPI_LIST},
    )
    def list(self, request):
        """List API keys - all keys for admins, user's own keys for regular users"""
        if request.user.is_staff:
            items = list(ApiKey.objects.all())
        else:
            items = list(ApiKey.objects.filter(user_id=request.user.id))

        items = self.paginate(items)
        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {"data": data}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(items, include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["API Keys"],
        summary="Retrieve an API key (own key, or any if staff)",
        responses={
            200: _JSONAPI_ITEM,
            403: OpenApiResponse(description="Forbidden"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    def retrieve(self, request, pk=None):
        """Get a specific API key"""
        obj = ApiKey.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Verify ownership or admin access
        if not request.user.is_staff and obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included([obj], include_rels, request)
        return Response(payload)

    @extend_schema(
        tags=["API Keys"],
        summary="Create an API key — plain key returned once in response attributes.key",
        request=inline_serializer(
            name="ApiKeyCreateRequest",
            fields={
                "name": drf_serializers.CharField(
                    help_text="Human-readable name for this key"
                ),
                "expires_days": drf_serializers.IntegerField(
                    required=False, help_text="Days until expiry (omit for no expiry)"
                ),
                "scopes": drf_serializers.ListField(
                    child=drf_serializers.CharField(),
                    required=False,
                    help_text="Defaults to ['read', 'write']",
                ),
            },
        ),
        responses={
            201: OpenApiResponse(
                description="API key created — includes plain key in attributes.key (only shown once)"
            ),
            400: OpenApiResponse(description="Missing name or invalid user"),
        },
    )
    def create(self, request):
        """Create a new API key"""
        data = request.data if isinstance(request.data, dict) else {}
        node = data.get("data") or {}
        attrs = node.get("attributes") or {}
        relationships = node.get("relationships") or {}

        name = attrs.get("name")
        if not name:
            return Response(
                {"errors": [{"detail": "Name is required"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

        expires_days = attrs.get("expires_days")
        scopes = attrs.get("scopes", ["read", "write"])  # Default scopes

        # Determine target user - admins can create keys for other users
        target_user_id = request.user.id  # Default to current user

        # Check if admin is specifying a different user
        user_rel = relationships.get("user") or relationships.get("users")
        if user_rel and request.user.is_staff:
            user_data = user_rel.get("data")
            if isinstance(user_data, dict) and "id" in user_data:
                try:
                    target_user_id = int(user_data["id"])
                    # Verify the target user exists
                    User = get_user_model()
                    if not User.objects.filter(id=target_user_id).exists():
                        return Response(
                            {"errors": [{"detail": "Invalid user ID"}]},
                            status=status.HTTP_400_BAD_REQUEST,
                        )
                except (TypeError, ValueError):
                    return Response(
                        {"errors": [{"detail": "Invalid user ID"}]},
                        status=status.HTTP_400_BAD_REQUEST,
                    )
        elif user_rel and not request.user.is_staff:
            return Response(
                {
                    "errors": [
                        {"detail": "Only admins can create API keys for other users"}
                    ]
                },
                status=status.HTTP_403_FORBIDDEN,
            )

        try:
            api_key_obj, plain_key = ApiKey.generate_key(
                name=name,
                user_id=target_user_id,
                expires_days=expires_days,
                scopes=scopes,
            )
        except Exception as e:
            return Response(
                {"errors": [{"detail": f"Failed to create API key: {str(e)}"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

        ser = self.get_serializer()
        resource = ser.to_resource(api_key_obj)

        # Include the plain key in the response (only time it's available)
        resource["attributes"]["key"] = plain_key

        payload = {"data": resource}
        include_rels = self._parse_include(request)
        if include_rels:
            payload["included"] = self._build_included(
                [api_key_obj], include_rels, request
            )

        return Response(payload, status=status.HTTP_201_CREATED)

    @extend_schema(
        tags=["API Keys"],
        summary="Revoke (delete) an API key",
        responses={
            204: OpenApiResponse(description="Revoked"),
            403: OpenApiResponse(description="Forbidden"),
        },
    )
    def destroy(self, request, pk=None):
        """Revoke an API key"""
        obj = ApiKey.objects.filter(pk=pk).first()
        if not obj:
            return Response(status=204)

        # Verify ownership or admin access
        if not request.user.is_staff and obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)

        obj.delete()
        return Response(status=204)

    @extend_schema(
        tags=["API Keys"],
        summary="Revoke an API key (alternative to DELETE)",
        responses={
            200: _JSONAPI_ITEM,
            403: OpenApiResponse(description="Forbidden"),
            404: OpenApiResponse(description="Not found"),
        },
    )
    @action(detail=True, methods=["post"])
    def revoke(self, request, pk=None):
        """Revoke an API key (alternative to DELETE)"""
        obj = ApiKey.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        # Verify ownership or admin access
        if not request.user.is_staff and obj.user_id != request.user.id:
            return Response({"errors": [{"detail": "Forbidden"}]}, status=403)

        obj.revoke()

        ser = self.get_serializer()
        payload = {"data": ser.to_resource(obj)}
        return Response(payload)


@extend_schema_view(
    list=extend_schema(tags=["Projects"], summary="List projects"),
    retrieve=extend_schema(tags=["Projects"], summary="Retrieve a project"),
    create=extend_schema(tags=["Projects"], summary="Create a project"),
    update=extend_schema(tags=["Projects"], summary="Update a project"),
    partial_update=extend_schema(tags=["Projects"], summary="Partial update a project"),
    destroy=extend_schema(tags=["Projects"], summary="Delete a project"),
)
class ProjectViewSet(BaseViewSet):
    model = Project
    serializer_class = ProjectSerializer

    @extend_schema(
        tags=["Projects"],
        summary="List descriptions for a project",
        responses={200: _JSONAPI_LIST},
    )
    @action(detail=True, methods=["get"])
    def descriptions(self, request, pk=None):
        obj = Project.objects.filter(pk=int(pk)).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        ser = DescriptionSerializer()
        ser.set_parent_context("project", obj.id, "descriptions")
        desc_ids = list(
            ProjectDescription.objects.filter(project_id=obj.id)
            .order_by("order")
            .values_list("description_id", flat=True)
        )
        desc_map = {d.id: d for d in Description.objects.filter(pk__in=desc_ids)}
        items = [desc_map[did] for did in desc_ids if did in desc_map]
        data = [ser.to_resource(d) for d in items]
        return Response({"data": data})


@extend_schema(
    tags=["Career Data"],
    summary="Get aggregated career data as an AI-ready prompt string",
    parameters=[
        OpenApiParameter(
            "user_id",
            OpenApiTypes.INT,
            OpenApiParameter.PATH,
            required=False,
            description="Target user ID (defaults to authenticated user)",
        ),
    ],
    responses={
        200: OpenApiResponse(
            description="Career data formatted as a prompt string",
            response=inline_serializer(
                name="CareerDataResponse",
                fields={"data": drf_serializers.CharField()},
            ),
        )
    },
)
@api_view(["GET"])
@permission_classes([IsAuthenticated])
def career_data(request, user_id=None):
    # Get aggregated career data for the authenticated user or specified user (with API key).
    # Determine which user's data to return
    target_user_id = user_id if user_id is not None else request.user.id

    # If accessing another user's data, ensure proper authorization
    if user_id is not None and user_id != request.user.id:
        # This would be for API key usage - add authorization logic here if needed
        # For now, allow access (you may want to add API key validation)
        pass

    career_data = CareerData.for_user(target_user_id)

    prompt_builder = ApplicationPromptBuilder(max_section_chars=60000)
    career_data_prompt = prompt_builder.build_from_career_data(career_data)
    return Response({"data": career_data_prompt, "meta": career_data.to_refs()})


@extend_schema(
    tags=["Career Data"],
    summary="Generate an AI prompt for answering a job application question",
    request=inline_serializer(
        name="GeneratePromptRequest",
        fields={
            "question_id": drf_serializers.IntegerField(
                help_text="Required. ID of the Question to answer."
            ),
            "job_post_id": drf_serializers.IntegerField(
                required=False, help_text="Optional job post context."
            ),
            "resume_id": drf_serializers.IntegerField(
                required=False, help_text="Optional resume to include."
            ),
            "instructions": drf_serializers.CharField(
                required=False, help_text="Custom instructions appended to the prompt."
            ),
        },
    ),
    responses={
        200: OpenApiResponse(
            description="Generated prompt and context metadata",
            response=inline_serializer(
                name="GeneratePromptResponse",
                fields={
                    "data": inline_serializer(
                        name="GeneratePromptData",
                        fields={
                            "prompt": drf_serializers.CharField(),
                            "context": drf_serializers.DictField(),
                        },
                    )
                },
            ),
        ),
        400: OpenApiResponse(
            description="Missing or invalid question_id / job_post_id / resume_id"
        ),
        403: OpenApiResponse(description="Resume not accessible"),
        404: OpenApiResponse(description="Question not found"),
    },
)
@api_view(["POST"])
@permission_classes([IsAuthenticated])
def generate_prompt(request):
    """
    Generate an AI prompt using ApplicationPromptBuilder for a specific question.

    Expects JSON payload with:
    - question_id: ID of the question to answer (required)
    - job_post_id: (optional) ID of the job post
    - resume_id: (optional) ID of the resume to use
    - instructions: (optional) Custom instructions for the prompt
    """
    data = request.data if isinstance(request.data, dict) else {}

    # Extract required question_id
    question_id = data.get("question_id")
    if not question_id:
        return Response(
            {"errors": [{"detail": "question_id is required"}]},
            status=status.HTTP_400_BAD_REQUEST,
        )

    try:
        question_id = int(question_id)
        question = Question.objects.filter(pk=question_id).first()
        if not question:
            return Response(
                {"errors": [{"detail": "Question not found"}]},
                status=status.HTTP_404_NOT_FOUND,
            )
    except (TypeError, ValueError):
        return Response(
            {"errors": [{"detail": "Invalid question_id"}]},
            status=status.HTTP_400_BAD_REQUEST,
        )

    # Extract optional parameters
    job_post_id = data.get("job_post_id")
    resume_id = data.get("resume_id")
    instructions = data.get("instructions")

    # Use AnswerService to load comprehensive context data
    from job_hunting.lib.services.answer_service import AnswerService

    answer_service = AnswerService(
        ai_client=None
    )  # No AI client needed for data aggregation
    context = answer_service.load_context_for_question(question)

    # Override context with specific parameters if provided
    if job_post_id:
        try:
            job_post_id = int(job_post_id)
            job_post = JobPost.objects.filter(pk=job_post_id).first()
            if job_post:
                context["job_post"] = job_post
                if hasattr(job_post, "company"):
                    context["company"] = job_post.company
        except (TypeError, ValueError):
            return Response(
                {"errors": [{"detail": "Invalid job_post_id"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

    if resume_id:
        try:
            resume_id = int(resume_id)
            resume = Resume.get(resume_id)
            if resume and resume.user_id == request.user.id:
                context["resume"] = resume
                context["resumes"] = [resume]
            elif resume and resume.user_id != request.user.id:
                return Response(
                    {"errors": [{"detail": "Resume not accessible"}]},
                    status=status.HTTP_403_FORBIDDEN,
                )
        except (TypeError, ValueError):
            return Response(
                {"errors": [{"detail": "Invalid resume_id"}]},
                status=status.HTTP_400_BAD_REQUEST,
            )

    # Create prompt builder and generate prompt
    builder = ApplicationPromptBuilder()
    prompt = builder.build(context, instructions)

    return Response(
        {
            "data": {
                "prompt": prompt,
                "context": {
                    "question_id": question.id,
                    "job_post_id": (
                        context.get("job_post").id if context.get("job_post") else None
                    ),
                    "resume_ids": [r.id for r in context.get("resumes", [])],
                    "cover_letter_count": len(context.get("cover_letters", [])),
                    "qa_count": len(context.get("qas", [])),
                },
            }
        }
    )


# ---------------------------------------------------------------------------
# Career-data export / import
# ---------------------------------------------------------------------------


@extend_schema(
    tags=["Career Data"],
    summary="Export career data as an Excel (.xlsx) file",
    responses={
        200: OpenApiResponse(
            description="Excel workbook with sheets: job-posts, job-applications, questions, answers"
        )
    },
)
@api_view(["GET"])
@permission_classes([IsAuthenticated])
def career_data_export(request):
    from openpyxl import Workbook

    wb = Workbook()

    # -- job-posts --
    ws = wb.active
    ws.title = "job-posts"
    jp_headers = [
        "id", "title", "company", "description", "link",
        "posted_date", "extraction_date", "salary_min", "salary_max",
        "location", "remote", "created_at",
    ]
    ws.append(jp_headers)
    for jp in JobPost.objects.select_related("company").order_by("id"):
        ws.append([
            jp.id, jp.title,
            jp.company.name if jp.company else None,
            jp.description, jp.link,
            str(jp.posted_date) if jp.posted_date else None,
            str(jp.extraction_date) if jp.extraction_date else None,
            float(jp.salary_min) if jp.salary_min is not None else None,
            float(jp.salary_max) if jp.salary_max is not None else None,
            jp.location,
            jp.remote,
            jp.created_at.isoformat() if jp.created_at else None,
        ])

    # -- job-applications --
    ws2 = wb.create_sheet("job-applications")
    ja_headers = [
        "id", "job_post_id", "company", "status",
        "applied_at", "tracking_url", "notes",
    ]
    ws2.append(ja_headers)
    for ja in JobApplication.objects.select_related("company").order_by("id"):
        ws2.append([
            ja.id, ja.job_post_id,
            ja.company.name if ja.company else None,
            ja.status,
            ja.applied_at.isoformat() if ja.applied_at else None,
            ja.tracking_url, ja.notes,
        ])

    # -- questions --
    ws3 = wb.create_sheet("questions")
    q_headers = ["id", "application_id", "company", "job_post_id", "content", "favorite", "created_at"]
    ws3.append(q_headers)
    for q in Question.objects.select_related("company").order_by("id"):
        ws3.append([
            q.id, q.application_id,
            q.company.name if q.company else None,
            q.job_post_id, q.content, q.favorite,
            q.created_at.isoformat() if q.created_at else None,
        ])

    # -- answers --
    ws4 = wb.create_sheet("answers")
    a_headers = ["id", "question_id", "content", "favorite", "status", "created_at"]
    ws4.append(a_headers)
    for a in Answer.objects.order_by("id"):
        ws4.append([
            a.id, a.question_id, a.content, a.favorite, a.status,
            a.created_at.isoformat() if a.created_at else None,
        ])

    from io import BytesIO

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = 'attachment; filename="career-caddy-export.xlsx"'
    return response


@extend_schema(
    tags=["Career Data"],
    summary="Import career data from an Excel (.xlsx) file",
    request={"multipart/form-data": {"type": "object", "properties": {"file": {"type": "string", "format": "binary"}}}},
    responses={
        200: OpenApiResponse(description="Import summary with created/skipped counts"),
        400: OpenApiResponse(description="Missing or invalid file"),
        403: OpenApiResponse(description="Superuser access required"),
    },
)
@api_view(["POST"])
@permission_classes([IsAuthenticated])
def career_data_import(request):
    if not request.user.is_superuser:
        return Response(
            {"errors": [{"detail": "Superuser access required"}]},
            status=status.HTTP_403_FORBIDDEN,
        )

    uploaded = request.FILES.get("file")
    if not uploaded:
        return Response(
            {"errors": [{"detail": "No file provided. Upload an .xlsx file as 'file'."}]},
            status=status.HTTP_400_BAD_REQUEST,
        )

    from openpyxl import load_workbook
    from django.db import transaction
    from datetime import datetime

    try:
        wbook = load_workbook(uploaded, read_only=True)
    except Exception:
        return Response(
            {"errors": [{"detail": "Could not read file as .xlsx"}]},
            status=status.HTTP_400_BAD_REQUEST,
        )

    stats = {"job-posts": {"created": 0, "skipped": 0},
             "job-applications": {"created": 0, "skipped": 0},
             "questions": {"created": 0, "skipped": 0},
             "answers": {"created": 0, "skipped": 0}}

    def _parse_datetime(val):
        if val is None:
            return None
        if isinstance(val, datetime):
            return val
        try:
            return datetime.fromisoformat(str(val))
        except (ValueError, TypeError):
            return None

    def _parse_date(val):
        if val is None:
            return None
        from datetime import date
        if isinstance(val, date):
            return val
        try:
            return date.fromisoformat(str(val))
        except (ValueError, TypeError):
            return None

    def _rows_as_dicts(sheet):
        rows = sheet.iter_rows(values_only=True)
        headers = next(rows, None)
        if not headers:
            return
        headers = [str(h).strip() if h else "" for h in headers]
        for row in rows:
            yield dict(zip(headers, row))

    with transaction.atomic():
        # Map old IDs → new IDs for relational integrity
        jp_id_map = {}  # old job_post id → new job_post id
        ja_id_map = {}  # old application id → new application id
        q_id_map = {}   # old question id → new question id

        # -- job-posts --
        if "job-posts" in wbook.sheetnames:
            from job_hunting.models import Company
            for row in _rows_as_dicts(wbook["job-posts"]):
                old_id = row.get("id")
                link = row.get("link")
                company = None
                if row.get("company"):
                    company, _ = Company.objects.get_or_create(name=row["company"])
                # Skip duplicate by link (unique constraint)
                existing = None
                if link:
                    existing = JobPost.objects.filter(link=link).first()
                # Fallback: match on title + company + created_by for linkless posts
                if not existing and row.get("title"):
                    existing = JobPost.objects.filter(
                        title=row["title"], company=company, created_by=request.user
                    ).first()
                if existing:
                    if old_id is not None:
                        jp_id_map[int(old_id)] = existing.id
                    stats["job-posts"]["skipped"] += 1
                    continue
                from decimal import Decimal
                jp = JobPost.objects.create(
                    title=row.get("title"),
                    company=company,
                    description=row.get("description"),
                    link=link,
                    posted_date=_parse_date(row.get("posted_date")),
                    extraction_date=_parse_date(row.get("extraction_date")),
                    salary_min=Decimal(str(row["salary_min"])) if row.get("salary_min") is not None else None,
                    salary_max=Decimal(str(row["salary_max"])) if row.get("salary_max") is not None else None,
                    location=row.get("location"),
                    remote=row.get("remote"),
                    created_by=request.user,
                )
                if old_id is not None:
                    jp_id_map[int(old_id)] = jp.id
                stats["job-posts"]["created"] += 1

        # -- job-applications --
        if "job-applications" in wbook.sheetnames:
            from job_hunting.models import Company
            for row in _rows_as_dicts(wbook["job-applications"]):
                old_id = row.get("id")
                old_jp_id = row.get("job_post_id")
                new_jp_id = jp_id_map.get(int(old_jp_id)) if old_jp_id is not None else None
                # Skip duplicate: same user + same job_post
                if new_jp_id and JobApplication.objects.filter(
                    user=request.user, job_post_id=new_jp_id
                ).exists():
                    if old_id is not None:
                        existing = JobApplication.objects.filter(
                            user=request.user, job_post_id=new_jp_id
                        ).first()
                        ja_id_map[int(old_id)] = existing.id
                    stats["job-applications"]["skipped"] += 1
                    continue
                company = None
                if row.get("company"):
                    company, _ = Company.objects.get_or_create(name=row["company"])
                ja = JobApplication.objects.create(
                    user=request.user,
                    job_post_id=new_jp_id,
                    company=company,
                    status=row.get("status"),
                    applied_at=_parse_datetime(row.get("applied_at")),
                    tracking_url=row.get("tracking_url"),
                    notes=row.get("notes"),
                )
                if old_id is not None:
                    ja_id_map[int(old_id)] = ja.id
                stats["job-applications"]["created"] += 1

        # -- questions --
        if "questions" in wbook.sheetnames:
            from job_hunting.models import Company
            for row in _rows_as_dicts(wbook["questions"]):
                old_id = row.get("id")
                old_app_id = row.get("application_id")
                new_app_id = ja_id_map.get(int(old_app_id)) if old_app_id is not None else None
                old_jp_id = row.get("job_post_id")
                new_jp_id = jp_id_map.get(int(old_jp_id)) if old_jp_id is not None else None
                content = row.get("content")
                # Skip duplicate: same content + same application
                if content and new_app_id and Question.objects.filter(
                    content=content, application_id=new_app_id
                ).exists():
                    if old_id is not None:
                        existing = Question.objects.filter(
                            content=content, application_id=new_app_id
                        ).first()
                        q_id_map[int(old_id)] = existing.id
                    stats["questions"]["skipped"] += 1
                    continue
                company = None
                if row.get("company"):
                    company, _ = Company.objects.get_or_create(name=row["company"])
                q = Question.objects.create(
                    application_id=new_app_id,
                    company=company,
                    created_by=request.user,
                    job_post_id=new_jp_id,
                    content=content,
                    favorite=bool(row.get("favorite")),
                )
                if old_id is not None:
                    q_id_map[int(old_id)] = q.id
                stats["questions"]["created"] += 1

        # -- answers --
        if "answers" in wbook.sheetnames:
            for row in _rows_as_dicts(wbook["answers"]):
                old_q_id = row.get("question_id")
                new_q_id = q_id_map.get(int(old_q_id)) if old_q_id is not None else None
                if new_q_id is None:
                    stats["answers"]["skipped"] += 1
                    continue
                content = row.get("content")
                # Skip duplicate: same content + same question
                if content and Answer.objects.filter(
                    content=content, question_id=new_q_id
                ).exists():
                    stats["answers"]["skipped"] += 1
                    continue
                Answer.objects.create(
                    question_id=new_q_id,
                    content=content,
                    favorite=bool(row.get("favorite")),
                    status=row.get("status"),
                )
                stats["answers"]["created"] += 1

    wbook.close()
    return Response({"data": stats})


class AiUsageViewSet(BaseViewSet):
    model = AiUsage
    serializer_class = AiUsageSerializer

    def _base_queryset(self, request):
        if request.user.is_staff:
            qs = AiUsage.objects.all()
            user_id = request.query_params.get("user_id")
            if user_id:
                qs = qs.filter(user_id=int(user_id))
            return qs
        return AiUsage.objects.filter(user=request.user)

    def _apply_filters(self, qs, request):
        for field in ("agent_name", "model_name", "trigger"):
            val = request.query_params.get(field)
            if val:
                qs = qs.filter(**{field: val})
        pipeline_run_id = request.query_params.get("pipeline_run_id")
        if pipeline_run_id:
            qs = qs.filter(pipeline_run_id=pipeline_run_id)
        date_from = request.query_params.get("date_from")
        if date_from:
            qs = qs.filter(created_at__gte=date_from)
        date_to = request.query_params.get("date_to")
        if date_to:
            qs = qs.filter(created_at__lte=date_to)
        return qs

    def list(self, request):
        qs = self._base_queryset(request)
        qs = self._apply_filters(qs, request)
        qs = qs.order_by("-created_at")

        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs[offset:offset + page_size])

        ser = self.get_serializer()
        data = [ser.to_resource(o) for o in items]
        payload = {
            "data": data,
            "meta": {
                "total": total,
                "page": page_number,
                "per_page": page_size,
                "total_pages": total_pages,
            },
            "links": {"next": None},
        }

        if page_number < total_pages:
            base = request.build_absolute_uri(request.path)
            qp = request.query_params.dict()
            qp["page"] = page_number + 1
            qp["per_page"] = page_size
            payload["links"]["next"] = base + "?" + "&".join(
                f"{k}={v}" for k, v in qp.items()
            )

        return Response(payload)

    def create(self, request):
        from job_hunting.lib.pricing import estimate_cost

        data = request.data if isinstance(request.data, dict) else {}
        ser = self.get_serializer()

        if "data" in data:
            try:
                attrs = ser.parse_payload(request.data)
            except ValueError as e:
                return Response({"errors": [{"detail": str(e)}]}, status=400)
        else:
            attrs = {k: data[k] for k in data if k in {
                "agent_name", "model_name", "trigger", "pipeline_run_id",
                "request_tokens", "response_tokens", "total_tokens", "request_count",
            }}

        attrs["user_id"] = request.user.id
        attrs.pop("estimated_cost_usd", None)
        attrs["estimated_cost_usd"] = estimate_cost(
            attrs.get("model_name", ""),
            int(attrs.get("request_tokens", 0)),
            int(attrs.get("response_tokens", 0)),
        )

        obj = AiUsage.objects.create(**attrs)
        return Response({"data": ser.to_resource(obj)}, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=["get"])
    def summary(self, request):
        from django.db.models import Sum, Count
        from django.db.models.functions import TruncDay, TruncWeek, TruncMonth

        qs = self._base_queryset(request)
        qs = self._apply_filters(qs, request)

        days = int(request.query_params.get("days", 30))
        from django.utils import timezone as tz
        cutoff = tz.now() - tz.timedelta(days=days)
        qs = qs.filter(created_at__gte=cutoff)

        period = request.query_params.get("period", "daily")
        trunc_fn = {"daily": TruncDay, "weekly": TruncWeek, "monthly": TruncMonth}.get(
            period, TruncDay
        )

        group_by = request.query_params.get("group_by")
        valid_groups = {"agent_name", "model_name", "trigger"}
        group_field = group_by if group_by in valid_groups else None

        values_fields = ["period"]
        if group_field:
            values_fields.append(group_field)

        buckets = (
            qs.annotate(period=trunc_fn("created_at"))
            .values(*values_fields)
            .annotate(
                total_tokens=Sum("total_tokens"),
                request_tokens=Sum("request_tokens"),
                response_tokens=Sum("response_tokens"),
                estimated_cost_usd=Sum("estimated_cost_usd"),
                request_count=Count("id"),
            )
            .order_by("period")
        )

        bucket_list = []
        for b in buckets:
            entry = {
                "period": b["period"].isoformat() if b["period"] else None,
                "total_tokens": b["total_tokens"] or 0,
                "request_tokens": b["request_tokens"] or 0,
                "response_tokens": b["response_tokens"] or 0,
                "estimated_cost_usd": str(b["estimated_cost_usd"] or 0),
                "request_count": b["request_count"],
            }
            if group_field:
                entry[group_field] = b[group_field]
            bucket_list.append(entry)

        totals = qs.aggregate(
            total_tokens=Sum("total_tokens"),
            request_tokens=Sum("request_tokens"),
            response_tokens=Sum("response_tokens"),
            estimated_cost_usd=Sum("estimated_cost_usd"),
            request_count=Count("id"),
        )

        return Response({
            "data": {
                "buckets": bucket_list,
                "totals": {
                    "total_tokens": totals["total_tokens"] or 0,
                    "request_tokens": totals["request_tokens"] or 0,
                    "response_tokens": totals["response_tokens"] or 0,
                    "estimated_cost_usd": str(totals["estimated_cost_usd"] or 0),
                    "request_count": totals["request_count"],
                },
            },
        })


class WaitlistViewSet(BaseViewSet):
    model = Waitlist
    serializer_class = WaitlistSerializer
    permission_classes = [IsAuthenticated, IsAdminUser]

    def list(self, request):
        qs = Waitlist.objects.all().order_by("-created_at")
        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs[offset: offset + page_size])
        ser = self.get_serializer()
        return Response({
            "data": [ser.to_resource(o) for o in items],
            "meta": {"total": total, "page": page_number, "per_page": page_size, "total_pages": total_pages},
        })

    def destroy(self, request, pk=None):
        obj = Waitlist.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        obj.delete()
        return Response(status=204)


class InvitationViewSet(BaseViewSet):
    model = Invitation
    serializer_class = InvitationSerializer
    permission_classes = [IsAuthenticated, IsAdminUser]

    def list(self, request):
        qs = Invitation.objects.all().order_by("-created_at")
        total = qs.count()
        page_number, page_size = self._page_params()
        total_pages = math.ceil(total / page_size) if page_size else 1
        offset = (page_number - 1) * page_size
        items = list(qs[offset: offset + page_size])
        ser = self.get_serializer()
        return Response({
            "data": [ser.to_resource(o) for o in items],
            "meta": {"total": total, "page": page_number, "per_page": page_size, "total_pages": total_pages},
        })

    def create(self, request):
        import json
        import secrets
        from django.utils import timezone as tz
        from django.core.mail import send_mail
        from django.template.loader import render_to_string

        try:
            payload = request.data
            if isinstance(payload, bytes):
                payload = json.loads(payload.decode("utf-8") or "{}")
        except Exception:
            payload = {}

        # Support both JSON:API envelope and flat body
        if "data" in payload and isinstance(payload["data"], dict):
            attrs = payload["data"].get("attributes", {})
        else:
            attrs = payload

        email = (attrs.get("email") or "").strip().lower()
        if not email or "@" not in email:
            return Response(
                {"errors": [{"detail": "A valid email address is required."}]},
                status=400,
            )

        token = secrets.token_urlsafe(32)
        expires_at = tz.now() + tz.timedelta(days=7)

        invitation = Invitation.objects.create(
            email=email,
            token=token,
            created_by=request.user,
            expires_at=expires_at,
        )

        # Auto-remove matching waitlist entry
        Waitlist.objects.filter(email=email).delete()

        # Send invitation email
        try:
            invite_url = f"{settings.FRONTEND_URL}/accept-invite?token={token}"
            body = render_to_string(
                "invitation_email.txt", {"invite_url": invite_url}
            )
            send_mail(
                subject="You're invited to Career Caddy",
                message=body,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[email],
            )
        except Exception:
            logger.warning("Failed to send invitation email to %s", email)

        ser = self.get_serializer()
        return Response(
            {"data": ser.to_resource(invitation)}, status=status.HTTP_201_CREATED
        )

    @action(detail=True, methods=["post"], url_path="resend")
    def resend(self, request, pk=None):
        from django.utils import timezone as tz
        from django.core.mail import send_mail
        from django.template.loader import render_to_string

        obj = Invitation.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)

        if obj.is_accepted:
            return Response(
                {"errors": [{"detail": "Invitation already accepted."}]}, status=400
            )

        # Reset expiry if expired
        if obj.is_expired:
            obj.expires_at = tz.now() + tz.timedelta(days=7)
            obj.save()

        invite_url = f"{settings.FRONTEND_URL}/accept-invite?token={obj.token}"
        body = render_to_string(
            "invitation_email.txt", {"invite_url": invite_url}
        )
        try:
            send_mail(
                subject="You're invited to Career Caddy",
                message=body,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[obj.email],
            )
        except Exception:
            logger.warning("Failed to resend invitation email to %s", obj.email)
            return Response(
                {"errors": [{"detail": "Failed to send email."}]}, status=502
            )

        ser = self.get_serializer()
        return Response({"data": ser.to_resource(obj)})

    def destroy(self, request, pk=None):
        obj = Invitation.objects.filter(pk=pk).first()
        if not obj:
            return Response({"errors": [{"detail": "Not found"}]}, status=404)
        obj.delete()
        return Response(status=204)


def _create_user_from_data(username, password, email, first_name="", last_name=""):
    """Shared user creation logic for registration and invitation acceptance.

    Returns (user, error_messages, status_code).
    If error_messages is not None, user creation failed.
    """
    from django.contrib.auth.password_validation import validate_password
    from django.core.exceptions import ValidationError
    from job_hunting.models import Profile

    User = get_user_model()

    if not username:
        return None, [{"detail": "Username is required."}], 400
    if not password:
        return None, [{"detail": "Password is required."}], 400

    if User.objects.filter(username=username).exists():
        return None, [{"detail": "Username already exists."}], 400
    if email and User.objects.filter(email__iexact=email).exists():
        return None, [{"detail": "An account with this email already exists."}], 400

    try:
        validate_password(password)
    except ValidationError as e:
        return None, [{"detail": msg} for msg in e.messages], 400

    user = User(
        username=username, email=email,
        first_name=first_name, last_name=last_name,
    )
    user.set_password(password)
    user.save()

    Profile.objects.get_or_create(user_id=user.id)

    _notify_admins_new_signup(username, email, method="registration")

    return user, None, None


def _notify_admins_new_signup(username, email, method="registration"):
    """Email all superusers when someone signs up."""
    from django.core.mail import send_mail
    from django.template.loader import render_to_string
    from django.utils import timezone as tz

    User = get_user_model()
    admin_emails = list(
        User.objects.filter(is_superuser=True)
        .exclude(email="")
        .values_list("email", flat=True)
    )
    if not admin_emails:
        return

    try:
        body = render_to_string(
            "admin_new_signup.txt",
            {
                "username": username,
                "email": email or "(none)",
                "method": method,
                "timestamp": tz.now().strftime("%Y-%m-%d %H:%M UTC"),
            },
        )
        send_mail(
            subject=f"Career Caddy: new signup — {username}",
            message=body,
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=admin_emails,
        )
    except Exception:
        logger.warning("Failed to send admin signup notification for %s", username)


@csrf_exempt
def accept_invite(request):
    """Accept an invitation and create an account."""
    if request.method != "POST":
        return JsonResponse({"error": "method not allowed"}, status=405)

    import json
    from django.utils import timezone as tz
    from django.core.mail import send_mail
    from django.template.loader import render_to_string

    try:
        data = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        data = {}

    token = (data.get("token") or "").strip()
    username = (data.get("username") or "").strip()
    password = data.get("password", "")
    first_name = (data.get("first_name") or "").strip()
    last_name = (data.get("last_name") or "").strip()

    if not token:
        return JsonResponse(
            {"errors": [{"detail": "Invitation token is required."}]}, status=400
        )

    invitation = Invitation.objects.filter(token=token).first()
    if not invitation or not invitation.is_valid:
        return JsonResponse(
            {"errors": [{"detail": "Invalid or expired invitation link."}]},
            status=400,
        )

    user, errors, error_status = _create_user_from_data(
        username=username,
        password=password,
        email=invitation.email,
        first_name=first_name,
        last_name=last_name,
    )
    if errors:
        return JsonResponse({"errors": errors}, status=error_status)

    invitation.accepted_at = tz.now()
    invitation.save()

    # Send welcome email
    try:
        login_url = f"{settings.FRONTEND_URL}/login"
        body = render_to_string(
            "welcome_email.txt",
            {"username": username, "login_url": login_url},
        )
        send_mail(
            subject="Welcome to Career Caddy",
            message=body,
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=[invitation.email],
        )
    except Exception:
        logger.warning("Failed to send welcome email to %s", invitation.email)

    ser = DjangoUserSerializer()
    return JsonResponse({"data": ser.to_resource(user)}, status=201)


@api_view(["POST"])
@permission_classes([IsAuthenticated, IsAdminUser])
def test_email(request):
    """Admin-only endpoint to verify email delivery."""
    from django.utils import timezone as tz
    from django.core.mail import send_mail
    from django.template.loader import render_to_string

    data = request.data or {}

    target_email = (data.get("email") or "").strip().lower()
    if not target_email:
        target_email = request.user.email

    if not target_email or "@" not in target_email:
        return Response(
            {"errors": [{"detail": "No valid email address available."}]},
            status=400,
        )

    body = render_to_string(
        "test_email.txt", {"timestamp": tz.now().isoformat()}
    )
    send_mail(
        subject="Test Email — Career Caddy",
        message=body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        recipient_list=[target_email],
    )

    return Response(
        {"message": f"Test email sent to {target_email}."}, status=200
    )
