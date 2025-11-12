from docx import Document
import re


class DocxParser:
    def __init__(self, docpath):
        self.path = docpath
        self.name = None
        self.phone = None
        self.email = None
        self.title = None
        self.document = Document(
            self.path or "Doug Headley -- Security Software Engineer.docx"
        )

    def to_markdown(self):
        header_text = []
        for section in self.document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                header_text.append(paragraph.text)
        markdown_content = "\n".join(header_text)
        # Extract and process each section
        current_section = None
        section_content = []
        current_section = None

        prior_val = None
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            if text in self._sections():  # a title
                # If there was a prior section, append it to the markdown content
                if current_section:
                    markdown_content += (
                        f"## {current_section}\n\n"
                        + "\n".join(section_content)
                        + "\n\n"
                    )

                # Start a new section
                current_section = text
                section_content = []
                experience_titles = {}
            else:
                # Otherwise, collect the content of the current section
                if current_section in ["Experience", "EXPERIENCE"]:
                    if self.is_bullet_paragraph(prior_val) and self.is_header(
                        paragraph
                    ):
                        section_content.append("")
                        text = f"### {text}"
                    if self.is_header(prior_val) and self.is_header(paragraph):
                        text = f"#### {text}"
                    if self.is_bullet_paragraph(paragraph):
                        text = f"- {text}"
                    # Otherwise, collect the content of the current section
                section_content.append(text)
            prior_val = paragraph
        # Append the last gathered section if any
        if current_section:
            markdown_content += f"## {current_section}\n\n"
            if current_section == "Experience":
                if section_content:
                    experience_titles.append((section_content[0], section_content[1:]))
                for title, content in experience_titles:
                    markdown_content += f"### {title}\n\n" + "\n".join(content) + "\n\n"
            else:
                markdown_content += "\n".join(section_content) + "\n\n"

        # Print or save the Markdown content
        print(markdown_content)
        return markdown_content

    def _extract_header(self):
        import re

        if getattr(self, "_header_extracted", False):
            return
        for para in self.document.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue

            m = re.match(r"(?i)name\s*[:\-]?\s*(.*)$", text)
            if m:
                self.name = m.group(1).strip() or None
            m = re.match(r"(?i)title\s*[:\-]?\s*(.*)$", text)
            if m:
                self.title = m.group(1).strip() or None
            m = re.match(r"(?i)phone\s*[:\-]?\s*(.*)$", text)
            if m:
                self.phone = m.group(1).strip() or None
            m = re.match(r"(?i)email\s*[:\-]?\s*(.*)$", text)
            if m:
                self.email = m.group(1).strip() or None
        self._header_extracted = True

    def to_html(self):
        import mammoth  # type: ignore

        with open(self.path, "rb") as f:
            result = mammoth.convert_to_html(f)
        return result

    def _sections(self):
        return [
            "Summary",
            "SUMMARY",
            "Education",
            "EDUCATION",
            "Experience",
            "EXPERIENCE",
            "Skills",
            "SKILLS",
            "Certifications",
            "CERTIFICATIONS",
        ]

    def is_bullet_paragraph(self, paragraph):
        """Determine if a paragraph has a list bullet or number."""
        if paragraph is None:
            return False
        list_style_regex = re.compile(r"(?i)list|bullet|number")
        paragraph_style_name = paragraph.style.name.lower() if paragraph.style else ""
        if list_style_regex.search(paragraph_style_name):
            return True

        for run in paragraph.runs:
            run_style_name = run.style.name.lower() if run.style else ""
            if list_style_regex.search(run_style_name):
                return True

        return False

    def is_header(self, paragraph):
        # Header in this instance is the heading of an experience
        return not self.is_bullet_paragraph(paragraph)
