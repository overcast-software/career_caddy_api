"""Render a resume docx programmatically from the smart-Python context.

No jinja template, no docxtpl. The .docx is built paragraph-by-paragraph
with python-docx using the context produced by
``resume_export_context.build_context``. Eliminates the class of
"Encountered unknown tag 'endif'" errors we hit with docxtpl, because
there are no template tags to split across Word runs — all branching
lives in Python.

Styling:
- Uses python-docx's built-in document defaults; call sites that want a
  branded look can pass ``base_template_path`` to a .docx whose *styles*
  (fonts, heading colors, bullet style) are preloaded. This function
  reuses the base's style table; it never copies or merges the base's
  content. That gives design control without needing jinja in Word.
"""
from __future__ import annotations

from io import BytesIO
from typing import Optional

from docx import Document

from job_hunting.models import Resume
from job_hunting.lib.services.resume_export_context import build_context


def _add_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    # Try a named style if the base template carries one; fall back to
    # inline bold so plain blank templates still look sensible.
    try:
        p.style = doc.styles["Heading 2"]
    except KeyError:
        pass


def _add_plain(doc, text: str) -> None:
    if text:
        doc.add_paragraph(text)


def _add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(text)
    try:
        p.style = doc.styles["List Bullet"]
    except KeyError:
        pass


def render_docx(
    resume: Resume, base_template_path: Optional[str] = None
) -> bytes:
    """Build the resume .docx from ``resume`` and return the bytes.

    ``base_template_path`` — optional path to a .docx whose paragraph/character
    styles we want to inherit (fonts, heading colors). Content inside the base
    is ignored; we open it as a Document and append our sections. If omitted,
    python-docx's default style set is used.
    """
    ctx = build_context(resume)

    if base_template_path:
        doc = Document(base_template_path)
        # Clear any pre-existing body paragraphs from the base so we render
        # only what our context produces. Leaves the style table intact.
        body = doc.element.body
        for child in list(body):
            if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
                body.remove(child)
    else:
        doc = Document()

    # ---------- Header ----------
    if ctx["header_name"]:
        p = doc.add_paragraph()
        run = p.add_run(ctx["header_name"])
        run.bold = True
        run.font.size = run.font.size  # keep default
    if ctx["header_title"]:
        doc.add_paragraph(ctx["header_title"])
    if ctx["header_contact"]:
        doc.add_paragraph(ctx["header_contact"])

    # ---------- Summary ----------
    if ctx["has_summary"]:
        _add_heading(doc, "Summary")
        _add_plain(doc, ctx["summary_body"])

    # ---------- Skills ----------
    if ctx["has_skills"]:
        _add_heading(doc, "Skills")
        for line in ctx["skills_lines"]:
            _add_plain(doc, line)

    # ---------- Experience ----------
    if ctx["has_experiences"]:
        _add_heading(doc, "Experience")
        for exp in ctx["experiences"]:
            if exp["header_line"]:
                p = doc.add_paragraph()
                p.add_run(exp["header_line"]).bold = True
            if exp["summary"]:
                _add_plain(doc, exp["summary"])
            for bullet in exp["descriptions"]:
                _add_bullet(doc, bullet)

    # ---------- Certifications ----------
    if ctx["has_certifications"]:
        _add_heading(doc, "Certifications")
        for cert in ctx["certifications"]:
            if cert["header_line"]:
                p = doc.add_paragraph()
                p.add_run(cert["header_line"]).bold = True
            if cert["content"]:
                _add_plain(doc, cert["content"])

    # ---------- Education ----------
    if ctx["has_educations"]:
        _add_heading(doc, "Education")
        for edu in ctx["educations"]:
            _add_plain(doc, edu["line"])

    # ---------- Projects ----------
    if ctx["has_projects"]:
        _add_heading(doc, "Projects")
        for proj in ctx["projects"]:
            if proj["header_line"]:
                p = doc.add_paragraph()
                p.add_run(proj["header_line"]).bold = True
            for bullet in proj["descriptions"]:
                _add_bullet(doc, bullet)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
